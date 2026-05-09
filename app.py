#!/usr/bin/env python3
"""
WordCraft Pro - API 核心层（Web 版业务逻辑）
"""

import json
import logging
import os
import re
import subprocess
import sys
import threading
import time
import uuid
import zipfile
from urllib.error import HTTPError, URLError
from urllib.request import Request, urlopen, urlretrieve
from functools import wraps

logger = logging.getLogger(__name__)

APP_VERSION = "1.0.0"


def cached_response(ttl_seconds: int = 300):
    def decorator(func):
        cache = {}
        cache_time = {}

        @wraps(func)
        def wrapper(self, *args, **kwargs):
            cache_key = f"{func.__name__}_{hash(str(args) + str(kwargs))}"
            now = time.time()
            if cache_key in cache and (now - cache_time.get(cache_key, 0)) < ttl_seconds:
                return cache[cache_key]
            result = func(self, *args, **kwargs)
            cache[cache_key] = result
            cache_time[cache_key] = now
            return result

        return wrapper
    return decorator


class Api:
    def __init__(self, supabase_enabled: bool = True) -> None:
        self._session = {}
        self._supabase = None
        self._supabase_enabled = supabase_enabled
        self._supabase_initialized = False
        self._qa_health: dict = {}
        self._qa_runtime_config: dict = {}
        self._qa_autofix_attempted = False
        self._qa_autofix_attempted_at = 0
        self._qa_last_aggressive_retry_at = 0
        self._local_usage_counters: dict[str, dict] = {}
        self._local_activation_codes: dict[str, dict] = {
            "PRO-LOCAL": {"plan_tier": "pro", "duration_days": 30, "redeemed": False},
        }
        self._local_known_users: dict[str, dict] = {
            "mock-user-001": {"id": "mock-user-001", "email": "test@example.com", "nickname": "test"},
            "team-member-001": {"id": "team-member-001", "email": "member@example.com", "nickname": "member"},
            "team-member-002": {"id": "team-member-002", "email": "editor@example.com", "nickname": "editor"},
        }
        self._local_teams: dict[str, dict] = {}
        self._local_team_rules: dict[str, dict] = {}
        self._local_team_activity_logs: dict[str, list[dict]] = {}
        self._local_team_batch_jobs: dict[str, dict] = {}
        self._qa_install_state: dict = {
            "status": "idle",
            "message": "",
            "updated_at": int(time.time()),
        }
        self._warmup_qa_capabilities(force=True)

    def _load_runtime_config(self) -> dict:
        cfg = {}
        try:
            import yaml
            cfg_path = os.path.join(os.path.dirname(__file__), "config.yaml")
            if os.path.exists(cfg_path):
                with open(cfg_path, "r", encoding="utf-8") as f:
                    cfg = yaml.safe_load(f) or {}
        except Exception as exc:
            logger.warning("读取 config.yaml 失败: %s", exc)
        return cfg

    def _probe_autocorrect(self) -> dict:
        result = {"enabled": False, "available": False, "detail": "", "command": ""}
        qa_cfg = (self._qa_runtime_config.get("qa") or {})
        ac_cfg = qa_cfg.get("autocorrect_check") or {}
        enabled = bool(ac_cfg.get("enabled", True))
        cmd = str(ac_cfg.get("command", "autocorrect"))
        result["enabled"] = enabled
        result["command"] = cmd
        if not enabled:
            result["detail"] = "disabled by config"
            return result
        try:
            proc = subprocess.run(
                [cmd, "--version"],
                capture_output=True,
                text=True,
                encoding="utf-8",
                errors="replace",
                timeout=5,
            )
            if proc.returncode == 0:
                result["available"] = True
                result["detail"] = (proc.stdout or proc.stderr).strip() or "ok"
            else:
                result["detail"] = (proc.stderr or proc.stdout).strip() or f"exit={proc.returncode}"
        except Exception as exc:
            result["detail"] = f"exec failed: {exc}"
        return result

    def _get_autocorrect_target_command(self) -> str:
        """返回项目内推荐 autocorrect 可执行路径。"""
        exe = "autocorrect.exe" if sys.platform == "win32" else "autocorrect"
        return os.path.join(os.path.dirname(__file__), "tools", "autocorrect", exe)

    def _ensure_autocorrect_command(self):
        """若配置缺失 command，自动填充项目内命令路径。"""
        qa_cfg = self._qa_runtime_config.setdefault("qa", {})
        ac_cfg = qa_cfg.setdefault("autocorrect_check", {})
        if not ac_cfg.get("command"):
            ac_cfg["command"] = self._get_autocorrect_target_command()

    def _auto_install_autocorrect(self) -> tuple[bool, str]:
        """自动安装 autocorrect（当前优先 Windows 二进制分发）。"""
        try:
            cmd = self._get_autocorrect_target_command()
            if os.path.isfile(cmd):
                return True, "binary already exists"

            if sys.platform != "win32":
                return False, "auto-install autocorrect currently supports win32 only"

            tools_dir = os.path.dirname(cmd)
            os.makedirs(tools_dir, exist_ok=True)
            zip_path = os.path.join(os.path.dirname(__file__), "tools", "autocorrect-windows-amd64.zip")
            url = "https://github.com/huacnlee/autocorrect/releases/download/v2.16.3/autocorrect-windows-amd64.zip"
            urlretrieve(url, zip_path)
            with zipfile.ZipFile(zip_path, "r") as zf:
                zf.extractall(tools_dir)
            if os.path.isfile(cmd):
                return True, "downloaded autocorrect binary"
            return False, "download finished but binary missing"
        except Exception as exc:
            return False, f"autocorrect install failed: {exc}"

    def _auto_fix_qa_dependencies(self, ac: dict, ignore_cooldown: bool = False) -> dict:
        """autocorrect 依赖自动修复（带冷却重试，避免频繁安装）。"""
        autofix = {"attempted": False, "actions": []}
        qa_cfg = self._qa_runtime_config.get("qa") or {}
        auto_install_enabled = bool(qa_cfg.get("auto_install_dependencies", True))
        retry_interval = int(qa_cfg.get("auto_install_retry_seconds", 300))
        if not auto_install_enabled:
            autofix["actions"].append({"name": "autofix", "ok": False, "detail": "disabled by config"})
            self._qa_install_state = {
                "status": "disabled",
                "message": "auto install disabled",
                "updated_at": int(time.time()),
            }
            return autofix
        if (not ignore_cooldown) and self._qa_autofix_attempted and (time.time() - self._qa_autofix_attempted_at) < retry_interval:
            autofix["actions"].append({"name": "autofix", "ok": False, "detail": "cooldown"})
            self._qa_install_state = {
                "status": "cooldown",
                "message": f"next retry in <= {retry_interval}s",
                "updated_at": int(time.time()),
            }
            return autofix

        self._qa_autofix_attempted = True
        self._qa_autofix_attempted_at = int(time.time())
        autofix["attempted"] = True
        self._qa_install_state = {
            "status": "installing",
            "message": "starting dependency auto-install" if not ignore_cooldown else "retrying dependency auto-install now",
            "updated_at": int(time.time()),
        }

        if ac.get("enabled") and not ac.get("available"):
            self._qa_install_state = {
                "status": "installing",
                "message": "installing autocorrect",
                "updated_at": int(time.time()),
            }
            ok, detail = self._auto_install_autocorrect()
            autofix["actions"].append({"name": "autocorrect", "ok": ok, "detail": detail})
            self._ensure_autocorrect_command()

        self._qa_install_state = {
            "status": "verifying",
            "message": "verifying capabilities",
            "updated_at": int(time.time()),
        }
        return autofix

    def _warmup_qa_capabilities(self, force: bool = False, aggressive_retry: bool = False) -> dict:
        if self._qa_health and not force:
            return self._qa_health

        self._qa_runtime_config = self._load_runtime_config()
        self._ensure_autocorrect_command()
        ac = self._probe_autocorrect()

        # 首次发现缺依赖时尝试自动修复，再重探测一次。
        autofix = self._auto_fix_qa_dependencies(ac, ignore_cooldown=aggressive_retry)
        if autofix.get("attempted"):
            ac = self._probe_autocorrect()

        # 硬门禁：启用的能力必须全部可用。
        required = []
        if ac["enabled"]:
            required.append(("autocorrect", ac["available"]))
        missing = [name for name, ok in required if not ok]

        self._qa_health = {
            "ready": len(missing) == 0,
            "checked_at": int(time.time()),
            "missing": missing,
            "autofix": autofix,
            "install_state": {
                **self._qa_install_state,
                "completed": len(missing) == 0,
            },
            "capabilities": {
                "autocorrect": ac,
            },
        }
        if len(missing) == 0:
            self._qa_install_state = {
                "status": "completed",
                "message": "qa dependencies ready",
                "updated_at": int(time.time()),
            }
            self._qa_health["install_state"] = {
                **self._qa_install_state,
                "completed": True,
            }
        elif self._qa_install_state.get("status") in ("idle", "verifying", "completed"):
            self._qa_install_state = {
                "status": "pending",
                "message": "waiting for dependencies",
                "updated_at": int(time.time()),
            }
            self._qa_health["install_state"] = {
                **self._qa_install_state,
                "completed": False,
            }
        return self._qa_health

    def getQAHealth(self) -> str:
        health = self._warmup_qa_capabilities(force=True)
        # 纯网页场景下，前端主要靠 health 轮询；这里按节流自动触发一次“立即重试安装”。
        if not health.get("ready", False):
            now_ts = int(time.time())
            qa_cfg = self._qa_runtime_config.get("qa") or {}
            aggressive_interval = int(qa_cfg.get("auto_install_aggressive_retry_seconds", 20))
            if (now_ts - self._qa_last_aggressive_retry_at) >= aggressive_interval:
                self._qa_last_aggressive_retry_at = now_ts
                health = self._warmup_qa_capabilities(force=True, aggressive_retry=True)
        return json.dumps({"success": True, **health}, ensure_ascii=False)

    def _ensure_supabase_initialized(self):
        if not self._supabase_initialized and self._supabase_enabled:
            self._init_supabase()
            self._supabase_initialized = True

    def _init_supabase(self):
        try:
            from core.supabase_client import SupabaseClient
            self._supabase = SupabaseClient()
        except Exception as e:
            logger.warning("Supabase 初始化失败，使用本地模式: %s", e)
            self._supabase = None

    # ------------------------------------------------------------------
    #  文件操作
    # ------------------------------------------------------------------

    def _local_user_id(self) -> str:
        return (self._session or {}).get("user_id") or "mock-user-001"

    def _local_profile(self) -> dict:
        user = dict((self._session or {}).get("user_info") or {})
        user.setdefault("id", self._local_user_id())
        user.setdefault("plan_tier", "free")
        user.setdefault("plan_status", "active")
        user.setdefault("plan_source", "local")
        user.setdefault("team_id", None)
        user.setdefault("feature_flags", {})
        return user

    def _get_usage_counter(self, user_id: str | None = None) -> dict:
        from core.entitlements import current_day_key, current_period_key

        user_id = user_id or self._local_user_id()
        period_key = current_period_key()
        day_key = current_day_key()
        if self._supabase and user_id:
            return self._supabase.get_or_create_usage_counter(user_id, period_key, day_key)
        key = f"{user_id}:{period_key}"
        row = self._local_usage_counters.setdefault(key, {
            "user_id": user_id,
            "period_key": period_key,
            "day_key": day_key,
            "ai_qa_used": 0,
            "ai_parse_used": 0,
            "rule_check_used_today": 0,
        })
        if row.get("day_key") != day_key:
            row["day_key"] = day_key
            row["rule_check_used_today"] = 0
        return row

    @staticmethod
    def _normalize_ai_usage_payload(usage) -> dict:
        if not usage:
            return {}
        if isinstance(usage, dict):
            return dict(usage)
        if hasattr(usage, "model_dump"):
            try:
                return dict(usage.model_dump())
            except Exception:
                pass
        if hasattr(usage, "dict"):
            try:
                return dict(usage.dict())
            except Exception:
                pass
        normalized = {}
        for key in ("total_tokens", "prompt_tokens", "completion_tokens"):
            value = getattr(usage, key, None)
            if value is not None:
                normalized[key] = value
        return normalized

    def _record_ai_token_usage(self, purpose: str, model: str, usage) -> None:
        usage_data = self._normalize_ai_usage_payload(usage)
        try:
            amount = int(usage_data.get("total_tokens") or 0)
        except (TypeError, ValueError):
            amount = 0
        if amount <= 0:
            return

        try:
            prompt_tokens = int(usage_data.get("prompt_tokens") or 0)
        except (TypeError, ValueError):
            prompt_tokens = 0
        try:
            completion_tokens = int(usage_data.get("completion_tokens") or 0)
        except (TypeError, ValueError):
            completion_tokens = 0

        try:
            self._ensure_supabase_initialized()
            user_id = self._session.get("user_id") if self._session else None
            access_token = self._session.get("access_token") if self._session else None
            if self._supabase and user_id:
                self._supabase.record_token_usage(
                    user_id,
                    amount,
                    purpose,
                    model=model,
                    prompt_tokens=prompt_tokens,
                    completion_tokens=completion_tokens,
                    access_token=access_token,
                )
                return
            from core.token_tracker import TokenTracker
            TokenTracker().record_usage(
                amount,
                purpose=purpose,
                model=model,
                prompt_tokens=prompt_tokens,
                completion_tokens=completion_tokens,
            )
        except Exception as exc:
            logger.warning("Failed to persist AI token usage: %s", exc)

    def _increment_usage(self, counter_name: str, amount: int = 1) -> dict:
        user_id = self._local_user_id()
        if self._supabase and self._session.get("user_id"):
            return self._supabase.increment_usage_counter(
                user_id,
                counter_name,
                amount,
                access_token=(self._session or {}).get("access_token"),
            )
        row = self._get_usage_counter(user_id)
        row[counter_name] = int(row.get(counter_name) or 0) + int(amount)
        return row

    def _current_entitlements(self) -> dict:
        from core.entitlements import build_entitlements

        self._ensure_supabase_initialized()
        uid = self._session.get("user_id") if self._session else None
        if self._supabase and uid:
            access_token = (self._session or {}).get("access_token")
            return self._supabase.get_user_entitlements(uid, access_token=access_token)
        return build_entitlements(self._local_profile(), self._get_usage_counter())

    def _current_team_id(self) -> str | None:
        user_info = (self._session or {}).get("user_info") or {}
        team_id = user_info.get("team_id")
        return str(team_id).strip() if team_id else None

    def _current_user_id(self) -> str:
        return (self._session or {}).get("user_id") or self._local_user_id()

    def _team_workspace_payload(self, team: dict | None = None, members: list | None = None, rules: dict | None = None) -> dict:
        team_id = (team or {}).get("id") if isinstance(team, dict) else self._current_team_id()
        return {
            "team": team,
            "members": members or [],
            "rules": rules,
            "invitations": [],
            "team_id": team_id,
            "plan": self._current_entitlements(),
            "activities": self._list_team_activities(team_id, limit=20) if team_id else [],
            "jobs": self._list_team_batch_jobs(team_id, limit=20) if team_id else [],
        }

    @staticmethod
    def _team_gate(reason_code: str, message: str) -> dict:
        return {
            "allowed": False,
            "reason_code": reason_code,
            "feature": "team_rule_share",
            "upgrade_target": "team",
            "message": message,
        }

    def _current_local_team(self) -> dict | None:
        team_id = self._current_team_id()
        return self._local_teams.get(team_id) if team_id else None

    def _is_local_team_owner(self, team: dict | None) -> bool:
        return bool(team) and team.get("owner_user_id") == self._current_user_id()

    def _ensure_local_known_user(self, email: str) -> dict:
        normalized = (email or "").strip().lower()
        for user in self._local_known_users.values():
            if str(user.get("email", "")).strip().lower() == normalized:
                return user
        user_id = f"local-user-{len(self._local_known_users) + 1:03d}"
        user = {
            "id": user_id,
            "email": normalized,
            "nickname": normalized.split("@")[0] if "@" in normalized else normalized,
        }
        self._local_known_users[user_id] = user
        return user

    def _local_pending_team_invitations(self, user_id: str) -> list[dict]:
        invitations: list[dict] = []
        for team in self._local_teams.values():
            for member in team.get("members", []):
                if member.get("user_id") == user_id and member.get("status") == "pending":
                    invitations.append({**member, "team": {k: v for k, v in team.items() if k != "members"}})
        return invitations

    def _can_access_team_workspace(self) -> bool:
        gate = self._check_feature_gate("team_rule_share")
        if gate.get("allowed"):
            return True
        uid = self._current_user_id()
        if self._supabase and uid:
            workspace = self._supabase.get_user_team_workspace(uid)
            return bool(workspace.get("invitations"))
        return bool(self._local_pending_team_invitations(uid))

    def _resolve_session_role(self, access_token: str | None) -> str | None:
        if not (self._supabase and access_token):
            return None
        try:
            auth_user = self._supabase.client.auth.get_user(access_token)
            user = getattr(auth_user, "user", None)
            app_metadata = getattr(user, "app_metadata", None) or getattr(user, "raw_app_meta_data", None) or {}
            if isinstance(app_metadata, dict):
                role = app_metadata.get("role")
                return str(role).strip() if role else None
        except Exception as exc:
            logger.warning("failed to resolve auth role from session: %s", exc)
        return None

    @staticmethod
    def _now_iso() -> str:
        return time.strftime("%Y-%m-%dT%H:%M:%SZ", time.gmtime())

    def _record_team_activity(
        self,
        team_id: str | None,
        event_type: str,
        status: str = "completed",
        summary: str = "",
        payload: dict | None = None,
        target_email: str | None = None,
        target_user_id: str | None = None,
        batch_job_id: str | None = None,
    ) -> dict | None:
        normalized_team_id = str(team_id or "").strip()
        if not normalized_team_id:
            return None
        actor_user_id = self._current_user_id()
        row = {
            "id": f"activity-{uuid.uuid4().hex[:12]}",
            "team_id": normalized_team_id,
            "actor_user_id": actor_user_id,
            "event_type": event_type,
            "status": status,
            "summary": summary,
            "payload": payload or {},
            "target_email": target_email,
            "target_user_id": target_user_id,
            "batch_job_id": batch_job_id,
            "created_at": self._now_iso(),
        }
        if self._supabase:
            try:
                return self._supabase.create_team_activity(row)
            except Exception as exc:
                logger.warning("记录团队活动失败，回退本地日志: %s", exc)
        bucket = self._local_team_activity_logs.setdefault(normalized_team_id, [])
        bucket.insert(0, row)
        del bucket[100:]
        return row

    def _list_team_activities(self, team_id: str | None, limit: int = 20) -> list[dict]:
        normalized_team_id = str(team_id or "").strip()
        if not normalized_team_id:
            return []
        if self._supabase:
            try:
                return self._supabase.list_team_activities(normalized_team_id, limit=limit)
            except Exception as exc:
                logger.warning("读取团队活动失败，回退本地日志: %s", exc)
        return list(self._local_team_activity_logs.get(normalized_team_id, [])[:limit])

    def _create_team_batch_job(self, team_id: str, files: list[dict], categories: list[str]) -> dict:
        normalized_team_id = str(team_id or "").strip()
        actor_user_id = self._current_user_id()
        row = {
            "id": f"job-{uuid.uuid4().hex[:12]}",
            "team_id": normalized_team_id,
            "created_by": actor_user_id,
            "job_type": "batch_qa",
            "status": "queued",
            "file_count": len(files or []),
            "categories": categories or [],
            "request_payload": {
                "files": [{"name": item.get("name", "untitled")} for item in (files or [])],
            },
            "result_payload": None,
            "summary": f"等待批量检查 {len(files or [])} 份文档",
            "created_at": self._now_iso(),
            "updated_at": self._now_iso(),
            "started_at": None,
            "finished_at": None,
        }
        if self._supabase:
            try:
                row = self._supabase.create_team_batch_job(row)
            except Exception as exc:
                logger.warning("创建团队任务失败，回退本地: %s", exc)
        self._local_team_batch_jobs[row["id"]] = row
        self._record_team_activity(
            normalized_team_id,
            "batch_qa_created",
            status="queued",
            summary=f"已创建批量检查任务，文档数 {len(files or [])}",
            payload={"job_id": row["id"], "file_count": len(files or []), "categories": categories or []},
            batch_job_id=row["id"],
        )
        return row

    def _update_team_batch_job(self, job_id: str, **updates) -> dict | None:
        normalized_job_id = str(job_id or "").strip()
        if not normalized_job_id:
            return None
        row = dict(self._local_team_batch_jobs.get(normalized_job_id) or {})
        if not row and self._supabase:
            row = self._supabase.get_team_batch_job(normalized_job_id) or {}
        if not row:
            return None
        row.update({k: v for k, v in updates.items() if v is not None})
        row["updated_at"] = self._now_iso()
        self._local_team_batch_jobs[normalized_job_id] = row
        if self._supabase:
            try:
                self._supabase.update_team_batch_job(normalized_job_id, row)
            except Exception as exc:
                logger.warning("更新团队任务失败: %s", exc)
        return row

    def _list_team_batch_jobs(self, team_id: str | None, limit: int = 20) -> list[dict]:
        normalized_team_id = str(team_id or "").strip()
        if not normalized_team_id:
            return []
        if self._supabase:
            try:
                rows = self._supabase.list_team_batch_jobs(normalized_team_id, limit=limit)
                for row in rows:
                    self._local_team_batch_jobs[row["id"]] = row
                return rows
            except Exception as exc:
                logger.warning("读取团队任务失败，回退本地: %s", exc)
        rows = [row for row in self._local_team_batch_jobs.values() if row.get("team_id") == normalized_team_id]
        rows.sort(key=lambda item: item.get("created_at") or "", reverse=True)
        return rows[:limit]

    def _build_team_invite_email_payload(self, team: dict, invite_email: str, role: str = "member") -> dict:
        team_name = team.get("name") or "WordCraft 团队"
        role_label = "owner" if str(role or "").strip().lower() == "owner" else "member"
        subject = f"邀请加入 {team_name}"
        text = "\n".join([
            f"你已被邀请加入 {team_name}。",
            f"角色：{role_label}",
            f"受邀邮箱：{invite_email}",
            "请登录 WordCraft，在团队工作区中接受邀请。",
        ])
        html = (
            f"<p>你已被邀请加入 <strong>{team_name}</strong>。</p>"
            f"<p>角色：<strong>{role_label}</strong><br>受邀邮箱：<strong>{invite_email}</strong></p>"
            "<p>请登录 WordCraft，在团队工作区中接受邀请。</p>"
        )
        return {"subject": subject, "text": text, "html": html}

    def _deliver_team_invite_email(self, invite_email: str, payload: dict) -> dict:
        webhook_url = os.environ.get("TEAM_INVITE_EMAIL_WEBHOOK_URL", "").strip()
        resend_key = os.environ.get("RESEND_API_KEY", "").strip()
        from_email = os.environ.get("TEAM_INVITE_FROM_EMAIL", "").strip() or os.environ.get("WC_EMAIL_FROM", "").strip()
        body = {
            "to": invite_email,
            "subject": payload.get("subject", "WordCraft 团队邀请"),
            "text": payload.get("text", ""),
            "html": payload.get("html", ""),
        }
        try:
            if webhook_url:
                req = Request(
                    webhook_url,
                    data=json.dumps(body).encode("utf-8"),
                    headers={"Content-Type": "application/json"},
                    method="POST",
                )
                with urlopen(req, timeout=15) as resp:
                    raw = resp.read().decode("utf-8", errors="replace")
                return {"success": True, "provider": "webhook", "response": raw}
            if resend_key and from_email:
                resend_body = {
                    "from": from_email,
                    "to": [invite_email],
                    "subject": body["subject"],
                    "html": body["html"],
                    "text": body["text"],
                }
                req = Request(
                    "https://api.resend.com/emails",
                    data=json.dumps(resend_body).encode("utf-8"),
                    headers={
                        "Content-Type": "application/json",
                        "Authorization": f"Bearer {resend_key}",
                    },
                    method="POST",
                )
                with urlopen(req, timeout=15) as resp:
                    raw = resp.read().decode("utf-8", errors="replace")
                parsed = json.loads(raw) if raw else {}
                return {"success": True, "provider": "resend", "response": parsed}
            return {
                "success": False,
                "error_code": "TEAM_EMAIL_NOT_CONFIGURED",
                "error": "正式邀请邮件通道未配置，请设置 TEAM_INVITE_EMAIL_WEBHOOK_URL 或 RESEND_API_KEY + TEAM_INVITE_FROM_EMAIL。",
            }
        except HTTPError as exc:
            detail = exc.read().decode("utf-8", errors="replace") if hasattr(exc, "read") else str(exc)
            return {"success": False, "error_code": "TEAM_EMAIL_DELIVERY_FAILED", "error": detail or str(exc)}
        except URLError as exc:
            return {"success": False, "error_code": "TEAM_EMAIL_DELIVERY_FAILED", "error": str(exc)}
        except Exception as exc:
            return {"success": False, "error_code": "TEAM_EMAIL_DELIVERY_FAILED", "error": str(exc)}

    def _run_team_batch_job(self, job_id: str, team_id: str, files: list[dict], categories: list[str]) -> None:
        started = self._now_iso()
        self._update_team_batch_job(job_id, status="running", summary=f"正在批量检查 {len(files)} 份文档", started_at=started)
        self._record_team_activity(
            team_id,
            "batch_qa_started",
            status="running",
            summary=f"开始批量检查 {len(files)} 份文档",
            payload={"job_id": job_id, "file_count": len(files)},
            batch_job_id=job_id,
        )
        try:
            results = []
            for item in files:
                content = (item or {}).get("content", "")
                name = (item or {}).get("name") or "untitled"
                elements = (item or {}).get("elements")
                result = self.runQA(
                    content,
                    json.dumps(categories or ["typo", "consistency", "logic", "format", "crossref"], ensure_ascii=False),
                    elements_json=json.dumps(elements, ensure_ascii=False) if elements is not None else None,
                )
                results.append({
                    "name": name,
                    "result": json.loads(result) if isinstance(result, str) else result,
                })
            issues = 0
            for item in results:
                issues += len((item.get("result") or {}).get("issues") or [])
            finished = self._now_iso()
            self._update_team_batch_job(
                job_id,
                status="completed",
                finished_at=finished,
                summary=f"批量检查完成，共 {len(results)} 份文档，发现 {issues} 条问题",
                result_payload={"results": results, "count": len(results), "issues": issues},
            )
            self._record_team_activity(
                team_id,
                "batch_qa_completed",
                status="completed",
                summary=f"批量检查完成，共 {len(results)} 份文档，发现 {issues} 条问题",
                payload={"job_id": job_id, "count": len(results), "issues": issues},
                batch_job_id=job_id,
            )
        except Exception as exc:
            finished = self._now_iso()
            self._update_team_batch_job(
                job_id,
                status="failed",
                finished_at=finished,
                summary=f"批量检查失败：{exc}",
                result_payload={"error": str(exc)},
            )
            self._record_team_activity(
                team_id,
                "batch_qa_failed",
                status="failed",
                summary=f"批量检查失败：{exc}",
                payload={"job_id": job_id, "error": str(exc)},
                batch_job_id=job_id,
            )

    def sync_session_from_access_token(self, access_token: str | None) -> None:
        if not access_token:
            return
        self._ensure_supabase_initialized()
        if not self._supabase:
            return
        try:
            auth_user = self._supabase.client.auth.get_user(access_token)
            user = getattr(auth_user, "user", None)
            if not user or not getattr(user, "id", None):
                return
            profile = self._supabase.get_user_plan(user.id, access_token=access_token)
            user_info = {
                "id": user.id,
                "email": getattr(user, "email", None),
                "nickname": (
                    (getattr(user, "user_metadata", None) or {}).get("nickname")
                    or (getattr(user, "email", "") or "").split("@")[0]
                ),
                "plan_tier": profile.get("plan_tier", "free"),
                "plan_status": profile.get("plan_status", "active"),
                "plan_source": profile.get("plan_source", "system"),
                "current_period_start": profile.get("current_period_start"),
                "current_period_end": profile.get("current_period_end"),
                "team_id": profile.get("team_id"),
                "feature_flags": profile.get("feature_flags", {}) if profile else {},
            }
            app_role = self._resolve_session_role(access_token)
            if app_role:
                user_info["role"] = app_role
            self._session = {
                "user_id": user.id,
                "access_token": access_token,
                "user_info": user_info,
            }
        except Exception as exc:
            logger.warning("failed to sync session from access token: %s", exc)

    @staticmethod
    def _deny_json(gate: dict) -> str:
        return json.dumps({
            "success": False,
            "error_code": gate.get("reason_code", "FEATURE_LOCKED"),
            "error": gate.get("message", "当前套餐不可使用该功能"),
            "feature": gate.get("feature"),
            "upgrade_target": gate.get("upgrade_target"),
            "entitlement_denied": True,
        }, ensure_ascii=False)

    def _check_feature_gate(self, feature: str, quota: str | None = None) -> dict:
        from core.entitlements import check_feature_access, check_quota_available

        entitlements = self._current_entitlements()
        gate = check_feature_access(entitlements, feature)
        if not gate.get("allowed"):
            return gate
        if quota:
            return check_quota_available(entitlements, quota)
        return gate

    def getCurrentPlan(self) -> str:
        try:
            return json.dumps({"success": True, "plan": self._current_entitlements()}, ensure_ascii=False)
        except Exception as exc:
            return json.dumps({"success": False, "error": str(exc)}, ensure_ascii=False)

    def getUsageAndPlan(self) -> str:
        return self.getCurrentPlan()

    def getTeamWorkspace(self) -> str:
        try:
            uid = self._session.get("user_id") if self._session else None
            gate = self._check_feature_gate("team_rule_share")
            if not gate.get("allowed"):
                if not self._can_access_team_workspace():
                    return self._deny_json(gate)
            if self._supabase and uid:
                workspace = self._supabase.get_user_team_workspace(uid)
                workspace.setdefault("plan", self._current_entitlements())
                return json.dumps({"success": True, **workspace}, ensure_ascii=False)

            team_id = self._current_team_id()
            invitations = self._local_pending_team_invitations(self._current_user_id())
            if not team_id:
                payload = self._team_workspace_payload()
                payload["invitations"] = invitations
                return json.dumps({"success": True, **payload}, ensure_ascii=False)
            team = self._local_teams.get(team_id)
            rules = self._local_team_rules.get(team_id)
            payload = self._team_workspace_payload(team, team.get("members") if team else [], rules)
            payload["invitations"] = invitations
            return json.dumps({"success": True, **payload}, ensure_ascii=False)
        except Exception as exc:
            return json.dumps({"success": False, "error": str(exc)}, ensure_ascii=False)

    def createTeamWorkspace(self, name: str, seat_limit: int = 5) -> str:
        try:
            gate = self._check_feature_gate("team_rule_share")
            if not gate.get("allowed"):
                return self._deny_json(gate)

            team_name = (name or "").strip()
            if not team_name:
                return json.dumps({"success": False, "error": "团队名称不能为空", "error_code": "TEAM_NAME_REQUIRED"}, ensure_ascii=False)

            seat_limit = max(int(seat_limit or 5), 1)
            uid = self._session.get("user_id") if self._session else self._local_user_id()
            if self._supabase and uid:
                result = self._supabase.create_team_workspace(uid, team_name, seat_limit)
                if result.get("success") and result.get("team"):
                    self._session.setdefault("user_info", {})["team_id"] = result["team"].get("id")
                return json.dumps(result, ensure_ascii=False)

            team_id = f"team-local-{len(self._local_teams) + 1:03d}"
            user_info = self._session.setdefault("user_info", self._local_profile())
            member = {
                "team_id": team_id,
                "user_id": uid,
                "role": "owner",
                "status": "active",
                "joined_at": time.strftime("%Y-%m-%dT%H:%M:%SZ", time.gmtime()),
                "accepted_at": time.strftime("%Y-%m-%dT%H:%M:%SZ", time.gmtime()),
            }
            team = {
                "id": team_id,
                "name": team_name,
                "owner_user_id": uid,
                "seat_limit": seat_limit,
                "status": "active",
                "created_at": member["joined_at"],
                "members": [member],
            }
            self._local_teams[team_id] = team
            user_info["team_id"] = team_id
            self._session["user_id"] = uid
            return json.dumps({"success": True, **self._team_workspace_payload(team, team["members"], self._local_team_rules.get(team_id))}, ensure_ascii=False)
        except Exception as exc:
            return json.dumps({"success": False, "error": str(exc)}, ensure_ascii=False)

    def addTeamMemberByEmail(self, email: str, role: str = "member") -> str:
        try:
            gate = self._check_feature_gate("team_rule_share")
            if not gate.get("allowed"):
                return self._deny_json(gate)

            team_id = self._current_team_id()
            if not team_id:
                return self._deny_json(self._team_gate("TEAM_REQUIRED", "当前账号尚未加入团队，无法添加成员。"))

            normalized_email = (email or "").strip().lower()
            if not normalized_email:
                return json.dumps({"success": False, "error": "成员邮箱不能为空", "error_code": "MEMBER_EMAIL_REQUIRED"}, ensure_ascii=False)
            member_role = "owner" if str(role or "").strip().lower() == "owner" else "member"
            uid = self._current_user_id()

            if self._supabase and uid:
                result = self._supabase.add_team_member_by_email(uid, team_id, normalized_email, member_role)
                if result.get("success"):
                    self._record_team_activity(
                        team_id,
                        "team_invite_created",
                        status="pending",
                        summary=f"已创建待接受邀请：{normalized_email}",
                        payload={"role": member_role},
                        target_email=normalized_email,
                        target_user_id=(result.get("added_member") or {}).get("user_id"),
                    )
                    workspace = self._supabase.get_user_team_workspace(uid)
                    workspace.setdefault("plan", self._current_entitlements())
                    return json.dumps({"success": True, "added_member": result.get("added_member"), **workspace}, ensure_ascii=False)
                return json.dumps(result, ensure_ascii=False)

            team = self._current_local_team()
            if not team:
                return self._deny_json(self._team_gate("TEAM_REQUIRED", "当前账号尚未加入团队，无法添加成员。"))
            if not self._is_local_team_owner(team):
                return self._deny_json(self._team_gate("TEAM_OWNER_REQUIRED", "只有团队所有者可以录入成员。"))

            user = self._ensure_local_known_user(normalized_email)
            members = team.setdefault("members", [])
            if any(m.get("user_id") == user["id"] for m in members):
                return json.dumps({"success": False, "error": "该成员已在团队中", "error_code": "TEAM_MEMBER_EXISTS"}, ensure_ascii=False)
            if len(members) >= int(team.get("seat_limit") or 0):
                return json.dumps({"success": False, "error": "团队席位已满", "error_code": "TEAM_SEAT_LIMIT_REACHED"}, ensure_ascii=False)

            member = {
                "team_id": team_id,
                "user_id": user["id"],
                "email": user["email"],
                "role": member_role,
                "status": "pending",
                "invite_email": user["email"],
                "invited_by": uid,
                "invited_at": time.strftime("%Y-%m-%dT%H:%M:%SZ", time.gmtime()),
                "joined_at": None,
                "accepted_at": None,
            }
            members.append(member)
            self._record_team_activity(
                team_id,
                "team_invite_created",
                status="pending",
                summary=f"已创建待接受邀请：{normalized_email}",
                payload={"role": member_role},
                target_email=normalized_email,
                target_user_id=user["id"],
            )
            return json.dumps({
                "success": True,
                "added_member": member,
                **self._team_workspace_payload(team, members, self._local_team_rules.get(team_id)),
            }, ensure_ascii=False)
        except Exception as exc:
            return json.dumps({"success": False, "error": str(exc)}, ensure_ascii=False)

    def acceptTeamInvite(self, team_id: str) -> str:
        try:
            normalized_team_id = str(team_id or "").strip()
            if not normalized_team_id:
                return json.dumps({"success": False, "error": "团队邀请不存在", "error_code": "TEAM_INVITE_NOT_FOUND"}, ensure_ascii=False)

            uid = self._current_user_id()
            if self._supabase and uid:
                result = self._supabase.accept_team_invite(uid, normalized_team_id)
                if result.get("success"):
                    self._session.setdefault("user_info", {})["team_id"] = normalized_team_id
                    self._session["user_info"]["plan_tier"] = "team"
                    result["plan"] = self._current_entitlements()
                    self._record_team_activity(
                        normalized_team_id,
                        "team_invite_accepted",
                        status="completed",
                        summary="成员已接受团队邀请",
                        target_user_id=uid,
                    )
                return json.dumps(result, ensure_ascii=False)

            team = self._local_teams.get(normalized_team_id)
            if not team:
                return json.dumps({"success": False, "error": "团队不存在", "error_code": "TEAM_NOT_FOUND"}, ensure_ascii=False)
            members = team.get("members", [])
            invite = next((m for m in members if m.get("user_id") == uid and m.get("status") == "pending"), None)
            if not invite:
                return json.dumps({"success": False, "error": "待接受邀请不存在", "error_code": "TEAM_INVITE_NOT_FOUND"}, ensure_ascii=False)

            now = time.strftime("%Y-%m-%dT%H:%M:%SZ", time.gmtime())
            invite["status"] = "active"
            invite["accepted_at"] = now
            invite["joined_at"] = invite.get("joined_at") or now
            user_info = self._session.setdefault("user_info", self._local_profile())
            user_info["team_id"] = normalized_team_id
            user_info["plan_tier"] = "team"
            payload = self._team_workspace_payload(team, members, self._local_team_rules.get(normalized_team_id))
            payload["invitations"] = self._local_pending_team_invitations(uid)
            self._record_team_activity(
                normalized_team_id,
                "team_invite_accepted",
                status="completed",
                summary="成员已接受团队邀请",
                target_user_id=uid,
            )
            return json.dumps({"success": True, **payload}, ensure_ascii=False)
        except Exception as exc:
            return json.dumps({"success": False, "error": str(exc)}, ensure_ascii=False)

    def cancelTeamInvite(self, team_id: str, email: str) -> str:
        try:
            normalized_team_id = str(team_id or "").strip()
            normalized_email = str(email or "").strip().lower()
            if not normalized_team_id:
                return json.dumps({"success": False, "error": "Team invite was not found.", "error_code": "TEAM_INVITE_NOT_FOUND"}, ensure_ascii=False)
            if not normalized_email:
                return json.dumps({"success": False, "error": "Member email is required.", "error_code": "MEMBER_EMAIL_REQUIRED"}, ensure_ascii=False)

            uid = self._current_user_id()
            if self._supabase and uid:
                result = self._supabase.cancel_team_invite(uid, normalized_team_id, normalized_email)
                if result.get("success"):
                    self._record_team_activity(
                        normalized_team_id,
                        "team_invite_canceled",
                        status="completed",
                        summary=f"已撤销邀请：{normalized_email}",
                        target_email=normalized_email,
                    )
                    workspace = self._supabase.get_user_team_workspace(uid)
                    workspace.setdefault("plan", self._current_entitlements())
                    return json.dumps({"success": True, **workspace}, ensure_ascii=False)
                return json.dumps(result, ensure_ascii=False)

            team = self._local_teams.get(normalized_team_id)
            if not team:
                return json.dumps({"success": False, "error": "Team was not found.", "error_code": "TEAM_NOT_FOUND"}, ensure_ascii=False)
            if not self._is_local_team_owner(team):
                return self._deny_json(self._team_gate("TEAM_OWNER_REQUIRED", "Only the team owner can cancel pending invites."))

            members = team.get("members", [])
            pending_index = next((
                idx for idx, member in enumerate(members)
                if member.get("status") == "pending"
                and str(member.get("invite_email") or member.get("email") or "").strip().lower() == normalized_email
            ), None)
            if pending_index is None:
                return json.dumps({"success": False, "error": "Pending invite was not found.", "error_code": "TEAM_INVITE_NOT_FOUND"}, ensure_ascii=False)

            members.pop(pending_index)
            self._record_team_activity(
                normalized_team_id,
                "team_invite_canceled",
                status="completed",
                summary=f"已撤销邀请：{normalized_email}",
                target_email=normalized_email,
            )
            return json.dumps({
                "success": True,
                **self._team_workspace_payload(team, members, self._local_team_rules.get(normalized_team_id)),
            }, ensure_ascii=False)
        except Exception as exc:
            return json.dumps({"success": False, "error": str(exc)}, ensure_ascii=False)

    def sendTeamInviteEmail(self, team_id: str, email: str, role: str = "member") -> str:
        try:
            gate = self._check_feature_gate("team_rule_share")
            if not gate.get("allowed"):
                return self._deny_json(gate)

            normalized_team_id = str(team_id or "").strip()
            normalized_email = str(email or "").strip().lower()
            if not normalized_team_id:
                return json.dumps({"success": False, "error": "团队不存在", "error_code": "TEAM_NOT_FOUND"}, ensure_ascii=False)
            if not normalized_email:
                return json.dumps({"success": False, "error": "成员邮箱不能为空", "error_code": "MEMBER_EMAIL_REQUIRED"}, ensure_ascii=False)

            if self._supabase:
                workspace = self._supabase.get_team_workspace_by_id(normalized_team_id)
            else:
                team = self._local_teams.get(normalized_team_id)
                workspace = self._team_workspace_payload(team, team.get("members") if team else [], self._local_team_rules.get(normalized_team_id)) if team else {}
            team = (workspace or {}).get("team") or {}
            members = (workspace or {}).get("members") or []
            if not team:
                return json.dumps({"success": False, "error": "团队不存在", "error_code": "TEAM_NOT_FOUND"}, ensure_ascii=False)
            pending = next((
                member for member in members
                if member.get("status") == "pending"
                and str(member.get("invite_email") or member.get("email") or "").strip().lower() == normalized_email
            ), None)
            if not pending:
                return json.dumps({"success": False, "error": "待发送的邀请不存在", "error_code": "TEAM_INVITE_NOT_FOUND"}, ensure_ascii=False)

            payload = self._build_team_invite_email_payload(team, normalized_email, role or pending.get("role") or "member")
            delivery = self._deliver_team_invite_email(normalized_email, payload)
            activity_status = "completed" if delivery.get("success") else "failed"
            self._record_team_activity(
                normalized_team_id,
                "team_invite_email_sent",
                status=activity_status,
                summary=(
                    f"正式邀请邮件已发送：{normalized_email}"
                    if delivery.get("success")
                    else f"正式邀请邮件发送失败：{normalized_email}"
                ),
                payload={"provider": delivery.get("provider"), "error": delivery.get("error")},
                target_email=normalized_email,
                target_user_id=pending.get("user_id"),
            )
            if delivery.get("success"):
                refreshed = self.getTeamWorkspace()
                data = json.loads(refreshed) if isinstance(refreshed, str) else refreshed
                return json.dumps({
                    "success": True,
                    "delivery": delivery,
                    **(data or {}),
                }, ensure_ascii=False)
            return json.dumps({"success": False, **delivery}, ensure_ascii=False)
        except Exception as exc:
            return json.dumps({"success": False, "error": str(exc)}, ensure_ascii=False)

    def startTeamBatchQA(self, files_json: str, categories_str: str = '["typo", "consistency", "logic", "format", "crossref"]') -> str:
        try:
            gate = self._check_feature_gate("batch_check")
            if not gate.get("allowed"):
                return self._deny_json(gate)
            team_id = self._current_team_id()
            if not team_id:
                return self._deny_json(self._team_gate("TEAM_REQUIRED", "当前账号尚未加入团队，无法创建批量任务。"))
            files = json.loads(files_json) if isinstance(files_json, str) else (files_json or [])
            categories = json.loads(categories_str) if isinstance(categories_str, str) else (categories_str or [])
            if not files:
                return json.dumps({"success": False, "error": "当前没有可批量检查的文档", "error_code": "EMPTY_BATCH_FILES"}, ensure_ascii=False)
            job = self._create_team_batch_job(team_id, files, categories)
            worker = threading.Thread(
                target=self._run_team_batch_job,
                args=(job["id"], team_id, files, categories),
                daemon=True,
            )
            worker.start()
            return json.dumps({"success": True, "job": job, "jobs": self._list_team_batch_jobs(team_id, limit=20)}, ensure_ascii=False)
        except Exception as exc:
            return json.dumps({"success": False, "error": str(exc)}, ensure_ascii=False)

    def getTeamBatchJobs(self) -> str:
        try:
            gate = self._check_feature_gate("batch_check")
            if not gate.get("allowed"):
                return self._deny_json(gate)
            team_id = self._current_team_id()
            jobs = self._list_team_batch_jobs(team_id, limit=20) if team_id else []
            return json.dumps({"success": True, "jobs": jobs, "team_id": team_id}, ensure_ascii=False)
        except Exception as exc:
            return json.dumps({"success": False, "error": str(exc)}, ensure_ascii=False)

    def redeemActivationCode(self, code: str) -> str:
        try:
            self._ensure_supabase_initialized()
            normalized_code = (code or "").strip().upper()
            uid = self._session.get("user_id") if self._session else None
            if self._supabase and uid:
                result = self._supabase.redeem_activation_code(uid, normalized_code)
                if result.get("success"):
                    profile = self._supabase.get_user_plan(uid)
                    self._session.setdefault("user_info", {}).update(profile)
                    result["plan"] = self._current_entitlements()
                return json.dumps(result, ensure_ascii=False)

            code_key = normalized_code
            local_code = self._local_activation_codes.get(code_key)
            if not local_code or local_code.get("redeemed"):
                return json.dumps({"success": False, "error": "激活码无效或已使用", "error_code": "CODE_INVALID"}, ensure_ascii=False)
            local_code["redeemed"] = True
            user_info = self._session.setdefault("user_info", self._local_profile())
            user_info.update({
                "plan_tier": local_code.get("plan_tier", "pro"),
                "plan_status": "active",
                "plan_source": "activation_code",
            })
            self._session["user_id"] = user_info.get("id", self._local_user_id())
            return json.dumps({"success": True, "plan": self._current_entitlements()}, ensure_ascii=False)
        except Exception as exc:
            return json.dumps({"success": False, "error": str(exc)}, ensure_ascii=False)

    def openFile(self, file_content: str = None, file_name: str = None) -> str:
        if not file_content or not file_name:
            return json.dumps({"cancelled": True, "error": "请提供文件内容和文件名"})

        import tempfile, base64
        ext = file_name.rsplit('.', 1)[-1].lower() if '.' in file_name else 'docx'
        with tempfile.NamedTemporaryFile(suffix=f".{ext}", delete=False) as tmp:
            # 支持 base64 编码内容（二进制文件）
            try:
                raw_bytes = base64.b64decode(file_content)
            except Exception:
                raw_bytes = file_content.encode('utf-8') if isinstance(file_content, str) else file_content
            tmp.write(raw_bytes)
            path = tmp.name

        name = file_name

        if not os.path.exists(path):
            return json.dumps({"success": False, "error": f"文件不存在: {path}", "error_code": "FILE_NOT_FOUND"})
        if not os.access(path, os.R_OK):
            return json.dumps({"success": False, "error": "无权限读取文件", "error_code": "PERMISSION_DENIED"})

        try:
            size = os.path.getsize(path)
        except Exception as e:
            return json.dumps({"success": False, "error": str(e), "error_code": "SIZE_ERROR"})

        if size > 100 * 1024 * 1024:
            return json.dumps({"success": False, "error": "文件过大（超过100MB）", "error_code": "FILE_TOO_LARGE"})

        try:
            from core.entitlements import check_file_size_allowed
            gate = check_file_size_allowed(self._current_entitlements(), size)
            if not gate.get("allowed"):
                try:
                    os.unlink(path)
                except Exception:
                    pass
                return self._deny_json(gate)
        except Exception as gate_err:
            logger.warning("file size entitlement check failed: %s", gate_err)

        sections = []
        styles_meta: dict = {}
        fielded_refs: list = []
        try:
            if ext == "docx":
                elements, sections, styles_meta = self._parse_docx(path)
                fielded_refs = Api._extract_ref_field_display_texts(path)
            elif ext in ("txt", "md"):
                elements = self._parse_text(path)
            else:
                try:
                    elements = self._parse_text(path)
                except Exception:
                    elements, sections, styles_meta = self._parse_docx(path)
                    fielded_refs = Api._extract_ref_field_display_texts(path)

            if not elements:
                return json.dumps({"success": False, "error": "文件为空或无可解析内容", "error_code": "EMPTY_FILE"})
        except Exception as exc:
            error_msg = str(exc)
            error_code = "PARSE_ERROR"
            if "not a valid zip file" in error_msg.lower():
                error_msg = "文件不是有效的Word文档(.docx)"
                error_code = "INVALID_DOCX"
            elif "encoding" in error_msg.lower() or "decode" in error_msg.lower():
                error_msg = "文件编码错误"
                error_code = "ENCODING_ERROR"
            try:
                os.unlink(path)
            except Exception:
                pass
            return json.dumps({"success": False, "error": f"文件解析失败: {error_msg}", "error_code": error_code})

        try:
            os.unlink(path)
        except Exception:
            pass

        # EMF→PNG 转换：将 docx 内的 EMF 图片替换为 PNG，前端 docxBuffers 存修改后版本
        docx_b64 = None
        if ext == "docx":
            try:
                from parsers.dispatcher import convert_emf_images_in_docx
                converted = convert_emf_images_in_docx(raw_bytes)
                if converted is not raw_bytes:
                    docx_b64 = base64.b64encode(converted).decode("ascii")
            except Exception as emf_err:
                logger.warning("EMF→PNG 转换失败，忽略: %s", emf_err)

        result: dict = {
            "success": True, "name": name, "size": size,
            "type": ext, "elements": elements, "element_count": len(elements),
            "sections": sections, "styles": styles_meta,
            "fielded_refs": fielded_refs,
        }
        if docx_b64:
            result["docx_b64"] = docx_b64
        return json.dumps(result, ensure_ascii=False)

    def saveFile(self, content: str, file_name: str = "document.html") -> str:
        try:
            return json.dumps({"success": True, "content": content, "file_name": file_name}, ensure_ascii=False)
        except Exception as exc:
            return json.dumps({"success": False, "error": str(exc)}, ensure_ascii=False)

    def exportDocx(self, content: str, format_params: str = "{}", file_name: str = "document.docx") -> str:
        try:
            import tempfile, base64
            with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
                temp_path = tmp.name
            params = json.loads(format_params) if format_params else {}
            self._write_docx(temp_path, content, params)
            with open(temp_path, "rb") as f:
                docx_content = base64.b64encode(f.read()).decode('utf-8')
            try:
                os.unlink(temp_path)
            except Exception:
                pass
            return json.dumps({"success": True, "content": docx_content, "file_name": file_name}, ensure_ascii=False)
        except Exception as exc:
            logger.error("导出失败: %s", exc)
            return json.dumps({"success": False, "error": str(exc)}, ensure_ascii=False)

    def refreshDocxFields(self, docx_b64: str) -> str:
        """
        用 Word COM (PowerShell) 刷新 docx 的域/目录编号。
        回退：LibreOffice headless 重存（不刷域，但不破坏文件）。
        两者均不可用时返回 error_code=NO_OFFICE_ENGINE。
        """
        import base64 as _b64
        import subprocess
        import tempfile
        import shutil
        from pathlib import Path as _Path

        try:
            docx_bytes = _b64.b64decode(docx_b64)
        except Exception as exc:
            return json.dumps({"error": f"base64 decode failed: {exc}"})

        tmp_path = None
        try:
            with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
                tmp.write(docx_bytes)
                tmp_path = tmp.name

            # 尝试 Word COM（PowerShell，Windows only）
            if sys.platform == "win32":
                src_ps = tmp_path.replace("'", "''")
                ps = (
                    "$ErrorActionPreference='Stop';"
                    "$w=$null;"
                    "try{"
                    f"$w=New-Object -ComObject Word.Application;"
                    "$w.Visible=$false;"
                    f"$d=$w.Documents.Open('{src_ps}');"
                    "$d.Fields.Update();"
                    "$d.TablesOfContents|ForEach-Object{$_.Update()};"
                    "$d.Save();"
                    "$d.Close();"
                    "Write-Output 'OK'"
                    "}catch{Write-Error $_.Exception.Message}"
                    "finally{if($w){$w.Quit()}}"
                )
                try:
                    r = subprocess.run(
                        ["powershell", "-NonInteractive", "-Command", ps],
                        capture_output=True, text=True, timeout=120,
                    )
                    if r.returncode == 0 and "OK" in r.stdout:
                        with open(tmp_path, "rb") as f:
                            out_b64 = _b64.b64encode(f.read()).decode("utf-8")
                        return json.dumps({"success": True, "content": out_b64})
                    logger.warning("Word COM Fields.Update 失败: %s", r.stderr[:200])
                except subprocess.TimeoutExpired:
                    return json.dumps({"error_code": "TIMEOUT", "error": "Word COM 超时（>120s）"})
                except Exception as ps_err:
                    logger.warning("PowerShell 调用失败: %s", ps_err)

            # 尝试 LibreOffice headless（重转存，不真正刷域，但保持文件完整）
            lo_candidates = [
                "libreoffice", "soffice",
                "/usr/bin/libreoffice", "/usr/bin/soffice",
                "C:\\Program Files\\LibreOffice\\program\\soffice.exe",
                "C:\\Program Files (x86)\\LibreOffice\\program\\soffice.exe",
                "/Applications/LibreOffice.app/Contents/MacOS/soffice",
            ]
            lo_cmd = None
            for cmd in lo_candidates:
                try:
                    chk = subprocess.run([cmd, "--version"], capture_output=True, text=True, timeout=10)
                    if chk.returncode == 0:
                        lo_cmd = cmd
                        break
                except (FileNotFoundError, subprocess.TimeoutExpired):
                    continue

            if lo_cmd:
                out_dir = tempfile.mkdtemp(prefix="wc_refresh_")
                try:
                    r = subprocess.run(
                        [lo_cmd, "--headless", "--convert-to", "docx",
                         "--outdir", out_dir, tmp_path],
                        capture_output=True, text=True, timeout=120,
                    )
                    stem = _Path(tmp_path).stem
                    lo_out = os.path.join(out_dir, stem + ".docx")
                    if not os.path.isfile(lo_out):
                        for fn in os.listdir(out_dir):
                            if fn.endswith(".docx"):
                                lo_out = os.path.join(out_dir, fn)
                                break
                    if r.returncode == 0 and os.path.isfile(lo_out):
                        with open(lo_out, "rb") as f:
                            out_b64 = _b64.b64encode(f.read()).decode("utf-8")
                        return json.dumps({"success": True, "content": out_b64})
                except Exception as lo_err:
                    logger.warning("LibreOffice refresh 失败: %s", lo_err)
                finally:
                    shutil.rmtree(out_dir, ignore_errors=True)

            return json.dumps({"error_code": "NO_OFFICE_ENGINE",
                               "error": "Word COM 和 LibreOffice 均不可用，请安装其中之一"})
        except Exception as exc:
            logger.error("refreshDocxFields 异常: %s", exc)
            return json.dumps({"error_code": "ERROR", "error": str(exc)})
        finally:
            if tmp_path:
                try:
                    os.unlink(tmp_path)
                except Exception:
                    pass

    @cached_response(ttl_seconds=600)
    def getSystemInfo(self) -> str:
        return json.dumps({"version": APP_VERSION, "platform": sys.platform, "python": sys.version.split()[0]})

    # ------------------------------------------------------------------
    #  用户认证
    # ------------------------------------------------------------------

    def login(self, email: str, password: str) -> str:
        try:
            self._ensure_supabase_initialized()
            if self._supabase:
                result = self._supabase.sign_in(email, password)
                if result.get("success"):
                    profile = self._supabase.get_profile(result["user_id"])
                    user_info = {
                        "id": result["user_id"], "email": result["email"],
                        "nickname": profile.get("nickname", email.split("@")[0]) if profile else email.split("@")[0],
                        "avatar_url": profile.get("avatar_url", "") if profile else "",
                        "token_quota": profile.get("token_quota", 100000) if profile else 100000,
                        "token_used": profile.get("token_used", 0) if profile else 0,
                        "plan_tier": profile.get("plan_tier", "free") if profile else "free",
                        "plan_status": profile.get("plan_status", "active") if profile else "active",
                        "plan_source": profile.get("plan_source", "system") if profile else "system",
                        "current_period_start": profile.get("current_period_start") if profile else None,
                        "current_period_end": profile.get("current_period_end") if profile else None,
                        "team_id": profile.get("team_id") if profile else None,
                        "feature_flags": profile.get("feature_flags", {}) if profile else {},
                    }
                    app_role = self._resolve_session_role(result.get("access_token"))
                    if app_role:
                        user_info["role"] = app_role
                    self._session = {"user_id": result["user_id"], "access_token": result.get("access_token"), "user_info": user_info}
                    return json.dumps({"success": True, "user": user_info, "token": result.get("access_token")}, ensure_ascii=False)
                else:
                    return json.dumps({"success": False, "error": result.get("error", "登录失败")}, ensure_ascii=False)
            else:
                mock_user = {"id": "mock-user-001", "email": email, "nickname": email.split("@")[0],
                             "avatar_url": "", "token_quota": 100000, "token_used": 12345,
                             "plan_tier": "free", "plan_status": "active", "plan_source": "local",
                             "feature_flags": {}, "role": "user"}
                self._session = {"user_id": "mock-user-001", "user_info": mock_user}
                return json.dumps({"success": True, "user": mock_user, "token": "mock-jwt-token"}, ensure_ascii=False)
        except Exception as exc:
            return json.dumps({"success": False, "error": str(exc)}, ensure_ascii=False)

    def logout(self) -> str:
        try:
            self._ensure_supabase_initialized()
            if self._supabase and self._session.get("access_token"):
                self._supabase.sign_out(self._session.get("access_token"))
            self._session = {}
            return json.dumps({"success": True})
        except Exception as exc:
            return json.dumps({"success": False, "error": str(exc)}, ensure_ascii=False)

    # ------------------------------------------------------------------
    #  Token 管理
    # ------------------------------------------------------------------

    def getTokenUsage(self) -> str:
        try:
            if self._supabase and self._session.get("user_id"):
                usage = self._supabase.get_token_usage(self._session["user_id"])
                usage["plan"] = self._current_entitlements()
                return json.dumps(usage, ensure_ascii=False)
            else:
                from core.token_tracker import TokenTracker
                usage = TokenTracker().get_quota()
                return json.dumps({
                    "token_quota": usage.get('quota', 100000), "token_used": usage.get('used', 0),
                    "token_remaining": usage.get('remaining', 100000),
                    "usage_percentage": usage.get('usage_percentage', 0),
                    "today_usage": 0, "month_usage": usage.get('used', 0),
                    "plan": self._current_entitlements(),
                }, ensure_ascii=False)
        except Exception:
            return json.dumps({"token_quota": 100000, "token_used": 0, "token_remaining": 100000,
                               "usage_percentage": 0, "today_usage": 0, "month_usage": 0,
                               "plan": self._current_entitlements()}, ensure_ascii=False)

    # ------------------------------------------------------------------
    #  模板管理
    # ------------------------------------------------------------------

    def getUserTemplates(self) -> str:
        try:
            if self._supabase and self._session.get("user_id"):
                templates = self._supabase.get_user_templates(self._session["user_id"])
                return json.dumps({"success": True, "templates": templates}, ensure_ascii=False)
            else:
                mock = [
                    {"id": "template-001", "name": "广东海洋大学毕业论文", "category": "thesis", "file_url": "", "created_at": "2026-04-15"},
                    {"id": "template-002", "name": "尖草坪区绩效评价报告", "category": "gov", "file_url": "", "created_at": "2026-04-16"},
                ]
                return json.dumps({"success": True, "templates": mock}, ensure_ascii=False)
        except Exception as exc:
            return json.dumps({"success": False, "error": str(exc)}, ensure_ascii=False)

    @staticmethod
    def _extract_docx_format_rules(file_path: str) -> dict:
        """从 .docx 文件提取排版参数，返回与前端 _aiParsedRules 格式相同的 dict。"""
        try:
            from docx import Document
            from docx.oxml.ns import qn
            doc = Document(file_path)
        except Exception:
            return {}

        def _pt(sz):
            try:
                return round(float(sz.pt), 1)
            except Exception:
                return None

        def _cm(emu):
            try:
                return round(emu / 360000.0, 2)
            except Exception:
                return None

        def _cn_font(run):
            try:
                rpr = run._r.find(qn('w:rPr'))
                if rpr is None:
                    return None
                rf = rpr.find(qn('w:rFonts'))
                if rf is None:
                    return None
                return rf.get(qn('w:eastAsia')) or rf.get(qn('w:hAnsi'))
            except Exception:
                return None

        ALIGN = {1: '左对齐', 2: '居中', 3: '右对齐', 4: '两端对齐'}
        rules = {}

        if doc.sections:
            s = doc.sections[0]
            rules.update({
                'marginTop': _cm(s.top_margin),
                'marginBottom': _cm(s.bottom_margin),
                'marginLeft': _cm(s.left_margin),
                'marginRight': _cm(s.right_margin),
            })

        found = {k: False for k in ('h1', 'h2', 'h3', 'body')}
        for para in doc.paragraphs[:300]:
            if all(found.values()):
                break
            sname = (para.style.name or '').lower().strip()
            for lvl in (1, 2, 3):
                key = f'h{lvl}'
                if found[key]:
                    continue
                if sname in (f'heading {lvl}', f'标题 {lvl}', f'标题{lvl}', f'heading{lvl}'):
                    if para.runs:
                        run = para.runs[0]
                        rules[f'h{lvl}Font'] = _cn_font(run) or run.font.name
                        rules[f'h{lvl}Size'] = _pt(run.font.size)
                        found[key] = True
                    break
            if not found['body'] and sname in ('normal', '正文', 'body text', 'default paragraph font'):
                if para.text.strip() and para.runs:
                    run = para.runs[0]
                    cn = _cn_font(run)
                    rules['bodyFont'] = cn or run.font.name
                    rules['bodySize'] = _pt(run.font.size)
                    rules['westFont'] = run.font.name
                    pf = para.paragraph_format
                    try:
                        if pf.first_line_indent and run.font.size and run.font.size.pt > 0:
                            rules['indent'] = round(float(pf.first_line_indent / run.font.size), 1)
                    except Exception:
                        pass
                    try:
                        if pf.line_spacing is not None:
                            rules['lineHeight'] = round(float(pf.line_spacing), 2)
                    except Exception:
                        pass
                    rules['align'] = ALIGN.get(pf.alignment)
                    found['body'] = True

        return {k: v for k, v in rules.items() if v is not None}

    @staticmethod
    def _extract_docx_text(file_path: str, max_chars: int = 4000) -> str:
        """从 .docx 文件提取纯文本内容，供 AI 理解模板描述"""
        try:
            from docx import Document
            doc = Document(file_path)
            lines = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
            text = "\n".join(lines)
            return text[:max_chars]
        except Exception:
            return ""

    def uploadTemplate(self, file_path: str, name: str) -> str:
        try:
            format_rules = {}
            doc_text = ""
            if file_path.lower().endswith('.docx') and os.path.exists(file_path):
                format_rules = self._extract_docx_format_rules(file_path)
                doc_text = self._extract_docx_text(file_path)

            if self._supabase and self._session.get("user_id"):
                user_id = self._session["user_id"]
                file_url = self._supabase.upload_template_file(user_id, file_path, name)
                if file_url:
                    record = self._supabase.insert_template(user_id, name, file_url, "custom")
                    return json.dumps({"success": True, "template_id": record.get("id", "") if record else "", "name": name, "file_url": file_url, "format_rules": format_rules, "doc_text": doc_text}, ensure_ascii=False)
                return json.dumps({"success": False, "error": "上传失败"})
            else:
                if not os.path.exists(file_path):
                    return json.dumps({"success": False, "error": "文件不存在"})
                return json.dumps({"success": True, "template_id": "template-new-001", "name": name, "format_rules": format_rules, "doc_text": doc_text}, ensure_ascii=False)
        except Exception as exc:
            return json.dumps({"success": False, "error": str(exc)}, ensure_ascii=False)

    def saveTemplateSettings(self, name: str, template_data: str) -> str:
        try:
            data = json.loads(template_data) if isinstance(template_data, str) else template_data
            if self._supabase and self._session.get("user_id"):
                record = self._supabase.insert_template(self._session["user_id"], name, "", "custom", data)
                return json.dumps({"success": True, "template_id": record.get("id", "") if record else "", "name": name}, ensure_ascii=False)
            return json.dumps({"success": True, "template_id": "local-" + name, "name": name}, ensure_ascii=False)
        except Exception as exc:
            return json.dumps({"success": False, "error": str(exc)}, ensure_ascii=False)

    def deleteTemplate(self, template_id: str) -> str:
        try:
            if self._supabase and self._session.get("user_id"):
                success = self._supabase.delete_template(template_id, self._session["user_id"])
                return json.dumps({"success": success})
            return json.dumps({"success": True})
        except Exception as exc:
            return json.dumps({"success": False, "error": str(exc)}, ensure_ascii=False)

    # ------------------------------------------------------------------
    #  用户设置
    # ------------------------------------------------------------------

    def getUserSettings(self) -> str:
        try:
            if self._supabase and self._session.get("user_id"):
                settings = self._supabase.get_user_settings(self._session["user_id"])
                return json.dumps({"success": True, "settings": settings}, ensure_ascii=False)
            else:
                from core.supabase_client import SupabaseClient
                return json.dumps({"success": True, "settings": SupabaseClient._DEFAULT_SETTINGS.copy()}, ensure_ascii=False)
        except Exception as exc:
            return json.dumps({"success": False, "error": str(exc)}, ensure_ascii=False)

    def saveUserSettings(self, settings_str: str) -> str:
        try:
            settings = json.loads(settings_str) if isinstance(settings_str, str) else settings_str
            if self._supabase and self._session.get("user_id"):
                ok = self._supabase.save_user_settings(self._session["user_id"], settings)
                return json.dumps({"success": ok}, ensure_ascii=False)
            return json.dumps({"success": True}, ensure_ascii=False)
        except Exception as exc:
            return json.dumps({"success": False, "error": str(exc)}, ensure_ascii=False)

    # ------------------------------------------------------------------
    #  排版规范保存/加载
    # ------------------------------------------------------------------

    _FORMAT_RULES_FILE = os.path.join(os.path.dirname(__file__), "web", "format_rules.json")
    _TRUSTED_FORMAT_RULE_SOURCES = {"manual", "team_manual", "ai_reviewed"}

    @classmethod
    def _normalize_format_rules(cls, rules: dict | None, default_source: str | None = None) -> dict | None:
        if not isinstance(rules, dict):
            return None
        normalized = dict(rules)
        if normalized.get("savedByUser") is True and default_source and not normalized.get("saveSource"):
            normalized["saveSource"] = default_source
        return normalized

    @classmethod
    def _is_trusted_format_rules(cls, rules: dict | None) -> bool:
        if not isinstance(rules, dict):
            return False
        source = str(rules.get("saveSource") or "").strip().lower()
        return bool(rules.get("savedByUser") is True and source in cls._TRUSTED_FORMAT_RULE_SOURCES)

    def saveFormatRequirements(self, rules_json: str, scope: str = "personal") -> str:
        """Save user format rules — Supabase primary, local JSON fallback."""
        try:
            rules = json.loads(rules_json) if isinstance(rules_json, str) else rules_json

            scope = str(scope or "personal").strip().lower()
            if scope not in {"personal", "team"}:
                scope = "personal"
            default_source = "team_manual" if scope == "team" else "manual"
            rules = self._normalize_format_rules(rules, default_source=default_source)
            if isinstance(rules, dict):
                rules["savedByUser"] = True
                rules["saveSource"] = str(rules.get("saveSource") or default_source).strip().lower()
            gate_feature = "team_rule_share" if scope == "team" else "personal_rule_library"
            gate = self._check_feature_gate(gate_feature)
            if not gate.get("allowed"):
                return self._deny_json(gate)

            if scope == "team":
                team_id = self._current_team_id()
                if not team_id:
                    return self._deny_json(self._team_gate("TEAM_REQUIRED", "当前账号尚未加入团队，无法保存团队规则库。"))
                uid = self._session.get("user_id") if self._session else None
                if self._supabase and uid:
                    result = self._supabase.save_team_format_rules(team_id, uid, rules)
                    return json.dumps({
                        "success": result.get("success", False),
                        "storage": result.get("storage", "supabase"),
                        "scope": "team",
                        "team_id": team_id,
                        "error": result.get("error"),
                    }, ensure_ascii=False)
                self._local_team_rules[team_id] = rules
                return json.dumps({"success": True, "storage": "local", "scope": "team", "team_id": team_id}, ensure_ascii=False)

            # ── Supabase（已登录时） ──
            uid = self._session.get("user_id") if self._session else None
            if self._supabase and uid:
                try:
                    self._supabase.client.table("user_format_rules").upsert(
                        {"user_id": uid, "rules_json": rules,
                         "updated_at": "now()"},
                        on_conflict="user_id",
                    ).execute()
                    # 同步本地文件作为离线缓存
                    with open(self._FORMAT_RULES_FILE, "w", encoding="utf-8") as f:
                        json.dump(rules, f, ensure_ascii=False, indent=2)
                    return json.dumps({"success": True, "storage": "supabase", "scope": "personal"}, ensure_ascii=False)
                except Exception as sb_err:
                    print(f"[FormatRules] Supabase save failed, falling back to local: {sb_err}")

            # ── 本地 JSON 兜底 ──
            with open(self._FORMAT_RULES_FILE, "w", encoding="utf-8") as f:
                json.dump(rules, f, ensure_ascii=False, indent=2)
            return json.dumps({"success": True, "storage": "local", "scope": "personal"}, ensure_ascii=False)
        except Exception as exc:
            return json.dumps({"success": False, "error": str(exc)}, ensure_ascii=False)

    def loadFormatRequirements(self, scope: str = "personal") -> str:
        """Load saved format rules — Supabase primary, local JSON fallback."""
        try:
            scope = str(scope or "personal").strip().lower()
            if scope not in {"personal", "team"}:
                scope = "personal"
            if scope == "team":
                gate = self._check_feature_gate("team_rule_share")
                if not gate.get("allowed"):
                    return self._deny_json(gate)
                team_id = self._current_team_id()
                if not team_id:
                    return self._deny_json(self._team_gate("TEAM_REQUIRED", "当前账号尚未加入团队，无法读取团队规则库。"))
                uid = self._session.get("user_id") if self._session else None
                if self._supabase and uid:
                    result = self._supabase.get_team_format_rules(team_id)
                    team_rules = self._normalize_format_rules(result.get("rules"))
                    return json.dumps({
                        "success": True,
                        "rules": team_rules if self._is_trusted_format_rules(team_rules) else None,
                        "storage": result.get("storage", "supabase"),
                        "scope": "team",
                        "team_id": team_id,
                    }, ensure_ascii=False)
                local_rules = self._normalize_format_rules(self._local_team_rules.get(team_id))
                return json.dumps({
                    "success": True,
                    "rules": local_rules if self._is_trusted_format_rules(local_rules) else None,
                    "storage": "local",
                    "scope": "team",
                    "team_id": team_id,
                }, ensure_ascii=False)

            # ── Supabase（已登录时） ──
            uid = self._session.get("user_id") if self._session else None
            if self._supabase and uid:
                try:
                    res = (
                        self._supabase.client
                        .table("user_format_rules")
                        .select("rules_json, updated_at")
                        .eq("user_id", uid)
                        .limit(1)
                        .execute()
                    )
                    if res.data:
                        rules = self._normalize_format_rules(res.data[0]["rules_json"])
                        trusted_rules = rules if self._is_trusted_format_rules(rules) else None
                        if trusted_rules:
                            with open(self._FORMAT_RULES_FILE, "w", encoding="utf-8") as f:
                                json.dump(trusted_rules, f, ensure_ascii=False, indent=2)
                        return json.dumps({"success": True, "rules": trusted_rules,
                                           "storage": "supabase", "scope": "personal"}, ensure_ascii=False)
                except Exception as sb_err:
                    print(f"[FormatRules] Supabase load failed, falling back to local: {sb_err}")

            # ── 本地 JSON 兜底 ──
            if os.path.exists(self._FORMAT_RULES_FILE):
                with open(self._FORMAT_RULES_FILE, "r", encoding="utf-8") as f:
                    rules = self._normalize_format_rules(json.load(f))
                return json.dumps({"success": True, "rules": rules if self._is_trusted_format_rules(rules) else None,
                                   "storage": "local", "scope": "personal"}, ensure_ascii=False)
            return json.dumps({"success": True, "rules": None, "scope": "personal"}, ensure_ascii=False)
        except Exception as exc:
            return json.dumps({"success": False, "error": str(exc)}, ensure_ascii=False)

    def runBatchQA(self, files_json: str, categories_str: str = '["typo", "consistency", "logic", "format", "crossref"]') -> str:
        try:
            gate = self._check_feature_gate("batch_check")
            if not gate.get("allowed"):
                return self._deny_json(gate)

            files = json.loads(files_json) if isinstance(files_json, str) else (files_json or [])
            categories = json.loads(categories_str) if isinstance(categories_str, str) else categories_str
            results = []
            for item in files:
                content = (item or {}).get("content", "")
                name = (item or {}).get("name") or "untitled"
                elements = (item or {}).get("elements")
                result = self.runQA(
                    content,
                    json.dumps(categories or ["typo", "consistency", "logic", "format", "crossref"], ensure_ascii=False),
                    elements_json=json.dumps(elements, ensure_ascii=False) if elements is not None else None,
                )
                results.append({
                    "name": name,
                    "result": json.loads(result) if isinstance(result, str) else result,
                })
            return json.dumps({"success": True, "results": results, "count": len(results)}, ensure_ascii=False)
        except Exception as exc:
            return json.dumps({"success": False, "error": str(exc)}, ensure_ascii=False)

    # ------------------------------------------------------------------
    #  质量检查
    # ------------------------------------------------------------------

    def runQA(self, content: str, categories_str: str = '["typo", "consistency", "logic", "format", "crossref"]',
              elements_json: str = None, format_rules_json: str = None) -> str:
        try:
            from core.entitlements import check_quota_available
            gate = check_quota_available(self._current_entitlements(), "rule_check")
            if not gate.get("allowed"):
                return self._deny_json(gate)

            qa_health = self._warmup_qa_capabilities(force=False)
            if not qa_health.get("ready", False):
                # 用户刚完成依赖安装时，给一次强制重检机会，避免必须重启后端。
                qa_health = self._warmup_qa_capabilities(force=True, aggressive_retry=True)
            if not qa_health.get("ready", False):
                return json.dumps({
                    "success": False,
                    "error_code": "QA_CAPABILITY_NOT_READY",
                    "error": "QA 能力未就绪，请联系管理员完成依赖安装。",
                    "missing": qa_health.get("missing", []),
                    "capabilities": qa_health.get("capabilities", {}),
                }, ensure_ascii=False)

            from core.qa_engine import QAEngine
            from core.document_model import DocumentModel, DocElement, ElementType, FontStyle
            from html.parser import HTMLParser

            class _StripHTML(HTMLParser):
                def __init__(self): super().__init__(); self._parts = []
                def handle_data(self, d): self._parts.append(d)
                def get_text(self): return ''.join(self._parts)

            doc = DocumentModel(title="检查文档", source_format="html")

            # ── 优先路径：前端传来结构化 elements（含 runs 字体信息）──
            if elements_json:
                try:
                    import re as _re
                    raw_elems = json.loads(elements_json) if isinstance(elements_json, str) else elements_json
                    _TYPE_MAP = {"h1": (ElementType.HEADING, 1), "h2": (ElementType.HEADING, 2),
                                 "h3": (ElementType.HEADING, 3), "p": (ElementType.PARAGRAPH, 0),
                                 "li": (ElementType.PARAGRAPH, 0)}
                    for el in (raw_elems or []):
                        t = el.get("type", "p")
                        etype, lvl = _TYPE_MAP.get(t, (ElementType.PARAGRAPH, 0))
                        txt = (el.get("text") or "").strip()
                        if not txt:
                            continue
                        fs = FontStyle()
                        runs = el.get("runs") or []
                        if runs:
                            # 按字符数加权取主导字体/字号（与前端 _getDominantEastAsiaFont 同逻辑）
                            font_w: dict[str, int] = {}
                            size_w: dict[float, int] = {}
                            for r in runs:
                                rlen = len((r.get("text") or "").strip())
                                if not rlen:
                                    continue
                                fn = r.get("font_eastAsia") or r.get("font_name") or ""
                                if fn:
                                    font_w[fn] = font_w.get(fn, 0) + rlen
                                sz = r.get("font_size_pt")
                                if sz:
                                    size_w[float(sz)] = size_w.get(float(sz), 0) + rlen
                            if font_w:
                                fs.font_name_cn = max(font_w, key=font_w.get)
                            if size_w:
                                fs.font_size_pt = max(size_w, key=size_w.get)
                        # 第十八批：保留前端传入的 metadata（含 structure_role 等结构识别字段）
                        meta = el.get("metadata") if isinstance(el.get("metadata"), dict) else {}
                        doc.elements.append(DocElement(
                            element_type=etype, content=txt, level=lvl,
                            font_style=fs, metadata=dict(meta)))
                except Exception as _e:
                    print(f"[runQA] elements_json 解析失败，回退 HTML 路径: {_e}")
                    doc.elements.clear()

            # ── 回退路径：从 HTML 解析（无字体信息）──
            if not doc.elements and content:
                import re as _re
                raw = content if isinstance(content, str) else str(content)
                # 逐个匹配 HTML 块元素（h1-h6, p, li）以及换行纯文本
                for m in _re.finditer(
                    r'<h([1-6])[^>]*>(.*?)</h[1-6]>|<(?:p|li)[^>]*>(.*?)</(?:p|li)>|([^\n<]+)',
                    raw, _re.IGNORECASE | _re.DOTALL
                ):
                    if m.group(1):  # heading
                        p = _StripHTML(); p.feed(m.group(2)); txt = p.get_text().strip()
                        if txt:
                            doc.elements.append(DocElement(element_type=ElementType.HEADING,
                                content=txt, level=int(m.group(1))))
                    elif m.group(3) is not None:  # p / li
                        p = _StripHTML(); p.feed(m.group(3)); txt = p.get_text().strip()
                        if txt:
                            elem_type = ElementType.PARAGRAPH
                            if _re.match(r"^(图|表)\s*\d+([\-\.]\d+)?", txt):
                                elem_type = ElementType.CAPTION
                            elif _re.match(r"^\[\d+\]", txt):
                                elem_type = ElementType.REFERENCE
                            doc.elements.append(DocElement(element_type=elem_type, content=txt))
                    elif m.group(4):  # 纯文本行
                        txt = m.group(4).strip()
                        if txt:
                            elem_type = ElementType.PARAGRAPH
                            if _re.match(r"^(图|表)\s*\d+([\-\.]\d+)?", txt):
                                elem_type = ElementType.CAPTION
                            elif _re.match(r"^\[\d+\]", txt):
                                elem_type = ElementType.REFERENCE
                            doc.elements.append(DocElement(element_type=elem_type, content=txt))

            categories = json.loads(categories_str) if isinstance(categories_str, str) else categories_str

            # 优先使用前端显式透传的受信规则；缺失时再回退到本地已保存规则。
            format_rules = None
            if "format" in categories:
                try:
                    from core.format_checker import FormatRules

                    _rules_dict = None
                    _frontend_explicit = format_rules_json is not None
                    if _frontend_explicit and format_rules_json != 'null':
                        _rules_dict = json.loads(format_rules_json) if isinstance(format_rules_json, str) else format_rules_json
                        _rules_dict = self._normalize_format_rules(_rules_dict)
                        if not self._is_trusted_format_rules(_rules_dict):
                            _rules_dict = None

                    # 只有前端完全没有传 format_rules 时，才 fallback 读磁盘；
                    # 前端明确传了 null 表示本次不使用排版规范，不再 fallback。
                    if not _frontend_explicit and _rules_dict is None and os.path.exists(self._FORMAT_RULES_FILE):
                        with open(self._FORMAT_RULES_FILE, "r", encoding="utf-8") as _f:
                            _rules_dict = json.load(_f)
                        if isinstance(_rules_dict, str):
                            _rules_dict = json.loads(_rules_dict)
                        _rules_dict = self._normalize_format_rules(_rules_dict)
                        if not self._is_trusted_format_rules(_rules_dict):
                            _rules_dict = None

                    _fr = FormatRules.from_dict(_rules_dict) if _rules_dict else None
                    if _fr and not _fr.is_empty():
                        format_rules = _fr
                except Exception:
                    pass

            engine = QAEngine(config=self._qa_runtime_config)
            report = engine.check(doc, categories, format_rules=format_rules)
            self._increment_usage("rule_check_used_today")
            return json.dumps({
                "success": True,
                "issues": [
                    {"id": i.issue_id,
                     "category": i.category.value if hasattr(i.category, 'value') else str(i.category),
                     "severity": i.severity.value if hasattr(i.severity, 'value') else str(i.severity),
                     "title": i.title, "description": i.description,
                     "suggestion": i.suggestion, "confidence": i.confidence,
                     "location_text": (i.location_text or '').strip(),
                     "element_index": getattr(i, "element_index", -1),
                     "start_pos": getattr(i, "start_pos", -1),
                     "end_pos": getattr(i, "end_pos", -1),
                     "rule_id": getattr(i, "rule_id", ""),
                     "checker": getattr(i, "checker", "")}
                    for i in report.issues
                ],
                "stats": {
                    "total": len(report.issues),
                    "errors": len([i for i in report.issues if i.severity.value == "error"]),
                    "warnings": len([i for i in report.issues if i.severity.value == "warning"]),
                    "infos": len([i for i in report.issues if i.severity.value == "info"]),
                },
            }, ensure_ascii=False)
        except Exception as exc:
            return json.dumps({"success": False, "error": str(exc)})

    def acceptSuggestion(self, content: str, issue_id: str, original_text: str, suggested_text: str) -> str:
        try:
            corrected = content.replace(original_text, suggested_text, 1)
            if corrected != content:
                return json.dumps({"success": True, "corrected_content": corrected,
                                   "message": f"已将 '{original_text}' 改为 '{suggested_text}'"}, ensure_ascii=False)
            try:
                from bs4 import BeautifulSoup
                soup = BeautifulSoup(content, 'html.parser')
                for string in soup.find_all(string=True):
                    if original_text in string:
                        string.replace_with(string.replace(original_text, suggested_text, 1))
                        return json.dumps({"success": True, "corrected_content": str(soup),
                                           "message": f"已将 '{original_text}' 改为 '{suggested_text}'"}, ensure_ascii=False)
            except ImportError:
                pass
            return json.dumps({"success": False, "error": f"未找到待替换文本: {original_text}"})
        except Exception as exc:
            return json.dumps({"success": False, "error": str(exc)})

    # ------------------------------------------------------------------
    #  交叉引用
    # ------------------------------------------------------------------

    def runXRef(self, content: str, fielded_refs: list = None, elements_json=None, deep: bool = False) -> str:
        try:
            if deep:
                gate = self._check_feature_gate("deep_xref")
                if not gate.get("allowed"):
                    return self._deny_json(gate)

            from core.document_model import DocumentModel, DocElement, ElementType
            from core.crossref_engine import TargetScanner, RefPointScanner, CrossRefMatcher
            from core.crossref_models import CrossRefStatus, RefTargetType
            from html.parser import HTMLParser

            class _Strip(HTMLParser):
                def __init__(self): super().__init__(); self._parts = []
                def handle_data(self, d): self._parts.append(d)
                def get_text(self): return ''.join(self._parts)

            doc = DocumentModel(title="xref", source_format="html")

            # ── 第十八批：优先走结构化路径（含 metadata.structure_role），让 TargetScanner
            #    跳过目录里的"第3章 模型 ......15"等条目 ──
            if elements_json:
                try:
                    raw_elems = json.loads(elements_json) if isinstance(elements_json, str) else elements_json
                    _XREF_TYPE_MAP = {
                        "h1": (ElementType.HEADING, 1),
                        "h2": (ElementType.HEADING, 2),
                        "h3": (ElementType.HEADING, 3),
                        "p":  (ElementType.PARAGRAPH, 0),
                        "li": (ElementType.PARAGRAPH, 0),
                        "caption": (ElementType.CAPTION, 0),
                        "ref": (ElementType.REFERENCE, 0),
                    }
                    for el in (raw_elems or []):
                        t = (el.get("type") or "p").lower()
                        etype, lvl = _XREF_TYPE_MAP.get(t, (ElementType.PARAGRAPH, 0))
                        txt = (el.get("text") or "").strip()
                        if not txt:
                            continue
                        meta = el.get("metadata") if isinstance(el.get("metadata"), dict) else {}
                        doc.elements.append(DocElement(
                            element_type=etype, content=txt, level=lvl,
                            metadata=dict(meta)))
                except Exception as _xref_struct_err:
                    print(f"[runXRef] elements_json 解析失败，回退 HTML 正则路径: {_xref_struct_err}")
                    doc.elements.clear()

            # ── 回退路径：仅有 HTML 时按现有逻辑解析（无 metadata） ──
            if not doc.elements:
                for m in re.finditer(
                    r'<h([1-6])[^>]*>(.*?)</h[1-6]>|<(?:p|li)[^>]*>(.*?)</(?:p|li)>',
                    content, re.IGNORECASE | re.DOTALL
                ):
                    if m.group(1):
                        p = _Strip(); p.feed(m.group(2)); txt = p.get_text().strip()
                        if txt:
                            doc.elements.append(DocElement(element_type=ElementType.HEADING,
                                content=txt, level=int(m.group(1))))
                    elif m.group(3) is not None:
                        p = _Strip(); p.feed(m.group(3)); txt = p.get_text().strip()
                        if txt:
                            et = ElementType.PARAGRAPH
                            if re.match(r"^\s*\[\d+\]", txt):
                                et = ElementType.REFERENCE
                            doc.elements.append(DocElement(element_type=et, content=txt))

            scanner = TargetScanner()
            targets = scanner.scan(doc)
            ref_scanner = RefPointScanner()
            ref_points = ref_scanner.scan(doc)  # 单遍扫描，scan_index 已在此赋值

            # runXRef 只输出参考文献交叉引用建议，图表/章节/公式由 QA CrossRefChecker 处理
            targets = [t for t in targets if t.target_type == RefTargetType.REFERENCE]
            ref_points = [rp for rp in ref_points if rp.target_type == RefTargetType.REFERENCE]

            matcher = CrossRefMatcher()
            report = matcher.match(targets, ref_points)

            status_map = {
                CrossRefStatus.VALID: "valid",
                CrossRefStatus.DANGLING: "dangling",
                CrossRefStatus.UNREFERENCED: "unreferenced",
                CrossRefStatus.MISMATCH: "dangling",
            }
            targets_out = [{"type": t.target_type.value, "text": t.number, "label": t.label, "title": t.title or ""} for t in targets]
            matches_out = [{"status": status_map.get(m.status, "dangling"),
                            "reference": m.ref_point.ref_text if m.ref_point.ref_text else "—",
                            "target": m.target.label if m.target else "",
                            "message": m.message,
                            "context": getattr(m.ref_point, 'context', ''),
                            "element_index": getattr(m.ref_point, 'element_index', None),
                            "start_pos": getattr(m.ref_point, 'start_pos', None),
                            "scan_index": getattr(m.ref_point, 'scan_index', -1)} for m in report.matches]
            # scan_index 是唯一排序键：扫描阶段按文档阅读顺序分配，禁止任何 reference 二次排序。
            # UNREFERENCED/DUPLICATE 等合成条目 scan_index=-1，统一排到末尾。
            matches_out.sort(key=lambda m: m["scan_index"] if m.get("scan_index", -1) >= 0 else 999999)

            # Build xref_issues for adoption UI (第十三批)
            # "unreferenced" = valid plain-text ref that can be adopted → REF field
            # "dangling"     = text ref without a matching target (warning only)
            # Filter out refs that are already Word REF fields (fielded_refs from openFile)
            already_fielded: set = set(fielded_refs) if fielded_refs else set()

            xref_issues_out = []
            seen_labels: set = set()
            for m in report.matches:
                if m.status == CrossRefStatus.VALID and m.ref_point.ref_text and m.ref_point.ref_text not in seen_labels:
                    # Skip if this label is already a Word REF field in the document
                    if m.ref_point.ref_text in already_fielded:
                        seen_labels.add(m.ref_point.ref_text)
                        continue
                    seen_labels.add(m.ref_point.ref_text)
                    xref_issues_out.append({
                        "type": "unreferenced",
                        "target_label": m.ref_point.ref_text,
                        "bookmark_name": m.target.bookmark_name if m.target else None,
                        "element_index": getattr(m.ref_point, "element_index", None),
                        "start_pos": getattr(m.ref_point, "start_pos", None),
                        "scan_index": getattr(m.ref_point, "scan_index", -1),
                        "title": f"可创建字段：{m.ref_point.ref_text}",
                        "description": f"「{m.ref_point.ref_text}」在正文中为纯文本，采纳后导出时自动转为 Word REF 字段",
                        "suggestion": m.ref_point.ref_text,
                    })
            for m in report.matches:
                if m.status == CrossRefStatus.DANGLING and m.ref_point.ref_text:
                    xref_issues_out.append({
                        "type": "dangling",
                        "target_label": m.ref_point.ref_text,
                        "bookmark_name": None,
                        "element_index": getattr(m.ref_point, "element_index", None),
                        "start_pos": getattr(m.ref_point, "start_pos", None),
                        "scan_index": getattr(m.ref_point, "scan_index", -1),
                        "title": f"悬空引用：{m.ref_point.ref_text}",
                        "description": f"「{m.ref_point.ref_text}」找不到对应目标，请检查引用是否正确",
                        "suggestion": None,
                    })
            # scan_index 是唯一排序键，与 matches_out 保持一致。
            xref_issues_out.sort(key=lambda x: x["scan_index"] if x.get("scan_index", -1) >= 0 else 999999)

            return json.dumps({
                "success": True, "targets": targets_out, "matches": matches_out,
                "xref_issues": xref_issues_out,
                "summary": {"total_targets": len(targets_out),
                             "valid_matches": sum(1 for m in matches_out if m["status"] == "valid"),
                             "dangling_references": sum(1 for m in matches_out if m["status"] == "dangling"),
                             "unreferenced": sum(1 for m in matches_out if m["status"] == "unreferenced")},
            }, ensure_ascii=False)
        except Exception as exc:
            return json.dumps({"success": False, "error": str(exc)})

    # ------------------------------------------------------------------
    #  排版
    # ------------------------------------------------------------------

    def applyFormat(self, content: str, rules_str: str = "{}") -> str:
        try:
            rules = json.loads(rules_str) if rules_str else {}
            styles = {
                "h1": {"font-family": rules.get("heading1_font", "宋体"), "font-size": f"{rules.get('heading1_size', '16')}pt", "line-height": rules.get("line_height", "1.5")},
                "h2": {"font-family": rules.get("heading2_font", "宋体"), "font-size": f"{rules.get('heading2_size', '14')}pt", "line-height": rules.get("line_height", "1.5")},
                "h3": {"font-family": rules.get("heading3_font", "宋体"), "font-size": f"{rules.get('heading3_size', '12')}pt", "line-height": rules.get("line_height", "1.5")},
                "p":  {"font-family": rules.get("body_font", "宋体"), "font-size": f"{rules.get('body_size', '10.5')}pt", "line-height": rules.get("line_height", "1.5")},
            }
            styled = self._apply_css_styles(content, styles)
            return json.dumps({"success": True, "message": "排版规则已应用", "rules_applied": len(rules), "formatted_content": styled}, ensure_ascii=False)
        except Exception as exc:
            return json.dumps({"success": False, "error": str(exc)})

    def _apply_css_styles(self, html_content: str, styles: dict) -> str:
        result = html_content
        for tag, tag_styles in styles.items():
            style_str = "; ".join(f"{k}: {v}" for k, v in tag_styles.items())
            def merge_styles(match, _tag=tag, _style=style_str):
                existing = match.group(1)
                if 'style=' in existing:
                    existing = re.sub(r'style="([^"]*)"', f'style="\\1; {_style}"', existing)
                    return f"<{_tag}{existing}>"
                return f'<{_tag}{existing} style="{_style}">'
            result = re.sub(f"<{tag}([^>]*)>", merge_styles, result, flags=re.IGNORECASE)
        return result

    # ------------------------------------------------------------------
    #  文档管理
    # ------------------------------------------------------------------

    def saveDocument(self, content: str, title: str = "未命名文档") -> str:
        try:
            if self._supabase and self._session.get("user_id"):
                doc_id = self._supabase.save_document(self._session["user_id"], title, {"html_content": content})
                if doc_id:
                    return json.dumps({"success": True, "message": "文档已保存到云端", "doc_id": doc_id}, ensure_ascii=False)
                return json.dumps({"success": False, "error": "保存失败"})
            else:
                import tempfile
                cache_dir = os.path.join(tempfile.gettempdir(), "wordcraft-cache")
                os.makedirs(cache_dir, exist_ok=True)
                doc_file = os.path.join(cache_dir, "current_doc.json")
                with open(doc_file, "w", encoding="utf-8") as f:
                    json.dump({"title": title, "content": content}, f, ensure_ascii=False)
                return json.dumps({"success": True, "message": "文档已保存到本地", "path": doc_file}, ensure_ascii=False)
        except Exception as exc:
            return json.dumps({"success": False, "error": str(exc)}, ensure_ascii=False)

    def loadDocument(self, doc_id: str = "current") -> str:
        try:
            if self._supabase and self._session.get("user_id"):
                doc = self._supabase.load_document(doc_id, self._session["user_id"])
                if doc:
                    return json.dumps({"success": True, "document": doc}, ensure_ascii=False)
                return json.dumps({"success": False, "error": "未找到文档"}, ensure_ascii=False)
            else:
                import tempfile
                doc_file = os.path.join(tempfile.gettempdir(), "wordcraft-cache", "current_doc.json")
                if os.path.exists(doc_file):
                    with open(doc_file, "r", encoding="utf-8") as f:
                        return json.dumps({"success": True, "document": json.load(f)}, ensure_ascii=False)
                return json.dumps({"success": False, "error": "未找到保存的文档"}, ensure_ascii=False)
        except Exception as exc:
            return json.dumps({"success": False, "error": str(exc)}, ensure_ascii=False)

    def updateDocument(self, doc_id: str, content: str, title: str = "未命名文档") -> str:
        try:
            if self._supabase and self._session.get("user_id"):
                ok = self._supabase.update_document(doc_id, self._session["user_id"],
                                                     {"content": {"html_content": content}, "title": title})
                return json.dumps({"success": ok, "doc_id": doc_id}, ensure_ascii=False)
            return json.dumps({"success": True, "doc_id": doc_id}, ensure_ascii=False)
        except Exception as exc:
            return json.dumps({"success": False, "error": str(exc)}, ensure_ascii=False)

    def getDocumentList(self) -> str:
        try:
            if self._supabase and self._session.get("user_id"):
                docs = self._supabase.get_user_documents(self._session["user_id"])
                return json.dumps({"success": True, "documents": docs}, ensure_ascii=False)
            else:
                return json.dumps({"success": True, "documents": [
                    {"id": "doc-001", "title": "益海嘉里项目综合效益评价", "updated_at": "2026-04-18 10:30", "word_count": 5678}
                ]}, ensure_ascii=False)
        except Exception as exc:
            return json.dumps({"success": False, "error": str(exc)}, ensure_ascii=False)

    # ------------------------------------------------------------------
    #  AI 调用
    # ------------------------------------------------------------------

    def callAI(self, system_prompt: str, user_message: str, config_str: str = "{}") -> str:
        config = {}
        if config_str:
            try:
                config = json.loads(config_str) if isinstance(config_str, str) else config_str
            except Exception:
                pass
        purpose = str(config.get("purpose") or "ai_parse").strip()
        if purpose not in ("ai_qa", "ai_parse"):
            purpose = "ai_parse"
        charge_quota = bool(config.get("charge_quota", True))
        gate = self._check_feature_gate(purpose, purpose if charge_quota else None)
        if not gate.get("allowed"):
            return self._deny_json(gate)

        model = str(config.get("model", "deepseek-v4-flash")).strip()
        if (not model) or (not model.lower().startswith("deepseek-")):
            # 兼容旧前端保存的 doubao/deepseek-v3 模型名，统一回退到可用默认模型。
            model = "deepseek-v4-flash"
        temperature = config.get("temperature", 0.3)
        max_tokens = int(config.get("max_tokens", 2048))
        reasoning_effort = config.get("reasoning_effort", "high")
        no_thinking = bool(config.get("no_thinking", False))
        # 优先环境变量；若运行时缓存未更新，则直接回读最新 config.yaml 做兜底。
        latest_cfg = self._load_runtime_config()
        latest_llm = latest_cfg.get("llm") or {}
        latest_deepseek = latest_llm.get("deepseek") or {}
        latest_api = latest_llm.get("api") or {}
        api_key = (
            os.environ.get("DEEPSEEK_API_KEY")
            or ((self._qa_runtime_config.get("llm") or {}).get("deepseek") or {}).get("api_key", "")
            or latest_deepseek.get("api_key", "")
            or latest_api.get("api_key", "")
        )
        if not api_key:
            return json.dumps({"error": "DEEPSEEK_API_KEY 未配置，已停用豆包通道。"}, ensure_ascii=False)

        try:
            from openai import OpenAI

            client = OpenAI(api_key=api_key, base_url="https://api.deepseek.com")
            req = {
                "model": model,
                "messages": [
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_message},
                ],
                "temperature": temperature,
                "max_tokens": max_tokens,
            }
            if model in ("deepseek-v4-flash", "deepseek-v4-pro"):
                if no_thinking:
                    req["extra_body"] = {"thinking": {"type": "disabled"}}
                else:
                    req["reasoning_effort"] = reasoning_effort
                    req["extra_body"] = {"thinking": {"type": "enabled"}}
            resp = client.chat.completions.create(**req)
            content = (resp.choices[0].message.content or "").strip()
            usage = getattr(resp, "usage", None)
            usage_payload = self._normalize_ai_usage_payload(usage)
            self._record_ai_token_usage(purpose, model, usage_payload)
            if charge_quota:
                if purpose == "ai_qa":
                    self._increment_usage("ai_qa_used")
                else:
                    self._increment_usage("ai_parse_used")
            return json.dumps({"content": content, "usage": usage_payload}, ensure_ascii=False)
        except Exception as exc:
            return json.dumps({"error": str(exc)}, ensure_ascii=False)

    # ------------------------------------------------------------------
    #  文件解析
    # ------------------------------------------------------------------

    @staticmethod
    @staticmethod
    def _extract_ref_field_display_texts(path: str) -> list:
        """Extract display texts of Word REF fields whose target bookmark actually exists.

        A REF field in OOXML looks like:
          <w:r><w:fldChar w:fldCharType="begin"/></w:r>
          <w:r><w:instrText> REF bookmark \\h </w:instrText></w:r>
          <w:r><w:fldChar w:fldCharType="separate"/></w:r>
          <w:r><w:t>[1]</w:t></w:r>   ← display text we want
          <w:r><w:fldChar w:fldCharType="end"/></w:r>

        Only returns refs whose bookmark actually exists in the document; broken REF
        fields (pointing to non-existent bookmarks) are excluded so the user can
        re-adopt them to create correct fields with superscript and bookmarks.
        """
        import zipfile, re as _re
        import xml.etree.ElementTree as ET
        W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        results = []
        try:
            with zipfile.ZipFile(path, 'r') as z:
                if 'word/document.xml' not in z.namelist():
                    return results
                xml_bytes = z.read('word/document.xml')
            root = ET.fromstring(xml_bytes)
            # Collect all bookmark names that actually exist in the document
            existing_bookmarks = {
                bm.get(f'{{{W}}}name', '')
                for bm in root.findall(f'.//{{{W}}}bookmarkStart')
            }
            # Flatten all runs in document order
            all_runs = root.findall(f'.//{{{W}}}r')
            i = 0
            while i < len(all_runs):
                r = all_runs[i]
                fc = r.find(f'{{{W}}}fldChar')
                if fc is not None and fc.get(f'{{{W}}}fldCharType') == 'begin':
                    # Scan forward: collect instrText, look for REF, then get display text
                    instr = ''
                    is_ref_field = False
                    display_parts = []
                    in_display = False
                    j = i + 1
                    while j < len(all_runs):
                        rj = all_runs[j]
                        fci = rj.find(f'{{{W}}}fldChar')
                        if fci is not None:
                            ftype = fci.get(f'{{{W}}}fldCharType', '')
                            if ftype == 'separate':
                                is_ref_field = Api._is_ref_field(instr)
                                in_display = True
                            elif ftype == 'end':
                                if is_ref_field and display_parts:
                                    txt = ''.join(display_parts).strip()
                                    if txt:
                                        # Only mark as "already fielded" if the target
                                        # bookmark actually exists; otherwise exclude so
                                        # the user can re-adopt to create a proper field.
                                        bm_match = _re.search(r'REF\s+(\S+)', instr, _re.IGNORECASE)
                                        bm_name = bm_match.group(1) if bm_match else ''
                                        if not bm_name or bm_name in existing_bookmarks:
                                            results.append(txt)
                                i = j
                                break
                        instr_el = rj.find(f'{{{W}}}instrText')
                        if instr_el is not None and instr_el.text:
                            instr += instr_el.text
                        if in_display:
                            t_el = rj.find(f'{{{W}}}t')
                            if t_el is not None and t_el.text:
                                display_parts.append(t_el.text)
                        j += 1
                i += 1
        except Exception:
            pass
        return results

    @staticmethod
    def _is_ref_field(instr: str) -> bool:
        """Return True only for standalone REF fields (not PAGEREF, NOTEREF, etc.)."""
        import re
        # REF appears at start of instruction text (after optional whitespace)
        return bool(re.match(r'\s*REF\s', instr, re.IGNORECASE))

    @staticmethod
    def _parse_docx(path: str):
        """Return (elements, sections, styles).

        - elements: paragraph / table list (each paragraph carries fmt + runs for round-trip export).
        - sections: page geometry per section.
        - styles: metadata for Normal / Heading 1-3 (font name, size, bold, rFonts) so the exporter
          can rewrite ``word/styles.xml`` to match the source. python-docx auto-materialises any
          missing style, so always emitting a 4-entry dict keeps the downstream code simple.
        """
        from docx import Document
        from docx.oxml.ns import qn
        doc = Document(path)
        elements = []

        # ── 样式字体缓存：沿继承链解析每个样式的 eastAsia 字体和字号 ──
        _style_cache: dict = {}

        def _resolve_style_fonts(style) -> dict:
            """返回 {eastAsia, ascii, size_pt}，沿 base_style 链向上合并。"""
            if style is None:
                return {}
            key = style.name
            if key in _style_cache:
                return _style_cache[key]
            # 占位防止循环继承
            _style_cache[key] = {}
            parent = getattr(style, "base_style", None)
            resolved = dict(_resolve_style_fonts(parent))
            # 读本层 rFonts
            el = style.element
            r_pr = el.find(qn("w:rPr")) if el is not None else None
            rf = r_pr.find(qn("w:rFonts")) if r_pr is not None else None
            if rf is not None:
                ea = rf.get(qn("w:eastAsia"))
                asc = rf.get(qn("w:ascii"))
                if ea:  resolved["eastAsia"] = ea
                if asc: resolved["ascii"]    = asc
            # 读本层字号
            try:
                sz = style.font.size
                if sz is not None:
                    resolved["size_pt"] = float(sz.pt)
            except Exception:
                pass
            _style_cache[key] = resolved
            return resolved

        # 预热缓存（防止首次查询时递归深度问题）
        for _s in doc.styles:
            try:
                _resolve_style_fonts(_s)
            except Exception:
                pass

        def _para_fmt(para):
            fmt = para.paragraph_format
            def _tw(length):
                if length is None:
                    return None
                try:
                    return int(length.twips)
                except Exception:
                    return None
            rule_name = str(fmt.line_spacing_rule).split()[0] if fmt.line_spacing_rule is not None else None
            line_spacing = None
            line_spacing_twips = None
            if fmt.line_spacing is not None:
                if rule_name in ("EXACTLY", "AT_LEAST"):
                    line_spacing_twips = _tw(fmt.line_spacing)
                else:
                    try:
                        line_spacing = float(fmt.line_spacing)
                    except Exception:
                        line_spacing = None
            return {
                "style": para.style.name if para.style else None,
                "alignment": str(para.alignment).split()[0] if para.alignment is not None else None,
                "first_line_indent_twips": _tw(fmt.first_line_indent),
                "left_indent_twips": _tw(fmt.left_indent),
                "right_indent_twips": _tw(fmt.right_indent),
                "space_before_twips": _tw(fmt.space_before),
                "space_after_twips": _tw(fmt.space_after),
                "line_spacing_rule": rule_name,
                "line_spacing": line_spacing,
                "line_spacing_twips": line_spacing_twips,
            }

        def _runs(para):
            out = []
            # Check if paragraph contains Word fields (REF, TOC, etc.)
            para_xml_str = str(para._element.xml)
            has_field_codes = 'fldChar' in para_xml_str or 'instrText' in para_xml_str
            # 段落样式继承字体（当 run 没有显式设置时用作回退）
            style_fonts = _resolve_style_fonts(para.style) if para.style else {}

            for r in para.runs:
                r_pr = getattr(r._element, "rPr", None)
                r_fonts = getattr(r_pr, "rFonts", None) if r_pr is not None else None
                # run 显式值
                r_eastAsia = r_fonts.get(qn("w:eastAsia")) if r_fonts is not None else None
                r_ascii    = r_fonts.get(qn("w:ascii"))    if r_fonts is not None else None
                r_hansi    = r_fonts.get(qn("w:hAnsi"))    if r_fonts is not None else None
                r_cs       = r_fonts.get(qn("w:cs"))       if r_fonts is not None else None
                r_size     = float(r.font.size.pt) if r.font.size else None
                r_fname    = r.font.name
                # 回退到继承链
                eff_eastAsia = r_eastAsia or style_fonts.get("eastAsia")
                eff_ascii    = r_ascii    or style_fonts.get("ascii")
                eff_size     = r_size     or style_fonts.get("size_pt")
                run_info = {
                    "text": r.text or "",
                    "font_name": r_fname or eff_ascii or eff_eastAsia,
                    "font_size_pt": eff_size,
                    "bold": bool(r.bold) if r.bold is not None else None,
                    "italic": bool(r.italic) if r.italic is not None else None,
                    "underline": bool(r.underline) if r.underline is not None else None,
                    "font_ascii":   r_ascii    or eff_ascii,
                    "font_hansi":   r_hansi    or style_fonts.get("ascii"),
                    "font_eastAsia": eff_eastAsia,
                    "font_cs": r_cs,
                    "is_in_field": has_field_codes,
                }
                out.append(run_info)
            return out

        for element in doc.element.body:
            tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag

            if tag == "p":
                from docx.text.paragraph import Paragraph
                para = Paragraph(element, doc)
                text = para.text or ""
                text_stripped = text.strip()
                style_name = ((para.style.name if para.style else "") or "").lower()
                elem_type = "p"
                if style_name in ("heading 1", "标题 1", "title") or style_name.startswith("heading 1 "):
                    elem_type = "h1"
                elif style_name in ("heading 2", "标题 2") or style_name.startswith("heading 2 "):
                    elem_type = "h2"
                elif style_name in ("heading 3", "标题 3", "heading 4", "标题 4") or re.match(r"^heading [34]", style_name):
                    elem_type = "h3"
                elif re.match(r"^heading [5-9]", style_name):
                    elem_type = "h3"
                elif style_name in ("caption", "题注") or "题注" in style_name:
                    elem_type = "caption"
                elif "list bullet" in style_name or "无序列表" in style_name:
                    elem_type = "li"
                elif "list number" in style_name or "有序列表" in style_name:
                    elem_type = "li"
                elif re.match(r"^\[\d+\]", text_stripped):
                    elem_type = "ref"
                # Detect section break contained inside this paragraph's pPr (w:sectPr). Word
                # stores section breaks on the paragraph that ends each section except the last,
                # which is captured by the body-level sectPr.
                sect_pr = element.find(qn("w:pPr") + "/" + qn("w:sectPr")) if False else None
                try:
                    p_pr = element.find(qn("w:pPr"))
                    sect_pr = p_pr.find(qn("w:sectPr")) if p_pr is not None else None
                except Exception:
                    sect_pr = None
                section_break_type = None
                if sect_pr is not None:
                    type_el = sect_pr.find(qn("w:type"))
                    if type_el is not None:
                        section_break_type = type_el.get(qn("w:val")) or "nextPage"
                    else:
                        section_break_type = "nextPage"
                elements.append({
                    "type": elem_type,
                    "text": text,
                    "fmt": _para_fmt(para),
                    "runs": _runs(para),
                    "section_break": section_break_type,
                })

            elif tag == "tbl":
                from docx.table import Table
                try:
                    table = Table(element, doc)
                    rows = [[cell.text.strip().replace("\n", " ") for cell in row.cells] for row in table.rows]
                    if rows:
                        html = "<table>" + "".join(
                            "<tr>" + "".join(f"<{'th' if i==0 else 'td'}>{c}</{'th' if i==0 else 'td'}>" for c in row) + "</tr>"
                            for i, row in enumerate(rows)
                        ) + "</table>"
                        elements.append({"type": "table", "text": html})
                except Exception:
                    pass

        if not elements:
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    elements.append({"type": "p", "text": text, "fmt": _para_fmt(para), "runs": _runs(para)})

        # 第十八批：注入结构识别 metadata（cover/toc/heading/body/reference/caption ...）
        try:
            from core.document_structure import classify_dicts
            classify_dicts(elements)
        except Exception as _structure_err:
            print(f"[_parse_docx] classify_dicts 失败，已忽略: {_structure_err}")

        sections = []
        for s in doc.sections:
            def _tw(v):
                try:
                    return int(v.twips) if v is not None else None
                except Exception:
                    return None
            sections.append({
                "page_width_twips": _tw(s.page_width),
                "page_height_twips": _tw(s.page_height),
                "left_margin_twips": _tw(s.left_margin),
                "right_margin_twips": _tw(s.right_margin),
                "top_margin_twips": _tw(s.top_margin),
                "bottom_margin_twips": _tw(s.bottom_margin),
            })

        tracked_styles = ("Normal", "Heading 1", "Heading 2", "Heading 3")
        styles_meta: dict = {}
        for style_name in tracked_styles:
            try:
                style = doc.styles[style_name]
            except (KeyError, Exception):
                continue
            r_fonts_el = None
            r_pr = style.element.find(qn("w:rPr")) if style.element is not None else None
            if r_pr is not None:
                r_fonts_el = r_pr.find(qn("w:rFonts"))
            rfonts = {
                "ascii": r_fonts_el.get(qn("w:ascii")) if r_fonts_el is not None else None,
                "hAnsi": r_fonts_el.get(qn("w:hAnsi")) if r_fonts_el is not None else None,
                "eastAsia": r_fonts_el.get(qn("w:eastAsia")) if r_fonts_el is not None else None,
                "cs": r_fonts_el.get(qn("w:cs")) if r_fonts_el is not None else None,
            }
            try:
                size_pt = float(style.font.size.pt) if style.font.size else None
            except Exception:
                size_pt = None
            styles_meta[style_name] = {
                "font_name": style.font.name,
                "font_size_pt": size_pt,
                "bold": bool(style.font.bold) if style.font.bold is not None else None,
                "italic": bool(style.font.italic) if style.font.italic is not None else None,
                "underline": bool(style.font.underline) if style.font.underline is not None else None,
                "rfonts": rfonts,
            }

        return elements, sections, styles_meta

    @staticmethod
    def _parse_text(path: str) -> list:
        content = None
        for enc in ("utf-8", "utf-8-sig", "gbk", "gb2312", "latin-1"):
            try:
                with open(path, "r", encoding=enc) as fh:
                    content = fh.read()
                break
            except (UnicodeDecodeError, UnicodeError):
                continue
        if content is None:
            raise RuntimeError(f"无法解码文件: {path}")

        elements = []
        lines = content.splitlines()
        i = 0
        while i < len(lines):
            line = lines[i].strip()
            if not line:
                i += 1
                continue
            if line.startswith("#### ") or line.startswith("### "):
                elements.append({"type": "h3", "text": line.lstrip("#").strip()})
            elif line.startswith("## "):
                elements.append({"type": "h2", "text": line[3:].strip()})
            elif line.startswith("# "):
                elements.append({"type": "h1", "text": line[2:].strip()})
            elif re.match(r"^[\-\*]\s+", line):
                elements.append({"type": "li", "text": "• " + re.sub(r"^[\-\*]\s+", "", line)})
            elif re.match(r"^\d+\.\s+", line):
                elements.append({"type": "li", "text": line})
            elif re.match(r"^\[\d+\]", line):
                elements.append({"type": "ref", "text": line})
            elif re.match(r"^(图|表|Figure|Table)\s*\d", line):
                elements.append({"type": "caption", "text": line})
            else:
                elements.append({"type": "p", "text": line})
            i += 1
        return elements

    # ------------------------------------------------------------------
    #  Word 导出（修复：安全设置中文字体）
    # ------------------------------------------------------------------

    @staticmethod
    def _set_run_font(run, cn_font: str, en_font: str, size_pt: float) -> None:
        """安全设置 run 的中英文字体和字号，避免 rPr 为 None 导致乱码。"""
        from docx.shared import Pt
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        run.font.name = en_font
        run.font.size = Pt(size_pt)

        rPr = run._element.get_or_add_rPr()
        # 移除已有 rFonts，重新构造，保证覆盖
        existing = rPr.find(qn('w:rFonts'))
        if existing is not None:
            rPr.remove(existing)
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), en_font)
        rFonts.set(qn('w:hAnsi'), en_font)
        rFonts.set(qn('w:eastAsia'), cn_font)
        rFonts.set(qn('w:cs'), cn_font)
        rPr.insert(0, rFonts)

    @staticmethod
    def _write_docx(path: str, html_content: str, format_params: dict = None) -> None:
        from docx import Document
        from docx.shared import Pt, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        fp = format_params or {}
        h1_font  = fp.get("h1Font",   "黑体");  h1_size  = float(fp.get("h1Size",  16))
        h2_font  = fp.get("h2Font",   "黑体");  h2_size  = float(fp.get("h2Size",  14))
        h3_font  = fp.get("h3Font",   "黑体");  h3_size  = float(fp.get("h3Size",  12))
        body_font= fp.get("bodyFont", "宋体");  body_size= float(fp.get("bodySize", 12))
        west_font= fp.get("westFont", "Times New Roman")
        indent   = float(fp.get("indent",     2))
        line_h   = float(fp.get("lineHeight", 1.5))
        align    = fp.get("align", "justify")
        margin_t = float(fp.get("marginTop",    2.5))
        margin_b = float(fp.get("marginBottom", 2.5))
        margin_l = float(fp.get("marginLeft",   3.0))
        margin_r = float(fp.get("marginRight",  2.5))

        doc = Document()
        sec = doc.sections[0]
        sec.page_width    = Cm(21);   sec.page_height     = Cm(29.7)
        sec.top_margin    = Cm(margin_t); sec.bottom_margin = Cm(margin_b)
        sec.left_margin   = Cm(margin_l); sec.right_margin  = Cm(margin_r)

        set_font = Api._set_run_font  # 使用安全字体设置方法

        for elem in Api._html_to_elements(html_content):
            t, text = elem["type"], elem["text"]

            if t in ("h1", "h2", "h3"):
                level = int(t[1])
                p = doc.add_heading(text, level=level)
                cn = h1_font if t == "h1" else (h2_font if t == "h2" else h3_font)
                sz = h1_size if t == "h1" else (h2_size if t == "h2" else h3_size)
                for run in p.runs:
                    set_font(run, cn, west_font, sz)

            elif t == "table":
                from html.parser import HTMLParser

                class _TableParser(HTMLParser):
                    def __init__(self):
                        super().__init__(); self.rows = []; self._row = []; self._cell = ""; self._in = False
                    def handle_starttag(self, tag, attrs):
                        if tag in ("th", "td"): self._in = True; self._cell = ""
                    def handle_endtag(self, tag):
                        if tag in ("th", "td"): self._row.append(self._cell); self._in = False
                        elif tag == "tr": self.rows.append(self._row); self._row = []
                    def handle_data(self, data):
                        if self._in: self._cell += data

                tp = _TableParser(); tp.feed(text)
                if tp.rows:
                    tbl = doc.add_table(rows=len(tp.rows), cols=len(tp.rows[0]))
                    tbl.style = "Table Grid"
                    for ri, row_data in enumerate(tp.rows):
                        for ci, cell_text in enumerate(row_data):
                            if ci < len(tbl.rows[ri].cells):
                                tbl.rows[ri].cells[ci].text = cell_text

            elif t == "ref":
                p = doc.add_paragraph(text)
                for run in p.runs:
                    set_font(run, body_font, west_font, 10.5)

            elif t == "caption":
                p = doc.add_paragraph(text)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    set_font(run, h3_font, west_font, 10.5)

            elif t == "li":
                p = doc.add_paragraph(text, style="List Bullet")
                p.paragraph_format.line_spacing = line_h
                for run in p.runs:
                    set_font(run, body_font, west_font, body_size)

            else:
                p = doc.add_paragraph(text)
                p.paragraph_format.first_line_indent = Cm(indent * 0.35 * body_size / 12)
                p.paragraph_format.line_spacing = line_h
                align_map = {"justify": WD_ALIGN_PARAGRAPH.JUSTIFY, "left": WD_ALIGN_PARAGRAPH.LEFT,
                             "center": WD_ALIGN_PARAGRAPH.CENTER, "right": WD_ALIGN_PARAGRAPH.RIGHT}
                p.alignment = align_map.get(align, WD_ALIGN_PARAGRAPH.JUSTIFY)
                for run in p.runs:
                    set_font(run, body_font, west_font, body_size)

        doc.save(path)

    @staticmethod
    def _html_to_elements(html: str) -> list:
        cleaned = re.sub(r"<br\s*/?>", "\n", html)
        pattern = re.compile(r"<(h[1-6]|p|div|caption|li)[^>]*>(.*?)</\1>|(<table>.*?</table>)",
                              re.DOTALL | re.IGNORECASE)
        elements = []
        for match in pattern.finditer(cleaned):
            if match.group(3):
                elements.append({"type": "table", "text": match.group(3)})
                continue
            tag = match.group(1)
            text = re.sub(r"<[^>]+>", "", match.group(2)).strip()
            if not text:
                continue
            if tag in ("h1", "h2", "h3", "h4", "h5", "h6"):
                elements.append({"type": tag if tag in ("h1","h2","h3") else "h3", "text": text})
            elif tag == "caption":
                elements.append({"type": "caption", "text": text})
            elif tag == "li":
                elements.append({"type": "li", "text": text})
            else:
                elements.append({"type": "p", "text": text})
        return elements
