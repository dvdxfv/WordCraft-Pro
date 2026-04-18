"""
DOCX 解析器

解析 .docx 文件，提取段落、标题、表格、图片、样式、页眉页脚等信息，
转换为统一的 DocumentModel。

注意：.doc（旧格式）需先通过 LibreOffice 转换为 .docx。
"""

from __future__ import annotations

import os
import re
import copy
from typing import Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from parsers.base import BaseParser
from core.document_model import (
    DocumentModel, DocElement, ElementType, FontStyle, ParagraphStyle,
    Alignment, LineSpacingType, NumberingType,
    TableData, TableCell, ImageData,
    PageSetup, SectionConfig, HeaderFooterConfig, PageNumberConfig, PageNumberFormat,
)


# Word 对齐方式映射
_ALIGN_MAP = {
    WD_ALIGN_PARAGRAPH.LEFT: Alignment.LEFT,
    WD_ALIGN_PARAGRAPH.CENTER: Alignment.CENTER,
    WD_ALIGN_PARAGRAPH.RIGHT: Alignment.RIGHT,
    WD_ALIGN_PARAGRAPH.JUSTIFY: Alignment.JUSTIFY,
}


class DocxParser(BaseParser):
    """DOCX 文件解析器"""

    supported_extensions = [".docx"]

    def parse(self, file_path: str, **kwargs) -> DocumentModel:
        """解析 .docx 文件"""
        abs_path = self._validate_file(file_path)
        doc = Document(abs_path)

        model = DocumentModel(
            source_file=abs_path,
            source_format="docx",
        )

        # 提取元数据
        self._extract_metadata(doc, model)

        # 提取页面设置（取第一个节的设置作为默认）
        self._extract_page_setup(doc, model)

        # 提取多节配置（页眉页脚）
        self._extract_sections(doc, model)

        # 提取文档内容
        self._extract_elements(doc, model)

        # 提取表格
        self._extract_tables(doc, model)

        # 提取图片
        self._extract_images(doc, model)

        return model

    # ---- 元数据 ----

    def _extract_metadata(self, doc: Document, model: DocumentModel):
        """提取文档元数据"""
        core_props = doc.core_properties
        if core_props.title:
            model.title = core_props.title
        if core_props.author:
            model.author = core_props.author
        if core_props.subject:
            model.subject = core_props.subject
        if core_props.keywords:
            model.keywords = core_props.keywords
        if core_props.created:
            model.created_date = core_props.created.isoformat()

    # ---- 页面设置 ----

    def _extract_page_setup(self, doc: Document, model: DocumentModel):
        """提取页面设置"""
        if not doc.sections:
            return
        section = doc.sections[0]

        # 页面尺寸（EMU → cm，1cm = 360000 EMU）
        emu_to_cm = 1.0 / 360000.0
        page_width = section.page_width * emu_to_cm
        page_height = section.page_height * emu_to_cm

        # 检测纸张大小
        paper = self._detect_paper_size(page_width, page_height)

        model.page_setup = PageSetup(
            paper_size=paper,
            orientation="landscape" if page_width > page_height else "portrait",
            margin_top_cm=section.top_margin * emu_to_cm,
            margin_bottom_cm=section.bottom_margin * emu_to_cm,
            margin_left_cm=section.left_margin * emu_to_cm,
            margin_right_cm=section.right_margin * emu_to_cm,
            gutter_cm=section.gutter * emu_to_cm,
            header_distance_cm=section.header_distance * emu_to_cm,
            footer_distance_cm=section.footer_distance * emu_to_cm,
        )

    @staticmethod
    def _detect_paper_size(width_cm: float, height_cm: float) -> str:
        """根据宽高检测纸张大小"""
        sizes = {
            "A4": (21.0, 29.7),
            "A3": (29.7, 42.0),
            "B5": (17.6, 25.0),
            "Letter": (21.59, 27.94),
        }
        for name, (w, h) in sizes.items():
            if abs(width_cm - w) < 0.5 and abs(height_cm - h) < 0.5:
                return name
            if abs(width_cm - h) < 0.5 and abs(height_cm - w) < 0.5:
                return name
        return f"{width_cm:.1f}x{height_cm:.1f}cm"

    # ---- 多节配置（页眉页脚） ----

    def _extract_sections(self, doc: Document, model: DocumentModel):
        """提取多节页眉页脚配置"""
        for i, section in enumerate(doc.sections):
            section_config = SectionConfig()

            # 首页不同
            sect_pr = section._sectPr
            title_pg = sect_pr.find(qn("w:titlePg"))
            section_config.first_page_different = title_pg is not None

            # 页眉
            header = section.header
            if header and not header.is_linked_to_previous:
                header_text = self._extract_header_footer_text(header)
                if header_text:
                    section_config.header = HeaderFooterConfig(text=header_text)
                    # 尝试提取页眉字体样式
                    self._extract_hf_style(header, section_config.header)

            # 首页页眉
            if section_config.first_page_different:
                first_header = section.first_page_header
                if first_header:
                    first_text = self._extract_header_footer_text(first_header)
                    if first_text:
                        section_config.first_page_header = HeaderFooterConfig(text=first_text)
                        self._extract_hf_style(first_header, section_config.first_page_header)

            # 页脚和页码
            footer = section.footer
            if footer and not footer.is_linked_to_previous:
                footer_text = self._extract_header_footer_text(footer)
                # 检测页码域
                has_page_num = self._has_page_number_field(footer)
                if has_page_num:
                    section_config.page_number = PageNumberConfig(enabled=True)
                    # 检测页码格式
                    pn_format = self._detect_page_number_format(footer)
                    section_config.page_number.number_format = pn_format
                    # 检测起始页码
                    pg_num_type = sect_pr.find(qn("w:pgNumType"))
                    if pg_num_type is not None:
                        start = pg_num_type.get(qn("w:start"))
                        if start:
                            section_config.page_number.start_from = int(start)
                        fmt = pg_num_type.get(qn("w:fmt"))
                        if fmt:
                            fmt_map = {"upperRoman": PageNumberFormat.UPPER_ROMAN,
                                       "lowerRoman": PageNumberFormat.LOWER_ROMAN,
                                       "decimal": PageNumberFormat.ARABIC}
                            section_config.page_number.number_format = fmt_map.get(fmt, PageNumberFormat.ARABIC)

            # 节名称（尝试从页眉推断）
            if section_config.header and section_config.header.text:
                section_config.name = section_config.header.text.strip()
            elif section_config.first_page_header and section_config.first_page_header.text:
                section_config.name = section_config.first_page_header.text.strip()
            else:
                section_config.name = f"第{i + 1}节"

            model.sections.append(section_config)

    def _extract_header_footer_text(self, header_or_footer) -> str:
        """提取页眉/页脚的纯文本"""
        texts = []
        for para in header_or_footer.paragraphs:
            text = para.text.strip()
            if text:
                texts.append(text)
        return " ".join(texts)

    def _extract_hf_style(self, header_or_footer, config: HeaderFooterConfig):
        """提取页眉/页脚的字体样式"""
        for para in header_or_footer.paragraphs:
            for run in para.runs:
                if run.font.name:
                    config.font_name = run.font.name
                if run.font.size:
                    config.font_size_pt = run.font.size.pt
                if run.font.bold:
                    config.bold = run.font.bold
                if run.font.size or run.font.name:
                    break
            if config.font_name or config.font_size_pt:
                break

    def _has_page_number_field(self, footer) -> bool:
        """检测页脚中是否包含页码域"""
        xml = footer._element.xml
        return "w:fldChar" in xml or "PAGE" in xml

    def _detect_page_number_format(self, footer) -> PageNumberFormat:
        """检测页码格式"""
        xml = footer._element.xml
        if "upperRoman" in xml:
            return PageNumberFormat.UPPER_ROMAN
        if "lowerRoman" in xml:
            return PageNumberFormat.LOWER_ROMAN
        return PageNumberFormat.ARABIC

    # ---- 文档内容 ----

    def _extract_elements(self, doc: Document, model: DocumentModel):
        """提取文档中的段落和标题"""
        for para in doc.paragraphs:
            element = self._parse_paragraph(para)
            if element is not None:
                model.elements.append(element)

    def _parse_paragraph(self, para) -> Optional[DocElement]:
        """解析单个段落为 DocElement"""
        text = para.text.strip()
        if not text and not para.runs:
            # 保留分页符/分节符
            if self._is_page_break(para):
                return DocElement(element_type=ElementType.PAGE_BREAK)
            if self._is_section_break(para):
                return DocElement(element_type=ElementType.SECTION_BREAK)
            return None

        # 判断元素类型
        elem_type, level = self._detect_element_type(para, text)

        # 提取字体样式
        font_style = self._extract_font_style(para)

        # 提取段落样式
        para_style = self._extract_paragraph_style(para)

        # 提取编号
        numbering_text = self._extract_numbering_text(para)

        element = DocElement(
            element_type=elem_type,
            content=text,
            level=level,
            font_style=font_style,
            paragraph_style=para_style,
            style_name=para.style.name if para.style else "",
            numbering_text=numbering_text,
        )

        return element

    def _detect_element_type(self, para, text: str) -> tuple[ElementType, int]:
        """检测段落类型（标题/正文/参考文献等）"""
        style_name = (para.style.name or "").lower()

        # 标题检测
        if style_name.startswith("heading") or style_name.startswith("标题"):
            # 提取标题级别
            level = 0
            match = re.search(r"(\d+)", style_name)
            if match:
                level = int(match.group(1))
            return ElementType.HEADING, level

        if style_name in ("title", "标题"):
            return ElementType.HEADING, 0

        # 参考文献检测
        if style_name.startswith("bibliography") or "参考" in style_name:
            return ElementType.REFERENCE, 0

        # 题注检测
        if style_name.startswith("caption") or "题注" in style_name:
            return ElementType.CAPTION, 0

        # 通过内容特征检测
        if re.match(r"^\[[\d,\-–]+\]", text):
            return ElementType.REFERENCE, 0

        if re.match(r"^(图|表|Figure|Table)\s*[\d\-–]", text):
            return ElementType.CAPTION, 0

        # 默认为正文段落
        return ElementType.PARAGRAPH, 0

    def _extract_font_style(self, para) -> FontStyle:
        """提取段落的字体样式（取第一个 run 的样式）"""
        style = FontStyle()
        for run in para.runs:
            if run.font.name:
                # 区分中英文字体
                rpr = run._element.find(qn("w:rPr"))
                if rpr is not None:
                    ea_font = rpr.find(qn("w:rFonts"))
                    if ea_font is not None:
                        style.font_name_cn = ea_font.get(qn("w:eastAsia")) or ""
                        style.font_name_en = ea_font.get(qn("w:ascii")) or ea_font.get(qn("w:hAnsi")) or ""
                if not style.font_name_cn and not style.font_name_en:
                    style.font_name_en = run.font.name
            if run.font.size:
                style.font_size_pt = run.font.size.pt
            if run.font.bold is not None:
                style.bold = run.font.bold
            if run.font.italic is not None:
                style.italic = run.font.italic
            if run.font.underline is not None:
                style.underline = run.font.underline
            if run.font.color and run.font.color.rgb:
                style.color = str(run.font.color.rgb)
            if run.font.superscript:
                style.super_script = True
            if run.font.subscript:
                style.sub_script = True
            # 取第一个有效 run 的样式即可
            if style.font_name_cn or style.font_name_en or style.font_size_pt:
                break

        # 如果 run 级别没有提取到，尝试段落级别
        if not style.font_name_cn and not style.font_name_en:
            pf = para.style.font if para.style else None
            if pf:
                if pf.name:
                    style.font_name_en = pf.name
                if pf.size:
                    style.font_size_pt = pf.size.pt

        return style

    def _extract_paragraph_style(self, para) -> ParagraphStyle:
        """提取段落格式"""
        pf = para.paragraph_format
        style = ParagraphStyle()

        # 对齐
        if pf.alignment is not None:
            style.alignment = _ALIGN_MAP.get(pf.alignment, Alignment.LEFT)

        # 缩进
        if pf.first_line_indent is not None:
            # EMU → cm
            style.first_indent_cm = pf.first_line_indent / 360000.0
        if pf.left_indent is not None:
            style.left_indent_cm = pf.left_indent / 360000.0
        if pf.right_indent is not None:
            style.right_indent_cm = pf.right_indent / 360000.0

        # 行距
        if pf.line_spacing is not None:
            if pf.line_spacing_rule is not None:
                rule = pf.line_spacing_rule
                # python-docx 的 line_spacing_rule 是枚举
                rule_name = rule.name if hasattr(rule, "name") else str(rule)
                if "EXACTLY" in rule_name or "AT_LEAST" in rule_name:
                    style.line_spacing_type = LineSpacingType.EXACT
                    style.line_spacing_value = pf.line_spacing.pt if hasattr(pf.line_spacing, "pt") else pf.line_spacing
                else:
                    style.line_spacing_type = LineSpacingType.MULTIPLE
                    style.line_spacing_value = pf.line_spacing
            else:
                style.line_spacing_type = LineSpacingType.MULTIPLE
                style.line_spacing_value = pf.line_spacing

        # 段前段后
        if pf.space_before is not None:
            style.space_before_pt = pf.space_before.pt if hasattr(pf.space_before, "pt") else pf.space_before
        if pf.space_after is not None:
            style.space_after_pt = pf.space_after.pt if hasattr(pf.space_after, "pt") else pf.space_after

        # 分页控制
        style.keep_with_next = pf.keep_with_next
        style.keep_lines_together = pf.keep_together
        style.page_break_before = pf.page_break_before

        return style

    def _extract_numbering_text(self, para) -> str:
        """提取段落的编号文本"""
        # 检查段落是否有编号
        numPr = para._element.find(qn("w:pPr"))
        if numPr is not None:
            numPr = numPr.find(qn("w:numPr"))
            if numPr is not None:
                # 尝试从段落的 numId 获取编号
                # python-docx 没有直接API，通过XML提取
                ilvl = numPr.find(qn("w:ilvl"))
                if ilvl is not None:
                    # 编号文本通常在段落的实际文本前面
                    text = para.text
                    # 尝试匹配常见编号模式
                    match = re.match(r"^([一二三四五六七八九十]+、|\d+(?:\.\d+)*\.?|\(\d+\)|\([一二三四五六七八九十]+\))\s*", text)
                    if match:
                        return match.group(1)
        return ""

    @staticmethod
    def _is_page_break(para) -> bool:
        """检测是否为分页符"""
        for run in para.runs:
            if run._element.xml.find("w:br") != -1:
                br = run._element.find(qn("w:br"))
                if br is not None and br.get(qn("w:type")) in ("page", None):
                    return True
        # 也检查段落级别的分页
        pPr = para._element.find(qn("w:pPr"))
        if pPr is not None:
            for child in pPr:
                if "pageBreakBefore" in child.tag:
                    return True
        return False

    @staticmethod
    def _is_section_break(para) -> bool:
        """检测是否为分节符"""
        pPr = para._element.find(qn("w:pPr"))
        if pPr is not None:
            sectPr = pPr.find(qn("w:sectPr"))
            if sectPr is not None:
                return True
        return False

    # ---- 表格 ----

    def _extract_tables(self, doc: Document, model: DocumentModel):
        """提取文档中的表格"""
        for table in doc.tables:
            table_data = self._parse_table(table)
            if table_data and table_data.num_rows > 0:
                model.tables.append(table_data)

    def _parse_table(self, table) -> Optional[TableData]:
        """解析单个表格"""
        rows = []
        for row_idx, row in enumerate(table.rows):
            cells = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                is_header = (row_idx == 0)

                # 提取单元格字体样式
                font_style = FontStyle()
                para_style = ParagraphStyle()
                if cell.paragraphs:
                    first_para = cell.paragraphs[0]
                    for run in first_para.runs:
                        if run.font.name:
                            font_style.font_name_en = run.font.name
                        if run.font.size:
                            font_style.font_size_pt = run.font.size.pt
                        break
                    pf = first_para.paragraph_format
                    if pf.alignment is not None:
                        para_style.alignment = _ALIGN_MAP.get(pf.alignment, Alignment.LEFT)

                cells.append(TableCell(
                    content=cell_text,
                    font_style=font_style,
                    paragraph_style=para_style,
                    is_header=is_header,
                ))
            if cells:
                rows.append(cells)

        if not rows:
            return None

        return TableData(rows=rows)

    # ---- 图片 ----

    def _extract_images(self, doc: Document, model: DocumentModel):
        """提取文档中的图片信息"""
        # 遍历文档中的 inline shapes
        for rel in doc.part.rels.values():
            if "image" in rel.reltype:
                image_data = ImageData(
                    file_path=rel.target_ref if hasattr(rel, "target_ref") else "",
                    metadata={"rel_id": rel.rId},
                )
                model.images.append(image_data)
