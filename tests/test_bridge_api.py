#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
WordCraft Pro Bridge API 单元测试

测试 pywebview 与后端的 API 桥接：
- 文件操作 (openFile, saveFile, exportDocx)
- 用户认证 (login, logout)
- 质量检查 (runQA)
- 交叉引用 (runXRef)
- 排版引擎 (applyFormat)
- 文档管理 (saveDocument, loadDocument)
- Token 管理 (getTokenUsage)
- 模板管理 (getUserTemplates, uploadTemplate, deleteTemplate)
"""

import json
import os
import sys
import tempfile
from unittest.mock import Mock, MagicMock, patch

import pytest

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from app import Api


# ============================================================
# 辅助函数
# ============================================================

def create_mock_window():
    """创建 Mock pywebview 窗口"""
    window = Mock()
    window.create_file_dialog = Mock()
    return window


# ============================================================
# 文件操作 API 测试
# ============================================================

class TestFileOperationAPI:
    """文件操作 API 测试"""

    def setup_method(self):
        """测试前准备"""
        self.api = Api(supabase_enabled=False)
        self.mock_window = create_mock_window()
        self.api.set_window(self.mock_window)
        self.temp_dir = tempfile.mkdtemp()

    def teardown_method(self):
        """测试后清理"""
        import shutil
        shutil.rmtree(self.temp_dir, ignore_errors=True)

    def test_openFile_cancelled(self):
        """测试打开文件 - 用户取消"""
        self.mock_window.create_file_dialog.return_value = None

        result = self.api.openFile()
        data = json.loads(result)

        assert data["cancelled"] is True
        print("✓ openFile 取消操作正确处理")

    def test_openFile_with_docx(self):
        """测试打开 DOCX 文件"""
        # 创建测试 DOCX 文件
        from docx import Document
        test_file = os.path.join(self.temp_dir, "test.docx")
        doc = Document()
        doc.add_paragraph("测试段落")
        doc.save(test_file)

        self.mock_window.create_file_dialog.return_value = [test_file]

        result = self.api.openFile()
        data = json.loads(result)

        assert "name" in data
        assert data["name"] == "test.docx"
        assert data["type"] == "docx"
        assert isinstance(data["elements"], list)
        print("✓ openFile 正确解析 DOCX 文件")

    def test_openFile_with_txt(self):
        """测试打开 TXT 文件"""
        # 创建测试 TXT 文件
        test_file = os.path.join(self.temp_dir, "test.txt")
        with open(test_file, "w", encoding="utf-8") as f:
            f.write("# 标题\n\n测试正文\n")

        self.mock_window.create_file_dialog.return_value = [test_file]

        result = self.api.openFile()
        data = json.loads(result)

        assert data["type"] == "txt"
        assert isinstance(data["elements"], list)
        print("✓ openFile 正确解析 TXT 文件")

    def test_openFile_with_markdown(self):
        """测试打开 MD 文件"""
        test_file = os.path.join(self.temp_dir, "test.md")
        with open(test_file, "w", encoding="utf-8") as f:
            f.write("# 标题\n\n- 列表项\n")

        self.mock_window.create_file_dialog.return_value = [test_file]

        result = self.api.openFile()
        data = json.loads(result)

        assert data["type"] == "md"
        assert isinstance(data["elements"], list)
        print("✓ openFile 正确解析 MD 文件")

    def test_saveFile_success(self):
        """测试保存文件"""
        save_path = os.path.join(self.temp_dir, "saved.html")
        self.mock_window.create_file_dialog.return_value = save_path

        result = self.api.saveFile("<h1>测试</h1>")
        data = json.loads(result)

        assert data["success"] is True
        assert os.path.exists(save_path)
        print("✓ saveFile 成功保存文件")

    def test_saveFile_cancelled(self):
        """测试保存文件 - 用户取消"""
        self.mock_window.create_file_dialog.return_value = None

        result = self.api.saveFile("<h1>测试</h1>")
        data = json.loads(result)

        assert data["cancelled"] is True
        print("✓ saveFile 取消操作正确处理")

    def test_exportDocx_success(self):
        """测试导出 DOCX"""
        export_path = os.path.join(self.temp_dir, "exported.docx")
        self.mock_window.create_file_dialog.return_value = export_path

        result = self.api.exportDocx("<h1>测试标题</h1><p>测试内容</p>")
        data = json.loads(result)

        assert data["success"] is True
        assert os.path.exists(export_path)
        print("✓ exportDocx 成功导出文件")

    def test_exportDocx_with_format_params(self):
        """测试导出 DOCX - 带格式参数"""
        export_path = os.path.join(self.temp_dir, "formatted.docx")
        self.mock_window.create_file_dialog.return_value = export_path

        format_params = json.dumps({
            "font": "宋体",
            "font_size": 12,
            "margin_top": 2.5
        })

        result = self.api.exportDocx("<p>内容</p>", format_params)
        data = json.loads(result)

        assert data["success"] is True
        print("✓ exportDocx 正确处理格式参数")

    def test_getSystemInfo(self):
        """测试获取系统信息"""
        result = self.api.getSystemInfo()
        data = json.loads(result)

        assert "version" in data
        assert "platform" in data
        assert "python" in data
        assert data["version"] == "1.0.0"
        print("✓ getSystemInfo 返回正确的系统信息")


# ============================================================
# 用户认证 API 测试
# ============================================================

class TestAuthenticationAPI:
    """用户认证 API 测试"""

    def setup_method(self):
        """测试前准备"""
        self.api = Api(supabase_enabled=False)
        self.api.set_window(create_mock_window())

    def test_login_local_mode(self):
        """测试本地模式登录"""
        result = self.api.login("test@example.com", "password123")
        data = json.loads(result)

        assert data["success"] is True
        assert "user" in data
        assert data["user"]["email"] == "test@example.com"
        assert "token" in data
        print("✓ login 本地模式成功")

    def test_login_stores_session(self):
        """测试登录后会话存储"""
        self.api.login("user@example.com", "password")

        assert self.api._session.get("user_id") is not None
        assert self.api._session.get("user_info") is not None
        print("✓ login 正确存储会话信息")

    def test_logout_clears_session(self):
        """测试登出清空会话"""
        # 先登录
        self.api.login("user@example.com", "password")
        assert self.api._session.get("user_id") is not None

        # 然后登出
        result = self.api.logout()
        data = json.loads(result)

        assert data["success"] is True
        assert self.api._session == {}
        print("✓ logout 正确清空会话")


# ============================================================
# 质量检查 API 测试
# ============================================================

class TestQAAPI:
    """质量检查 API 测试"""

    def setup_method(self):
        """测试前准备"""
        self.api = Api(supabase_enabled=False)
        self.api.set_window(create_mock_window())

    def test_runQA_returns_json(self):
        """测试质量检查返回 JSON"""
        result = self.api.runQA("这是一个正确的文本。")
        data = json.loads(result)

        assert "success" in data
        assert isinstance(data, dict)
        print("✓ runQA 返回有效的 JSON 格式")

    def test_runQA_with_categories(self):
        """测试质量检查 - 指定检查类别"""
        categories = json.dumps(["typo", "consistency"])
        result = self.api.runQA("测试文本", categories)
        data = json.loads(result)

        assert isinstance(data, dict)
        print("✓ runQA 正确处理类别参数")

    def test_runQA_returns_issues_structure(self):
        """测试质量检查返回问题结构"""
        result = self.api.runQA("请输入正确的帐号", '["typo"]')
        data = json.loads(result)

        if data.get("success") and data.get("issues"):
            for issue in data["issues"]:
                assert "category" in issue
                assert "severity" in issue
                assert "title" in issue

        assert "stats" in data
        print("✓ runQA 返回正确的问题结构")


# ============================================================
# 交叉引用 API 测试
# ============================================================

class TestCrossRefAPI:
    """交叉引用 API 测试"""

    def setup_method(self):
        """测试前准备"""
        self.api = Api(supabase_enabled=False)
        self.api.set_window(create_mock_window())

    def test_runXRef_returns_json(self):
        """测试交叉引用返回 JSON"""
        result = self.api.runXRef("如图3-1所示")
        data = json.loads(result)

        assert isinstance(data, dict)
        print("✓ runXRef 返回有效的 JSON 格式")

    def test_runXRef_returns_structure(self):
        """测试交叉引用返回结构"""
        result = self.api.runXRef("第一章\n如图1-1所示")
        data = json.loads(result)

        assert "targets" in data or "references" in data or "results" in data
        print("✓ runXRef 返回正确的结构")


# ============================================================
# 排版引擎 API 测试
# ============================================================

class TestFormattingAPI:
    """排版引擎 API 测试"""

    def setup_method(self):
        """测试前准备"""
        self.api = Api(supabase_enabled=False)
        self.api.set_window(create_mock_window())

    def test_applyFormat_returns_json(self):
        """测试排版返回 JSON"""
        result = self.api.applyFormat("<p>内容</p>")
        data = json.loads(result)

        assert isinstance(data, dict)
        print("✓ applyFormat 返回有效的 JSON 格式")

    def test_applyFormat_with_rules(self):
        """测试排版 - 带规则参数"""
        rules = json.dumps({
            "font": "宋体",
            "font_size": 12
        })
        result = self.api.applyFormat("<p>内容</p>", rules)
        data = json.loads(result)

        assert isinstance(data, dict)
        print("✓ applyFormat 正确处理规则参数")


# ============================================================
# 文档管理 API 测试
# ============================================================

class TestDocumentManagementAPI:
    """文档管理 API 测试"""

    def setup_method(self):
        """测试前准备"""
        self.api = Api(supabase_enabled=False)
        self.api.set_window(create_mock_window())

    def test_saveDocument_returns_json(self):
        """测试保存文档返回 JSON"""
        result = self.api.saveDocument("<p>测试内容</p>", "测试文档")
        data = json.loads(result)

        assert "success" in data
        assert isinstance(data, dict)
        print("✓ saveDocument 返回有效的 JSON 格式")

    def test_saveDocument_local_mode(self):
        """测试文档保存 - 本地模式"""
        result = self.api.saveDocument("<p>内容</p>", "文档标题")
        data = json.loads(result)

        assert data["success"] is True
        print("✓ saveDocument 本地模式成功保存")

    def test_loadDocument_returns_json(self):
        """测试加载文档返回 JSON"""
        result = self.api.loadDocument()
        data = json.loads(result)

        assert isinstance(data, dict)
        print("✓ loadDocument 返回有效的 JSON 格式")


# ============================================================
# Token 管理 API 测试
# ============================================================

class TestTokenManagementAPI:
    """Token 管理 API 测试"""

    def setup_method(self):
        """测试前准备"""
        self.api = Api(supabase_enabled=False)
        self.api.set_window(create_mock_window())

    def test_getTokenUsage_returns_json(self):
        """测试 Token 用量返回 JSON"""
        result = self.api.getTokenUsage()
        data = json.loads(result)

        assert "token_quota" in data
        assert "token_used" in data
        assert isinstance(data, dict)
        print("✓ getTokenUsage 返回有效的 JSON 格式")

    def test_getTokenUsage_structure(self):
        """测试 Token 用量返回结构"""
        result = self.api.getTokenUsage()
        data = json.loads(result)

        assert "token_quota" in data
        assert "token_used" in data
        assert "token_remaining" in data
        assert "usage_percentage" in data
        print("✓ getTokenUsage 返回正确的结构")


# ============================================================
# 模板管理 API 测试
# ============================================================

class TestTemplateManagementAPI:
    """模板管理 API 测试"""

    def setup_method(self):
        """测试前准备"""
        self.api = Api(supabase_enabled=False)
        self.api.set_window(create_mock_window())
        self.temp_dir = tempfile.mkdtemp()

    def teardown_method(self):
        """测试后清理"""
        import shutil
        shutil.rmtree(self.temp_dir, ignore_errors=True)

    def test_getUserTemplates_returns_json(self):
        """测试获取用户模板返回 JSON"""
        result = self.api.getUserTemplates()
        data = json.loads(result)

        assert "success" in data
        assert "templates" in data
        print("✓ getUserTemplates 返回有效的 JSON 格式")

    def test_getUserTemplates_returns_list(self):
        """测试获取用户模板返回列表"""
        result = self.api.getUserTemplates()
        data = json.loads(result)

        assert isinstance(data["templates"], list)
        if len(data["templates"]) > 0:
            template = data["templates"][0]
            assert "id" in template
            assert "name" in template
        print("✓ getUserTemplates 返回正确的模板列表")

    def test_uploadTemplate_local_mode(self):
        """测试上传模板 - 本地模式"""
        # 创建测试模板文件
        template_file = os.path.join(self.temp_dir, "template.docx")
        with open(template_file, "w") as f:
            f.write("test")

        result = self.api.uploadTemplate(template_file, "测试模板")
        data = json.loads(result)

        assert data["success"] is True
        print("✓ uploadTemplate 本地模式成功上传")

    def test_uploadTemplate_nonexistent_file(self):
        """测试上传模板 - 文件不存在"""
        result = self.api.uploadTemplate("/nonexistent/file.docx", "模板")
        data = json.loads(result)

        assert data["success"] is False
        print("✓ uploadTemplate 正确处理文件不存在的情况")

    def test_deleteTemplate(self):
        """测试删除模板"""
        result = self.api.deleteTemplate("template-001")
        data = json.loads(result)

        assert "success" in data
        print("✓ deleteTemplate 返回有效的 JSON 格式")


# ============================================================
# API 错误处理测试
# ============================================================

class TestAPIErrorHandling:
    """API 错误处理测试"""

    def setup_method(self):
        """测试前准备"""
        self.api = Api(supabase_enabled=False)
        self.api.set_window(create_mock_window())

    def test_api_without_window(self):
        """测试没有窗口对象的 API 调用"""
        api_no_window = Api()
        # 不设置 window

        result = api_no_window.openFile()
        data = json.loads(result)

        assert data["cancelled"] is True
        print("✓ API 正确处理无窗口对象的情况")

    def test_openFile_parse_error(self):
        """测试文件解析错误"""
        # 创建无效的 DOCX 文件
        test_file = os.path.join(tempfile.gettempdir(), "invalid.docx")
        with open(test_file, "wb") as f:
            f.write(b"invalid content")

        self.api.mock_window = create_mock_window()
        self.api.mock_window.create_file_dialog.return_value = [test_file]
        self.api.set_window(self.api.mock_window)

        result = self.api.openFile()
        data = json.loads(result)

        # 应该要么返回 error，要么返回空 elements
        assert "error" in data or "elements" in data
        print("✓ openFile 正确处理解析错误")

        # 清理
        if os.path.exists(test_file):
            os.remove(test_file)


# ============================================================
# 集成测试：完整 API 流程
# ============================================================

class TestAPIIntegration:
    """API 集成测试"""

    def setup_method(self):
        """测试前准备"""
        self.api = Api(supabase_enabled=False)
        self.api.set_window(create_mock_window())
        self.temp_dir = tempfile.mkdtemp()

    def teardown_method(self):
        """测试后清理"""
        import shutil
        shutil.rmtree(self.temp_dir, ignore_errors=True)

    def test_complete_workflow(self):
        """测试完整工作流：登录 → 打开 → 检查 → 保存"""
        # 1. 登录
        login_result = self.api.login("user@test.com", "password")
        login_data = json.loads(login_result)
        assert login_data["success"] is True

        # 2. 获取模板
        templates_result = self.api.getUserTemplates()
        templates_data = json.loads(templates_result)
        assert templates_data["success"] is True

        # 3. 运行质量检查
        qa_result = self.api.runQA("测试文档内容", '["typo"]')
        qa_data = json.loads(qa_result)
        assert isinstance(qa_data, dict)

        # 4. 保存文档
        save_result = self.api.saveDocument("<p>内容</p>", "测试文档")
        save_data = json.loads(save_result)
        assert save_data["success"] is True

        print("✓ 完整工作流测试通过")

    def test_all_apis_return_valid_json(self):
        """验证所有 API 返回有效的 JSON"""
        api_methods = [
            (self.api.getSystemInfo, []),
            (self.api.login, ["test@test.com", "password"]),
            (self.api.logout, []),
            (self.api.runQA, ["测试"]),
            (self.api.runXRef, ["测试"]),
            (self.api.applyFormat, ["测试"]),
            (self.api.getTokenUsage, []),
            (self.api.getUserTemplates, []),
            (self.api.saveDocument, ["<p>测试</p>", "文档"]),
            (self.api.loadDocument, []),
        ]

        for method, args in api_methods:
            result = method(*args)
            try:
                data = json.loads(result)
                assert isinstance(data, dict)
            except json.JSONDecodeError:
                pytest.fail(f"{method.__name__} 返回无效的 JSON: {result}")

        print("✓ 所有 API 返回有效的 JSON")


if __name__ == "__main__":
    pytest.main([__file__, "-v", "-s"])
