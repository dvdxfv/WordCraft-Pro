# -*- coding: utf-8 -*-
"""
第七批 A-③ refreshDocxFields 单元测试

测试场景：
1. 传入有效 docx_b64 → 在无 Office 引擎的环境下返回 NO_OFFICE_ENGINE（正常降级）
2. 传入损坏的 base64 → 返回 error（不抛异常）
3. 传入空字符串 → 返回 error（不抛异常）
4. 返回格式符合约定（有 error_code 或 success+content）
5. Flask 路由 /api/refreshDocxFields 可访问
"""

from __future__ import annotations

import base64
import io
import json
import sys
import os
from pathlib import Path

import pytest
from docx import Document

ROOT = Path(__file__).resolve().parents[1]
SAMPLE_DOC = ROOT / "samples" / "南海鸢乌贼捕捞量智能反演文献综述.docx"

sys.path.insert(0, str(ROOT))
from app import Api


# ── 工具 ─────────────────────────────────────────────────────────────────────

def _make_minimal_docx_b64() -> str:
    doc = Document()
    doc.add_paragraph("测试段落，用于刷新域测试。")
    buf = io.BytesIO()
    doc.save(buf)
    return base64.b64encode(buf.getvalue()).decode("utf-8")


def _parse_result(raw: str) -> dict:
    return json.loads(raw)


# ── 测试 Api.refreshDocxFields ────────────────────────────────────────────────

class TestRefreshDocxFields:
    """测试 app.py 的 refreshDocxFields 方法（无 Office 引擎时的降级行为）。"""

    @pytest.fixture(scope="class")
    def api(self):
        return Api()

    def test_valid_docx_returns_known_keys(self, api):
        """有效 docx → 返回 JSON，含 success+content 或 error_code。"""
        b64 = _make_minimal_docx_b64()
        raw = api.refreshDocxFields(b64)
        result = _parse_result(raw)
        has_success = "success" in result and result["success"] and "content" in result
        has_error_code = "error_code" in result
        assert has_success or has_error_code, \
            f"期望 success+content 或 error_code，实际: {list(result.keys())}"

    def test_no_office_engine_or_success(self, api):
        """
        在 CI / 无 Office 环境下，应返回 error_code=NO_OFFICE_ENGINE（不抛异常）。
        若本机装了 Word 或 LibreOffice，则可能返回 success=True，也合法。
        """
        b64 = _make_minimal_docx_b64()
        raw = api.refreshDocxFields(b64)
        result = _parse_result(raw)
        # 两种合法结果：成功刷新 或 NO_OFFICE_ENGINE
        ok = (result.get("success") is True and "content" in result) or \
             result.get("error_code") in ("NO_OFFICE_ENGINE", "TIMEOUT", "ERROR")
        assert ok, f"不期望的结果: {result}"

    def test_invalid_base64_returns_error(self, api):
        """损坏的 base64 → 返回 error，不抛异常。"""
        raw = api.refreshDocxFields("这不是base64!!@@##")
        result = _parse_result(raw)
        assert "error" in result, f"损坏 base64 应返回 error，实际: {result}"

    def test_empty_string_returns_error(self, api):
        """空字符串 → 返回 error，不抛异常。"""
        raw = api.refreshDocxFields("")
        result = _parse_result(raw)
        assert "error" in result, f"空字符串应返回 error，实际: {result}"

    def test_success_content_is_valid_docx(self, api):
        """如果返回 success，content 必须能被 python-docx 正常打开。"""
        b64 = _make_minimal_docx_b64()
        raw = api.refreshDocxFields(b64)
        result = _parse_result(raw)
        if not (result.get("success") and "content" in result):
            pytest.skip("此环境无 Office 引擎，跳过 content 格式验证")
        docx_bytes = base64.b64decode(result["content"])
        doc = Document(io.BytesIO(docx_bytes))
        assert len(doc.paragraphs) > 0

    @pytest.mark.skipif(not SAMPLE_DOC.exists(), reason="样本文件不存在")
    def test_real_sample_doc(self, api):
        """对真实样本文档调用，不抛异常，返回合法结构。"""
        b64 = base64.b64encode(SAMPLE_DOC.read_bytes()).decode("utf-8")
        raw = api.refreshDocxFields(b64)
        result = _parse_result(raw)
        ok = (result.get("success") is True and "content" in result) or \
             "error_code" in result
        assert ok, f"样本文档刷新返回不合法结构: {result}"


# ── 测试 Flask 路由 ───────────────────────────────────────────────────────────

class TestRefreshDocxFieldsRoute:
    """测试 /api/refreshDocxFields 路由的 HTTP 行为。"""

    @pytest.fixture(scope="class")
    def client(self):
        """创建 Flask 测试客户端。"""
        sys.path.insert(0, str(ROOT / "web"))
        from flask_app import app
        app.config["TESTING"] = True
        with app.test_client() as c:
            yield c

    def test_route_exists_and_returns_json(self, client):
        """路由存在，返回 JSON，不报 404/500。"""
        b64 = _make_minimal_docx_b64()
        resp = client.post(
            "/api/refreshDocxFields",
            json={"docx_b64": b64},
            content_type="application/json",
        )
        assert resp.status_code == 200
        data = json.loads(resp.data)
        assert isinstance(data, dict)

    def test_route_missing_param_returns_error(self, client):
        """缺少 docx_b64 参数 → 返回 error，不崩溃。"""
        resp = client.post(
            "/api/refreshDocxFields",
            json={},
            content_type="application/json",
        )
        assert resp.status_code == 200
        data = json.loads(resp.data)
        assert "error" in data or "error_code" in data

    def test_route_invalid_b64_returns_error(self, client):
        """非法 base64 → 返回 error，不崩溃。"""
        resp = client.post(
            "/api/refreshDocxFields",
            json={"docx_b64": "!!!not_base64!!!"},
            content_type="application/json",
        )
        assert resp.status_code == 200
        data = json.loads(resp.data)
        assert "error" in data


if __name__ == "__main__":
    import pytest as _pytest
    _pytest.main([__file__, "-v"])
