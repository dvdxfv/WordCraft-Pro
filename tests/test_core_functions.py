#!/usr/bin/env python3
"""
WordCraft Pro 核心函数单元测试
测试 app.py 中 Api 类的静态方法：_parse_docx, _parse_text, _html_to_elements, _write_docx
"""

import os
import sys
import pytest

# 确保项目根目录在 sys.path 中
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from app import Api


# ======================================================================
#  测试模板文件路径
# ======================================================================
TEMPLATE_DOCX = os.path.join(
    os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
    "web", "test_template.docx",
)


# ======================================================================
#  1. _parse_docx 测试
# ======================================================================
class TestParseDocx:
    """DOCX 文件解析测试"""

    @pytest.fixture(scope="class")
    def elements(self):
        """解析测试模板，整个类共享结果"""
        return Api._parse_docx(TEMPLATE_DOCX)

    def test_parse_docx_returns_non_empty_list(self, elements):
        """解析返回非空列表"""
        assert isinstance(elements, list)
        assert len(elements) > 0

    def test_parse_docx_contains_h1(self, elements):
        """包含 h1 类型元素（一级标题）"""
        types = [e["type"] for e in elements]
        assert "h1" in types

    def test_parse_docx_contains_h2(self, elements):
        """包含 h2 类型元素（二级标题）"""
        types = [e["type"] for e in elements]
        assert "h2" in types

    def test_parse_docx_contains_h3(self, elements):
        """包含 h3 类型元素（三级标题）"""
        types = [e["type"] for e in elements]
        assert "h3" in types

    def test_parse_docx_contains_p(self, elements):
        """包含 p 类型元素（正文段落）"""
        types = [e["type"] for e in elements]
        assert "p" in types

    def test_parse_docx_contains_table(self, elements):
        """包含 table 类型元素（表格）"""
        types = [e["type"] for e in elements]
        assert "table" in types

    def test_parse_docx_contains_ref(self, elements):
        """包含 ref 类型元素（参考文献）"""
        types = [e["type"] for e in elements]
        assert "ref" in types

    def test_parse_docx_h1_contains_chapter_text(self, elements):
        """h1 元素中包含 '第一章'"""
        h1_texts = [e["text"] for e in elements if e["type"] == "h1"]
        assert any("第一章" in t for t in h1_texts), (
            f"未在 h1 元素中找到 '第一章'，实际 h1 文本: {h1_texts}"
        )

    def test_parse_docx_table_starts_with_table_tag(self, elements):
        """表格元素的 text 以 '<table>' 开头"""
        table_elems = [e for e in elements if e["type"] == "table"]
        assert len(table_elems) > 0, "未找到 table 类型元素"
        for elem in table_elems:
            assert elem["text"].startswith("<table>"), (
                f"表格文本不以 '<table>' 开头: {elem['text'][:50]}"
            )


# ======================================================================
#  2. _parse_text 测试
# ======================================================================
class TestParseText:
    """文本/Markdown 文件解析测试"""

    def test_parse_md_h1(self, tmp_path):
        """解析 Markdown 一级标题 (# → h1)"""
        f = tmp_path / "test.md"
        f.write_text("# 一级标题\n正文内容\n", encoding="utf-8")
        elements = Api._parse_text(str(f))
        assert elements[0]["type"] == "h1"
        assert elements[0]["text"] == "一级标题"

    def test_parse_md_h2(self, tmp_path):
        """解析 Markdown 二级标题 (## → h2)"""
        f = tmp_path / "test.md"
        f.write_text("## 二级标题\n正文内容\n", encoding="utf-8")
        elements = Api._parse_text(str(f))
        assert elements[0]["type"] == "h2"
        assert elements[0]["text"] == "二级标题"

    def test_parse_md_h3(self, tmp_path):
        """解析 Markdown 三级标题 (### → h3)"""
        f = tmp_path / "test.md"
        f.write_text("### 三级标题\n正文内容\n", encoding="utf-8")
        elements = Api._parse_text(str(f))
        assert elements[0]["type"] == "h3"
        assert elements[0]["text"] == "三级标题"

    def test_parse_md_unordered_list(self, tmp_path):
        """解析无序列表 (- item → li)"""
        f = tmp_path / "test.md"
        f.write_text("- 列表项一\n- 列表项二\n", encoding="utf-8")
        elements = Api._parse_text(str(f))
        assert len(elements) == 2
        assert all(e["type"] == "li" for e in elements)
        assert "列表项一" in elements[0]["text"]
        assert "列表项二" in elements[1]["text"]

    def test_parse_md_ordered_list(self, tmp_path):
        """解析有序列表 (1. item → li)"""
        f = tmp_path / "test.md"
        f.write_text("1. 第一项\n2. 第二项\n", encoding="utf-8")
        elements = Api._parse_text(str(f))
        assert len(elements) == 2
        assert all(e["type"] == "li" for e in elements)
        assert "第一项" in elements[0]["text"]
        assert "第二项" in elements[1]["text"]

    def test_parse_md_table(self, tmp_path):
        """解析 Markdown 表格"""
        f = tmp_path / "test.md"
        f.write_text("| 姓名 | 年龄 |\n| --- | --- |\n| 张三 | 25 |\n| 李四 | 30 |\n", encoding="utf-8")
        elements = Api._parse_text(str(f))
        table_elems = [e for e in elements if e["type"] == "table"]
        assert len(table_elems) == 1
        assert "<table>" in table_elems[0]["text"]
        assert "张三" in table_elems[0]["text"]
        assert "李四" in table_elems[0]["text"]

    def test_parse_md_ref(self, tmp_path):
        """解析参考文献 ([1] xxx → ref)"""
        f = tmp_path / "test.md"
        f.write_text("[1] 张三. 某某论文. 2024.\n[2] 李四. 另一篇论文. 2025.\n", encoding="utf-8")
        elements = Api._parse_text(str(f))
        assert len(elements) == 2
        assert all(e["type"] == "ref" for e in elements)

    def test_parse_md_caption(self, tmp_path):
        """解析题注 (图3-1 xxx → caption)"""
        f = tmp_path / "test.md"
        f.write_text("图3-1 系统架构图\n正文内容\n", encoding="utf-8")
        elements = Api._parse_text(str(f))
        assert elements[0]["type"] == "caption"
        assert "系统架构图" in elements[0]["text"]

    def test_parse_md_empty_lines_skipped(self, tmp_path):
        """空行被跳过"""
        f = tmp_path / "test.md"
        f.write_text("\n\n# 标题\n\n\n正文\n\n", encoding="utf-8")
        elements = Api._parse_text(str(f))
        # 空行不应出现在结果中
        assert all(e["text"] for e in elements)
        # 只应有标题和正文两个元素
        assert len(elements) == 2

    def test_parse_md_plain_text(self, tmp_path):
        """普通文本 → p 类型"""
        f = tmp_path / "test.md"
        f.write_text("这是一段普通正文文本。\n", encoding="utf-8")
        elements = Api._parse_text(str(f))
        assert len(elements) == 1
        assert elements[0]["type"] == "p"
        assert elements[0]["text"] == "这是一段普通正文文本。"

    def test_parse_gbk_txt(self, tmp_path):
        """GBK 编码的 .txt 文件能正确解析"""
        f = tmp_path / "test_gbk.txt"
        content = "这是GBK编码的测试文件。\n# 一级标题\n正文段落。\n"
        f.write_bytes(content.encode("gbk"))
        elements = Api._parse_text(str(f))
        assert len(elements) >= 2
        # 验证内容正确解码
        texts = [e["text"] for e in elements]
        assert any("GBK" in t for t in texts)
        assert any("一级标题" in t for t in texts)


# ======================================================================
#  3. _html_to_elements 测试
# ======================================================================
class TestHtmlToElements:
    """HTML 转元素列表测试"""

    def test_empty_string_returns_empty_list(self):
        """空字符串返回空列表"""
        assert Api._html_to_elements("") == []

    def test_h1_tag(self):
        """<h1>标题</h1> → [{type: 'h1', text: '标题'}]"""
        result = Api._html_to_elements("<h1>标题</h1>")
        assert len(result) == 1
        assert result[0]["type"] == "h1"
        assert result[0]["text"] == "标题"

    def test_p_tag(self):
        """<p>段落</p> → [{type: 'p', text: '段落'}]"""
        result = Api._html_to_elements("<p>段落</p>")
        assert len(result) == 1
        assert result[0]["type"] == "p"
        assert result[0]["text"] == "段落"

    def test_table_tag(self):
        """<table><tr><th>A</th></tr></table> → [{type: 'table', text: '<table>...'}]"""
        html = "<table><tr><th>A</th></tr></table>"
        result = Api._html_to_elements(html)
        assert len(result) == 1
        assert result[0]["type"] == "table"
        assert result[0]["text"].startswith("<table>")
        assert "A" in result[0]["text"]

    def test_li_tag(self):
        """<li>列表项</li> → [{type: 'li', text: '列表项'}]"""
        result = Api._html_to_elements("<li>列表项</li>")
        assert len(result) == 1
        assert result[0]["type"] == "li"
        assert result[0]["text"] == "列表项"

    def test_mixed_html(self):
        """混合 HTML 正确分类"""
        html = "<h1>标题</h1><p>段落一</p><p>段落二</p><li>列表项</li>"
        result = Api._html_to_elements(html)
        assert len(result) == 4
        assert result[0]["type"] == "h1"
        assert result[0]["text"] == "标题"
        assert result[1]["type"] == "p"
        assert result[1]["text"] == "段落一"
        assert result[2]["type"] == "p"
        assert result[2]["text"] == "段落二"
        assert result[3]["type"] == "li"
        assert result[3]["text"] == "列表项"

    def test_html_tags_stripped(self):
        """HTML 中的标签被正确去除（只保留纯文本）"""
        html = "<h1>标题<strong>加粗</strong>部分</h1>"
        result = Api._html_to_elements(html)
        assert len(result) == 1
        # 内部标签 <strong> 应被去除，只保留纯文本
        assert "<strong>" not in result[0]["text"]
        assert "</strong>" not in result[0]["text"]
        assert result[0]["text"] == "标题加粗部分"


# ======================================================================
#  4. _write_docx 测试
# ======================================================================
class TestWriteDocx:
    """Word 导出测试"""

    def test_write_docx_creates_file(self, tmp_path):
        """能生成 .docx 文件"""
        out = tmp_path / "output.docx"
        html = "<h1>测试标题</h1><p>测试段落</p>"
        Api._write_docx(str(out), html)
        assert out.exists()

    def test_write_docx_file_size_positive(self, tmp_path):
        """文件大小 > 0"""
        out = tmp_path / "output.docx"
        html = "<h1>测试标题</h1><p>测试段落</p>"
        Api._write_docx(str(out), html)
        assert out.stat().st_size > 0

    def test_write_docx_can_be_reopened(self, tmp_path):
        """生成的文件可以用 python-docx 重新打开"""
        from docx import Document

        out = tmp_path / "output.docx"
        html = "<h1>测试标题</h1><p>测试段落内容</p>"
        Api._write_docx(str(out), html)
        doc = Document(str(out))
        # 能正常读取段落说明文件有效
        paragraphs = doc.paragraphs
        assert len(paragraphs) > 0

    def test_write_docx_contains_correct_content(self, tmp_path):
        """打开后包含正确的段落内容"""
        from docx import Document

        out = tmp_path / "output.docx"
        html = "<h1>一级标题</h1><p>这是正文段落。</p><h2>二级标题</h2>"
        Api._write_docx(str(out), html)
        doc = Document(str(out))
        texts = [p.text for p in doc.paragraphs]
        # 验证标题和段落内容都在文档中
        all_text = "".join(texts)
        assert "一级标题" in all_text
        assert "二级标题" in all_text
        assert "这是正文段落。" in all_text
