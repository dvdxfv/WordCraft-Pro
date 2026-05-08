# -*- coding: utf-8 -*-
"""Batch 7 browser regressions for export and QA accept/undo behavior."""

from __future__ import annotations

from pathlib import Path

import pytest

try:
    from playwright.sync_api import Page
except ImportError:
    pytest.skip("Playwright not installed", allow_module_level=True)

from .conftest import upload_document


TYPO_CAT = "\u9519\u522b\u5b57"
ADVISORY_CAT = "\u6570\u636e\u4e0d\u4e00\u81f4"
UNDO_TEXT = "\u64a4\u9500"


def _read_downloaded_docx_texts(download_path: str) -> list[str]:
    from docx import Document

    doc = Document(download_path)
    return [p.text for p in doc.paragraphs]


def _inject_fake_qa_issue(page: Page, original: str, suggestion: str, cat: str = TYPO_CAT) -> None:
    page.evaluate(
        """([orig, sugg, cat]) => {
            window.currentQAData = [{
                id: 0,
                sev: 'error',
                cat,
                title: 'test issue',
                desc: 'synthetic issue for e2e',
                suggestion: sugg,
                hlText: orig,
                hlClass: 'hl-error',
                status: 'pending',
            }];
            if (openTabs.length > 0) {
                const fn = openTabs[activeTabIdx].name;
                tabQAState[fn] = window.currentQAData;
            }
        }""",
        [original, suggestion, cat],
    )


def _inject_accepted_edit(page: Page, tab_name: str, original: str, suggestion: str) -> None:
    page.evaluate(
        """([fn, orig, sugg]) => {
            if (!tabAcceptedEdits[fn]) tabAcceptedEdits[fn] = [];
            tabAcceptedEdits[fn].push({ original: orig, suggestion: sugg, issueIdx: 0 });
        }""",
        [tab_name, original, suggestion],
    )


def _get_tab_name(page: Page) -> str:
    return page.evaluate("openTabs[activeTabIdx]?.name || ''")


def _get_first_nonempty_para(page: Page) -> str:
    return page.evaluate(
        """() => {
            const tab = openTabs[activeTabIdx];
            if (!tab) return '';
            const content = docContents[tab.name] || [];
            for (const it of content) {
                const text = (it.text || '').replace(/<[^>]+>/g, '').trim();
                if (text.length >= 2) return text;
            }
            return '';
        }"""
    )


def _setup_qa_with_span(page: Page, original: str, suggestion: str, cat: str = TYPO_CAT) -> None:
    _inject_fake_qa_issue(page, original, suggestion, cat)
    page.evaluate(
        """([orig, sugg]) => {
            renderQAIssues(window.currentQAData);
            const docPage = document.getElementById('docPage');
            docPage.innerHTML = '<p>test paragraph</p>';
            const span = document.createElement('span');
            span.className = 'hl-error';
            span.setAttribute('data-issue-id', '0');
            span.setAttribute('data-tip', JSON.stringify({ suggestion: sugg, title: 'test' }));
            span.textContent = orig;
            docPage.querySelector('p').appendChild(span);
        }""",
        [original, suggestion],
    )


@pytest.fixture()
def logged_in_with_sample(logged_in_page: Page, sample_doc_path: Path) -> Page:
    page = logged_in_page
    upload_document(page, sample_doc_path)
    page.wait_for_timeout(2000)
    return page


@pytest.mark.e2e
class TestCloneExportE2E:
    def test_inject_edit_export_text_changed(
        self, logged_in_with_sample: Page, e2e_artifact_dir: Path
    ) -> None:
        page = logged_in_with_sample

        first_para = _get_first_nonempty_para(page)
        assert len(first_para) >= 2, "expected uploaded sample to expose paragraph text"
        original_snippet = first_para[:2]
        replacement = "fix"

        tab_name = _get_tab_name(page)
        assert tab_name, "expected an active tab after upload"

        _inject_accepted_edit(page, tab_name, original_snippet, replacement)

        with page.expect_download() as dl_info:
            page.evaluate("exportDocAsDocx(openTabs[activeTabIdx])")
        download = dl_info.value
        save_path = str(e2e_artifact_dir / download.suggested_filename)
        download.save_as(save_path)

        texts = _read_downloaded_docx_texts(save_path)
        all_text = " ".join(texts)
        assert replacement in all_text, f"expected exported docx to contain {replacement!r}"

    def test_nonexistent_edit_skipped(
        self, logged_in_with_sample: Page, e2e_artifact_dir: Path
    ) -> None:
        page = logged_in_with_sample
        tab_name = _get_tab_name(page)

        _inject_accepted_edit(page, tab_name, "not-present-snippet-xyz", "replacement")

        with page.expect_download() as dl_info:
            page.evaluate("exportDocAsDocx(openTabs[activeTabIdx])")
        download = dl_info.value
        save_path = str(e2e_artifact_dir / download.suggested_filename)
        download.save_as(save_path)

        from docx import Document

        doc = Document(save_path)
        assert len(doc.paragraphs) > 0

    def test_undo_means_no_edit_applied(
        self, logged_in_with_sample: Page, e2e_artifact_dir: Path
    ) -> None:
        page = logged_in_with_sample
        tab_name = _get_tab_name(page)

        first_para = _get_first_nonempty_para(page)
        assert len(first_para) >= 2
        original_snippet = first_para[:2]

        _inject_accepted_edit(page, tab_name, original_snippet, "undo-value")
        page.evaluate("""(fn) => { tabAcceptedEdits[fn] = []; }""", tab_name)

        with page.expect_download() as dl_info:
            page.evaluate("exportDocAsDocx(openTabs[activeTabIdx])")
        download = dl_info.value
        save_path = str(e2e_artifact_dir / download.suggested_filename)
        download.save_as(save_path)

        texts = _read_downloaded_docx_texts(save_path)
        all_text = " ".join(texts)
        assert "undo-value" not in all_text


@pytest.mark.e2e
@pytest.mark.no_login
class TestAcceptUndoUI:
    def test_hl_accepted_class_after_accept(self, anonymous_index_page: Page) -> None:
        page = anonymous_index_page
        _setup_qa_with_span(page, "typo", "fixed")
        page.evaluate("switchRightTab('qa')")
        page.wait_for_timeout(300)
        page.evaluate("acceptIssue(0)")
        page.wait_for_timeout(300)

        hl_class = page.evaluate(
            """() => {
                const el = document.querySelector('[data-issue-id="0"]') ||
                           document.querySelector('[data-accepted-id="0"]');
                return el ? el.className : 'NOT_FOUND';
            }"""
        )
        assert "hl-accepted" in hl_class

    def test_undo_restores_original_class(self, anonymous_index_page: Page) -> None:
        page = anonymous_index_page
        _setup_qa_with_span(page, "typo", "fixed")
        page.evaluate("acceptIssue(0)")
        page.wait_for_timeout(200)
        page.evaluate("undoAcceptedIssue(0)")
        page.wait_for_timeout(200)

        hl_class = page.evaluate(
            """() => {
                const el = document.querySelector('[data-issue-id="0"]');
                return el ? el.className : 'NOT_FOUND';
            }"""
        )
        assert any(name in hl_class for name in ("hl-error", "hl-warn", "hl-info"))

    def test_undo_button_appears_after_accept(self, anonymous_index_page: Page) -> None:
        page = anonymous_index_page
        _setup_qa_with_span(page, "typo", "fixed")
        page.evaluate("acceptIssue(0)")
        page.wait_for_timeout(300)

        undo_btn = page.evaluate(
            """() => {
                const card = document.getElementById('issue-0');
                if (!card) return 'NO_CARD';
                const btn = card.querySelector('.undo-accept');
                return btn ? btn.textContent : 'NO_BTN';
            }"""
        )
        assert UNDO_TEXT in undo_btn

    def test_advisory_issue_replaces_text_when_suggestion_is_direct_replacement(
        self, anonymous_index_page: Page
    ) -> None:
        page = anonymous_index_page
        _setup_qa_with_span(page, "old-value", "new-value", ADVISORY_CAT)
        page.evaluate("acceptIssue(0)")
        page.wait_for_timeout(200)

        result = page.evaluate(
            """() => {
                const el = document.querySelector('[data-accepted-id="0"]');
                return el ? { text: el.textContent, className: el.className } : null;
            }"""
        )
        assert result is not None, "accepted highlight not found"
        assert result["text"] == "new-value"
        assert "hl-accepted" in result["className"]

    def test_typo_issue_replaces_text(self, anonymous_index_page: Page) -> None:
        page = anonymous_index_page
        _setup_qa_with_span(page, "original-typo", "fixed-typo", TYPO_CAT)
        page.evaluate("acceptIssue(0)")
        page.wait_for_timeout(200)

        result = page.evaluate(
            """() => {
                const el = document.querySelector('[data-accepted-id="0"]');
                return el ? { text: el.textContent, className: el.className } : null;
            }"""
        )
        assert result is not None, "accepted highlight not found"
        assert result["text"] == "fixed-typo"
        assert "hl-accepted" in result["className"]

    def test_explanatory_advisory_does_not_replace_text(self, anonymous_index_page: Page) -> None:
        page = anonymous_index_page
        _setup_qa_with_span(
            page,
            "old-value",
            "建议统一为 new-value，并检查上下文中的其他同类写法是否保持一致。",
            ADVISORY_CAT,
        )
        page.evaluate("acceptIssue(0)")
        page.wait_for_timeout(200)

        result = page.evaluate(
            """() => {
                const el = document.querySelector('[data-accepted-id="0"]');
                return el ? { text: el.textContent, className: el.className } : null;
            }"""
        )
        assert result is not None, "accepted highlight not found"
        assert result["text"] == "old-value"
        assert "hl-accepted" in result["className"]
