#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
P2 unit tests: core.xref_fixer.apply_xref_fix

Covers:
  - Single [n] citation → REF field + superscript + bookmark
  - Multi-number [n,m] citation → superscript only (fallback)
  - No reference section → superscript only (B6)
  - Citation not found → error (B6)
  - Fingerprint mismatch → error (Decision A)
  - Bookmark not overwritten when paragraph already has one
"""
from __future__ import annotations

import base64
import io
import os
import sys

import pytest

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))))

from docx import Document
from docx.oxml.ns import qn

from core.xref_fixer import apply_xref_fix


# ─── helpers ────────────────────────────────────────────────────────────────

def _make_docx_with_citation(citation_text: str = "[1]", has_ref_section: bool = True) -> tuple[str, str]:
    """Build a DOCX with a body paragraph containing citation_text and optionally a ref section."""
    doc = Document()
    body_p = doc.add_paragraph()
    body_p.add_run("研究结果显示系统性能优良")
    body_p.add_run(citation_text)
    body_p.add_run("，符合预期。")
    fingerprint = (body_p.runs[0].text + body_p.runs[1].text + body_p.runs[2].text)[:40].strip()

    if has_ref_section:
        doc.add_heading("参考文献", level=1)
        ref_p = doc.add_paragraph()
        ref_p.add_run(f"[1] 贺建海,胡晓惠. 城市给排水工程设计规范. 2020.")
        if "[2]" in citation_text or citation_text in ("[1,2]", "[2]"):
            ref_p2 = doc.add_paragraph()
            ref_p2.add_run("[2] 另一条参考文献. 2021.")

    buf = io.BytesIO()
    doc.save(buf)
    return base64.b64encode(buf.getvalue()).decode(), fingerprint


def _open_b64(b64: str) -> Document:
    return Document(io.BytesIO(base64.b64decode(b64)))


def _has_superscript_run(para_p) -> bool:
    for r in para_p.findall(qn("w:r")):
        rPr = r.find(qn("w:rPr"))
        if rPr is not None:
            va = rPr.find(qn("w:vertAlign"))
            if va is not None and va.get(qn("w:val")) == "superscript":
                return True
    return False


def _count_fld_chars(para_p) -> int:
    return len(list(para_p.iter(qn("w:fldChar"))))


def _get_instr_text(para_p) -> str:
    parts = []
    for t in para_p.iter(qn("w:instrText")):
        parts.append(t.text or "")
    return "".join(parts)


# ─── Single-number citation with ref section ────────────────────────────────

class TestSingleNumberCitation:
    def test_inserts_ref_field_and_superscript(self):
        """Single [1] with ref section → REF field (3 fldChar elements) + superscript."""
        b64, fp = _make_docx_with_citation("[1]")
        result = apply_xref_fix(b64, {
            "type": "xref", "para_idx": 0, "para_fingerprint": fp,
            "citation_text": "[1]", "ref_numbers": [1], "ref_element_indices": [2],
        })
        assert "error" not in result, result.get("error")
        out = _open_b64(result["docx_b64"])
        for p in out.paragraphs:
            if "研究结果" in p.text:
                assert _has_superscript_run(p._p), "Expected superscript run"
                assert _count_fld_chars(p._p) == 3, "Expected 3 fldChar (begin/sep/end)"
                assert "REF _WCRef_1" in _get_instr_text(p._p)
                return
        pytest.fail("Body paragraph not found")

    def test_creates_bookmark_in_ref_paragraph(self):
        b64, fp = _make_docx_with_citation("[1]")
        result = apply_xref_fix(b64, {
            "type": "xref", "para_idx": 0, "para_fingerprint": fp,
            "citation_text": "[1]", "ref_numbers": [1], "ref_element_indices": [2],
        })
        assert "error" not in result
        out = _open_b64(result["docx_b64"])
        found_bookmark = False
        for p in out.paragraphs:
            if "贺建海" in p.text:
                bm = p._p.find(qn("w:bookmarkStart"))
                assert bm is not None, "Expected bookmark in ref paragraph"
                assert bm.get(qn("w:name")) == "_WCRef_1"
                found_bookmark = True
                break
        assert found_bookmark, "Reference paragraph not found"

    def test_preserves_surrounding_text(self):
        """Text before and after citation must survive the replacement."""
        b64, fp = _make_docx_with_citation("[1]")
        result = apply_xref_fix(b64, {
            "type": "xref", "para_idx": 0, "para_fingerprint": fp,
            "citation_text": "[1]", "ref_numbers": [1], "ref_element_indices": [2],
        })
        assert "error" not in result
        out = _open_b64(result["docx_b64"])
        for p in out.paragraphs:
            if "研究结果" in p.text:
                full_text = p.text
                assert "研究结果显示系统性能优良" in full_text
                assert "符合预期" in full_text
                return
        pytest.fail("Body paragraph not found")

    def test_does_not_overwrite_existing_bookmark(self):
        """If ref paragraph already has a bookmark, it is reused, not duplicated."""
        doc = Document()
        body_p = doc.add_paragraph()
        body_p.add_run("段落文字[1]后面内容。")
        fp = body_p.text[:40].strip()
        doc.add_heading("参考文献", level=1)
        ref_p = doc.add_paragraph()
        ref_p.add_run("[1] 已有书签的参考文献条目。")
        # Manually add an existing bookmark to ref_p
        from lxml import etree
        _W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        bm_s = etree.SubElement(ref_p._p, f"{{{_W}}}bookmarkStart")
        bm_s.set(f"{{{_W}}}id", "99")
        bm_s.set(f"{{{_W}}}name", "_ExistingRef")
        bm_e = etree.SubElement(ref_p._p, f"{{{_W}}}bookmarkEnd")
        bm_e.set(f"{{{_W}}}id", "99")

        buf = io.BytesIO()
        doc.save(buf)
        b64 = base64.b64encode(buf.getvalue()).decode()

        result = apply_xref_fix(b64, {
            "type": "xref", "para_idx": 0, "para_fingerprint": fp,
            "citation_text": "[1]", "ref_numbers": [1], "ref_element_indices": [2],
        })
        assert "error" not in result
        out = _open_b64(result["docx_b64"])
        for p in out.paragraphs:
            if "已有书签" in p.text:
                # Should only have one bookmarkStart (the existing one)
                bm_starts = p._p.findall(qn("w:bookmarkStart"))
                assert len(bm_starts) == 1
                assert bm_starts[0].get(qn("w:name")) == "_ExistingRef"
                return
        pytest.fail("Ref paragraph not found")


# ─── Multi-number citation (superscript fallback) ────────────────────────────

class TestMultiNumberCitation:
    def test_multi_number_applies_superscript_only(self):
        """[1,2] → no REF field, just superscript."""
        b64, fp = _make_docx_with_citation("[1,2]")
        result = apply_xref_fix(b64, {
            "type": "xref", "para_idx": 0, "para_fingerprint": fp,
            "citation_text": "[1,2]", "ref_numbers": [1, 2], "ref_element_indices": [2],
        })
        assert "error" not in result
        out = _open_b64(result["docx_b64"])
        for p in out.paragraphs:
            if "研究结果" in p.text:
                assert _has_superscript_run(p._p), "Expected superscript"
                assert _count_fld_chars(p._p) == 0, "Should NOT have REF field for multi-number"
                return
        pytest.fail("Body paragraph not found")


# ─── No reference section (B6 fallback) ─────────────────────────────────────

class TestNoRefSection:
    def test_no_ref_section_applies_superscript_only(self):
        """B6: no 参考文献 heading → superscript only, no crash."""
        b64, fp = _make_docx_with_citation("[1]", has_ref_section=False)
        result = apply_xref_fix(b64, {
            "type": "xref", "para_idx": 0, "para_fingerprint": fp,
            "citation_text": "[1]", "ref_numbers": [1], "ref_element_indices": [],
        })
        assert "error" not in result
        out = _open_b64(result["docx_b64"])
        for p in out.paragraphs:
            if "研究结果" in p.text:
                assert _has_superscript_run(p._p)
                assert _count_fld_chars(p._p) == 0, "No REF field when no ref section"
                return
        pytest.fail("Body paragraph not found")


# ─── Error cases ─────────────────────────────────────────────────────────────

class TestErrorCases:
    def test_fingerprint_mismatch_returns_error(self):
        b64, _ = _make_docx_with_citation("[1]")
        result = apply_xref_fix(b64, {
            "type": "xref", "para_idx": 0, "para_fingerprint": "根本不存在的文字",
            "citation_text": "[1]", "ref_numbers": [1], "ref_element_indices": [2],
        })
        assert "error" in result

    def test_citation_text_not_in_paragraph_returns_error(self):
        """Citation text doesn't appear in the matched paragraph → error."""
        b64, fp = _make_docx_with_citation("[1]")
        result = apply_xref_fix(b64, {
            "type": "xref", "para_idx": 0, "para_fingerprint": fp,
            "citation_text": "[99]",  # doesn't exist in text
            "ref_numbers": [99], "ref_element_indices": [],
        })
        assert "error" in result

    def test_missing_citation_text_returns_error(self):
        b64, fp = _make_docx_with_citation("[1]")
        result = apply_xref_fix(b64, {
            "type": "xref", "para_idx": 0, "para_fingerprint": fp,
            "ref_numbers": [1], "ref_element_indices": [2],
            # citation_text missing
        })
        assert "error" in result
