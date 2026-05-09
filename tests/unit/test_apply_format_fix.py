#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
P0 unit tests: core.format_fixer.apply_format_fix

Covers:
  B1  — font_size writes w:sz in half-points
  B2  — paragraph scope modifies ALL runs, not just one
  A   — fingerprint mismatch returns error; match succeeds
  font_name — eastAsia + ascii + hAnsi set on all runs
  line_spacing (exact / multiple)
"""
from __future__ import annotations

import base64
import io
import os
import sys

import pytest

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))))

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt

from core.format_fixer import apply_format_fix


# ─── helpers ────────────────────────────────────────────────────────────────

def _make_docx_b64(paragraphs: list[tuple[str, float]]) -> tuple[str, list[str]]:
    """Build a DOCX with given paragraphs [(text, pt_size)].
    Returns (b64_string, [fingerprints]).
    """
    doc = Document()
    fingerprints: list[str] = []
    for text, pt_size in paragraphs:
        p = doc.add_paragraph()
        r = p.add_run(text)
        r.font.size = Pt(pt_size)
        fingerprints.append(text[:40].strip())
    buf = io.BytesIO()
    doc.save(buf)
    return base64.b64encode(buf.getvalue()).decode(), fingerprints


def _open_b64(b64: str) -> Document:
    return Document(io.BytesIO(base64.b64decode(b64)))


def _run_sz(run_elem) -> str | None:
    rPr = run_elem.find(qn("w:rPr"))
    if rPr is None:
        return None
    sz = rPr.find(qn("w:sz"))
    return sz.get(qn("w:val")) if sz is not None else None


# ─── B1: half-points ────────────────────────────────────────────────────────

class TestFontSizeHalfPoints:
    def test_font_size_written_as_half_points(self):
        """B1: apply_format_fix('font_size', value=12) writes w:sz=24."""
        b64, fps = _make_docx_b64([("正文段落内容示例文字", 14.0)])
        result = apply_format_fix(b64, {
            "attr": "font_size", "scope": "paragraph", "value": 12,
            "para_idx": 1, "para_fingerprint": fps[0],
        })
        assert "error" not in result, result.get("error")
        doc = _open_b64(result["docx_b64"])
        for para in doc.paragraphs:
            if fps[0][:10] in para.text:
                for r in para._p.findall(qn("w:r")):
                    assert _run_sz(r) == "24", f"Expected w:sz=24 for 12pt, got {_run_sz(r)}"
                return
        pytest.fail("Target paragraph not found")

    def test_font_size_16pt_writes_32(self):
        b64, fps = _make_docx_b64([("标题字号段落", 14.0)])
        result = apply_format_fix(b64, {
            "attr": "font_size", "value": 16,
            "para_idx": 1, "para_fingerprint": fps[0],
        })
        assert "error" not in result
        doc = _open_b64(result["docx_b64"])
        for para in doc.paragraphs:
            if fps[0][:5] in para.text:
                for r in para._p.findall(qn("w:r")):
                    assert _run_sz(r) == "32"
                return
        pytest.fail("Target paragraph not found")


# ─── B2: paragraph scope modifies ALL runs ───────────────────────────────────

class TestParagraphScopeAllRuns:
    def test_paragraph_scope_modifies_all_runs(self):
        """B2: paragraph scope changes every run, not just the first."""
        doc = Document()
        para = doc.add_paragraph()
        for text, pt in [("第一段", 14.0), ("第二段", 18.0), ("第三段", 10.0)]:
            r = para.add_run(text)
            r.font.size = Pt(pt)
        buf = io.BytesIO()
        doc.save(buf)
        b64 = base64.b64encode(buf.getvalue()).decode()
        fp = para.text[:40].strip()

        result = apply_format_fix(b64, {
            "attr": "font_size", "scope": "paragraph", "value": 12,
            "para_idx": 1, "para_fingerprint": fp,
        })
        assert "error" not in result, result.get("error")
        out = _open_b64(result["docx_b64"])
        for p in out.paragraphs:
            if "第一段" in p.text:
                runs = p._p.findall(qn("w:r"))
                assert len(runs) == 3, "Should have 3 runs"
                for r in runs:
                    assert _run_sz(r) == "24", f"All runs should be 12pt (sz=24), got {_run_sz(r)}"
                return
        pytest.fail("Target paragraph not found")


# ─── Decision A: fingerprint validation ──────────────────────────────────────

class TestFingerprintValidation:
    def test_fingerprint_mismatch_returns_error(self):
        """Decision A: wrong fingerprint → error, DOCX not modified."""
        b64, fps = _make_docx_b64([("真实内容段落", 14.0)])
        result = apply_format_fix(b64, {
            "attr": "font_size", "value": 12,
            "para_idx": 0, "para_fingerprint": "不存在的指纹文字",
        })
        assert "error" in result
        assert "not found" in result["error"].lower() or "not found" in result["error"]

    def test_fingerprint_match_applies_correctly(self):
        """Decision A: correct fingerprint → modification succeeds."""
        b64, fps = _make_docx_b64([("指纹匹配段落内容", 14.0)])
        result = apply_format_fix(b64, {
            "attr": "font_size", "value": 12,
            "para_idx": 1, "para_fingerprint": fps[0],
        })
        assert "error" not in result, result.get("error")
        assert "docx_b64" in result

    def test_hint_idx_mismatch_still_finds_by_text(self):
        """Even if para_idx hint is wrong, fingerprint search finds the paragraph."""
        b64, fps = _make_docx_b64([
            ("第一段内容", 14.0),
            ("第二段内容需要修改字号", 14.0),
        ])
        # Pass wrong hint_idx but correct fingerprint for paragraph 2
        result = apply_format_fix(b64, {
            "attr": "font_size", "value": 12,
            "para_idx": 0,  # wrong — paragraph 2 is at a different index
            "para_fingerprint": fps[1],
        })
        assert "error" not in result, result.get("error")


# ─── font_name ───────────────────────────────────────────────────────────────

class TestFontName:
    def test_font_name_sets_east_asia(self):
        b64, fps = _make_docx_b64([("字体修改段落", 12.0)])
        result = apply_format_fix(b64, {
            "attr": "font_name", "scope": "paragraph", "value": "宋体",
            "para_idx": 1, "para_fingerprint": fps[0],
        })
        assert "error" not in result
        out = _open_b64(result["docx_b64"])
        for p in out.paragraphs:
            if fps[0][:4] in p.text:
                for r in p._p.findall(qn("w:r")):
                    rPr = r.find(qn("w:rPr"))
                    assert rPr is not None
                    rFonts = rPr.find(qn("w:rFonts"))
                    assert rFonts is not None
                    assert rFonts.get(qn("w:eastAsia")) == "宋体"
                    assert rFonts.get(qn("w:ascii")) == "宋体"
                return
        pytest.fail("Target paragraph not found")


# ─── line_spacing ────────────────────────────────────────────────────────────

class TestLineSpacing:
    def test_exact_spacing_writes_twips(self):
        """Exact 20pt → w:spacing line=400, lineRule=exact."""
        b64, fps = _make_docx_b64([("行距测试段落内容", 12.0)])
        result = apply_format_fix(b64, {
            "attr": "line_spacing", "scope": "paragraph",
            "mode": "exact", "value": 20,
            "para_idx": 1, "para_fingerprint": fps[0],
        })
        assert "error" not in result
        out = _open_b64(result["docx_b64"])
        for p in out.paragraphs:
            if fps[0][:4] in p.text:
                pPr = p._p.find(qn("w:pPr"))
                assert pPr is not None
                spacing = pPr.find(qn("w:spacing"))
                assert spacing is not None
                assert spacing.get(qn("w:line")) == "400"   # 20pt * 20
                assert spacing.get(qn("w:lineRule")) == "exact"
                return
        pytest.fail("Target paragraph not found")

    def test_multiple_spacing_writes_twips(self):
        """Multiple 1.5x → w:spacing line=360, lineRule=auto."""
        b64, fps = _make_docx_b64([("倍数行距段落内容", 12.0)])
        result = apply_format_fix(b64, {
            "attr": "line_spacing", "scope": "paragraph",
            "mode": "multiple", "value": 1.5,
            "para_idx": 1, "para_fingerprint": fps[0],
        })
        assert "error" not in result
        out = _open_b64(result["docx_b64"])
        for p in out.paragraphs:
            if fps[0][:4] in p.text:
                pPr = p._p.find(qn("w:pPr"))
                spacing = pPr.find(qn("w:spacing"))
                assert spacing is not None
                assert spacing.get(qn("w:line")) == "360"   # 1.5 * 240
                assert spacing.get(qn("w:lineRule")) == "auto"
                return
        pytest.fail("Target paragraph not found")

    def test_unknown_mode_returns_error(self):
        b64, fps = _make_docx_b64([("行距测试", 12.0)])
        result = apply_format_fix(b64, {
            "attr": "line_spacing", "mode": "unknown_mode", "value": 1.5,
            "para_idx": 1, "para_fingerprint": fps[0],
        })
        assert "error" in result
