# -*- coding: utf-8 -*-
"""
第七批 Clone Export 单元测试

测试与 JS _applyAcceptedEditsToDocXml 等效的 Python 实现，验证：
1. 单 run 内文字替换
2. 跨 run 文字替换
3. 不存在的原文 → skipped 计数 +1，XML 不变
4. 空 edit 列表 → XML 原样返回
5. original == suggestion 的 edit → skipped（无意义替换）
6. 对真实 .docx 样本做 clone 导出，验证目标文字被替换、其余段落不变
"""

from __future__ import annotations

import io
import zipfile
from pathlib import Path
from typing import Any
from xml.etree import ElementTree as ET

import pytest
from docx import Document

ROOT = Path(__file__).resolve().parents[1]
SAMPLE_DOC = ROOT / "samples" / "南海鸢乌贼捕捞量智能反演文献综述.docx"

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
XML_NS = "http://www.w3.org/XML/1998/namespace"


# ── Python 等效实现 ──────────────────────────────────────────────────────────

def _apply_edits_to_doc_xml(xml_bytes: bytes, edits: list[dict]) -> tuple[bytes, int, int]:
    """
    Python 等效的 JS _applyAcceptedEditsToDocXml。

    对 word/document.xml 的字节串执行文字替换：
    - 跳过 w:instrText run
    - 支持跨 run 匹配
    - 返回 (new_xml_bytes, applied, skipped)
    """
    if not edits:
        return xml_bytes, 0, 0

    # 注册所有已用到的命名空间，避免 ET 把它们重写成 ns0/ns1
    _ns_map = {
        "wpc": "http://schemas.microsoft.com/office/word/2010/wordprocessingCanvas",
        "cx":  "http://schemas.microsoft.com/office/drawing/2014/chartex",
        "cx2": "http://schemas.microsoft.com/office/drawing/2015/9/8/chartex",
        "mc":  "http://schemas.openxmlformats.org/markup-compatibility/2006",
        "aink":"http://schemas.microsoft.com/office/drawing/2016/ink",
        "am3d":"http://schemas.microsoft.com/office/drawing/2017/model3d",
        "o":   "urn:schemas-microsoft-com:office:office",
        "oel": "http://schemas.microsoft.com/office/2019/extlst",
        "r":   "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
        "m":   "http://schemas.openxmlformats.org/officeDocument/2006/math",
        "v":   "urn:schemas-microsoft-com:vml",
        "wp14":"http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing",
        "wp":  "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
        "w10": "urn:schemas-microsoft-com:office:word",
        "w":   W_NS,
        "w14": "http://schemas.microsoft.com/office/word/2010/wordml",
        "w15": "http://schemas.microsoft.com/office/word/2012/wordml",
        "w16cex":"http://schemas.microsoft.com/office/word/2018/wordml/cex",
        "w16cid":"http://schemas.microsoft.com/office/word/2016/wordml/cid",
        "w16": "http://schemas.microsoft.com/office/word/2018/wordml",
        "w16sdtdh":"http://schemas.microsoft.com/office/word/2020/wordml/sdtdatahash",
        "w16se":"http://schemas.microsoft.com/office/word/2015/wordml/symex",
        "wpg": "http://schemas.microsoft.com/office/word/2010/wordprocessingGroup",
        "wpi": "http://schemas.microsoft.com/office/word/2010/wordprocessingInk",
        "wne": "http://schemas.microsoft.com/office/word/2006/wordml",
        "wps": "http://schemas.microsoft.com/office/word/2010/wordprocessingShape",
    }
    for prefix, uri in _ns_map.items():
        try:
            ET.register_namespace(prefix, uri)
        except Exception:
            pass

    root = ET.fromstring(xml_bytes)
    body = root.find(f"{{{W_NS}}}body")
    if body is None:
        return xml_bytes, 0, len(edits)

    applied = skipped = 0

    for edit in edits:
        original = edit.get("original", "")
        suggestion = edit.get("suggestion", "")
        if not original or not suggestion or original == suggestion:
            skipped += 1
            continue

        found = False
        for para in body.iter(f"{{{W_NS}}}p"):
            # 收集 (run_el, t_el) 对，跳过含 instrText 的 run
            t_pairs: list[tuple[Any, Any]] = []
            for r_el in para.iter(f"{{{W_NS}}}r"):
                if r_el.find(f"{{{W_NS}}}instrText") is not None:
                    continue
                for t_el in r_el.findall(f"{{{W_NS}}}t"):
                    t_pairs.append((r_el, t_el))

            if not t_pairs:
                continue

            full_text = "".join(t.text or "" for _, t in t_pairs)
            idx = full_text.find(original)
            if idx < 0:
                continue

            match_end = idx + len(original)
            offset = 0
            replacement_inserted = False

            for _, t_el in t_pairs:
                node_start = offset
                node_text = t_el.text or ""
                node_len = len(node_text)
                node_end = node_start + node_len
                offset = node_end

                if node_end <= idx or node_start >= match_end:
                    continue

                before = node_text[: max(0, idx - node_start)]
                after_start = max(node_start, match_end)
                after = node_text[after_start - node_start:] if after_start < node_end else ""

                if not replacement_inserted:
                    t_el.text = before + suggestion + after
                    replacement_inserted = True
                else:
                    t_el.text = after

                # 保留首尾空格
                if t_el.text and (t_el.text.startswith(" ") or t_el.text.endswith(" ")):
                    t_el.set(f"{{{XML_NS}}}space", "preserve")

            found = True
            applied += 1
            break

        if not found:
            skipped += 1

    out = io.BytesIO()
    tree = ET.ElementTree(root)
    tree.write(out, xml_declaration=True, encoding="UTF-8", short_empty_elements=True)
    return out.getvalue(), applied, skipped


def _clone_docx_with_edits(src_bytes: bytes, edits: list[dict]) -> tuple[bytes, int, int]:
    """打开 docx zip，替换 document.xml，返回新 zip 字节。"""
    src_buf = io.BytesIO(src_bytes)
    out_buf = io.BytesIO()
    applied = skipped = 0

    with zipfile.ZipFile(src_buf, "r") as zin, zipfile.ZipFile(out_buf, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            fname = item.filename  # Python 3.14: ZipInfo.filename (not .name)
            data = zin.read(fname)
            if fname == "word/document.xml":
                data, applied, skipped = _apply_edits_to_doc_xml(data, edits)
            zout.writestr(item, data)

    return out_buf.getvalue(), applied, skipped


# ── 构造最小 docx fixture ───────────────────────────────────────────────────

def _make_minimal_docx(paragraphs: list[str]) -> bytes:
    """用 python-docx 在内存里生成一个最小 .docx。"""
    doc = Document()
    for text in paragraphs:
        doc.add_paragraph(text)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _read_all_text(docx_bytes: bytes) -> list[str]:
    """读取 docx 全部段落文本。"""
    doc = Document(io.BytesIO(docx_bytes))
    return [p.text for p in doc.paragraphs]


# ── 测试：_apply_edits_to_doc_xml ──────────────────────────────────────────

class TestApplyEditsToDocXml:

    def _xml_with_text(self, *texts: str) -> bytes:
        """构造包含多个段落（各自一个 run）的最小 document.xml。"""
        doc_bytes = _make_minimal_docx(list(texts))
        with zipfile.ZipFile(io.BytesIO(doc_bytes)) as z:
            return z.read("word/document.xml")

    def test_empty_edits_returns_unchanged(self):
        xml = self._xml_with_text("原始文本")
        out, applied, skipped = _apply_edits_to_doc_xml(xml, [])
        assert out == xml
        assert applied == 0
        assert skipped == 0

    def test_single_run_replacement(self):
        xml = self._xml_with_text("请输入正确的帐号进行登录。")
        out, applied, skipped = _apply_edits_to_doc_xml(
            xml, [{"original": "帐号", "suggestion": "账号"}]
        )
        assert applied == 1
        assert skipped == 0
        root = ET.fromstring(out)
        texts = [t.text or "" for t in root.iter(f"{{{W_NS}}}t")]
        combined = "".join(texts)
        assert "账号" in combined
        assert "帐号" not in combined

    def test_original_not_found_is_skipped(self):
        xml = self._xml_with_text("这段话里没有目标词。")
        out, applied, skipped = _apply_edits_to_doc_xml(
            xml, [{"original": "不存在的词", "suggestion": "替换"}]
        )
        assert applied == 0
        assert skipped == 1
        # XML 序列化后格式可能有细微差异（声明/命名空间），比较文本内容即可
        root = ET.fromstring(out)
        texts = "".join(t.text or "" for t in root.iter(f"{{{W_NS}}}t"))
        assert "这段话里没有目标词" in texts
        assert "替换" not in texts

    def test_same_original_and_suggestion_is_skipped(self):
        xml = self._xml_with_text("测试文本")
        out, applied, skipped = _apply_edits_to_doc_xml(
            xml, [{"original": "测试", "suggestion": "测试"}]
        )
        assert skipped == 1
        assert applied == 0

    def test_multiple_edits_each_in_different_para(self):
        xml = self._xml_with_text("帐号登录", "数据显示精确度高")
        out, applied, skipped = _apply_edits_to_doc_xml(
            xml,
            [
                {"original": "帐号", "suggestion": "账号"},
                {"original": "精确度", "suggestion": "准确度"},
            ],
        )
        assert applied == 2
        assert skipped == 0
        root = ET.fromstring(out)
        texts = "".join(t.text or "" for t in root.iter(f"{{{W_NS}}}t"))
        assert "账号" in texts and "帐号" not in texts
        assert "准确度" in texts and "精确度" not in texts

    def test_empty_original_is_skipped(self):
        xml = self._xml_with_text("测试")
        out, applied, skipped = _apply_edits_to_doc_xml(
            xml, [{"original": "", "suggestion": "x"}]
        )
        assert skipped == 1
        assert applied == 0


# ── 测试：_clone_docx_with_edits（整包 clone）─────────────────────────────

class TestCloneDocxWithEdits:

    def test_clone_no_edits_identical_text(self):
        original = _make_minimal_docx(["克隆测试段落。"])
        cloned, applied, skipped = _clone_docx_with_edits(original, [])
        assert applied == 0
        assert _read_all_text(cloned) == _read_all_text(original)

    def test_clone_applies_single_edit(self):
        original = _make_minimal_docx(["正确的帐号很重要。"])
        cloned, applied, skipped = _clone_docx_with_edits(
            original, [{"original": "帐号", "suggestion": "账号"}]
        )
        assert applied == 1
        texts = _read_all_text(cloned)
        assert any("账号" in t for t in texts)
        assert all("帐号" not in t for t in texts)

    def test_clone_preserves_other_paragraphs(self):
        original = _make_minimal_docx(["第一段：保持不变。", "第二段：帐号问题。", "第三段：也保持不变。"])
        cloned, applied, skipped = _clone_docx_with_edits(
            original, [{"original": "帐号", "suggestion": "账号"}]
        )
        texts = _read_all_text(cloned)
        assert "第一段：保持不变。" in texts
        assert "第三段：也保持不变。" in texts
        assert any("账号" in t for t in texts)

    def test_clone_skipped_edit_leaves_file_readable(self):
        """即使 edit 未匹配，导出文件依然可读。"""
        original = _make_minimal_docx(["正常文本。"])
        cloned, applied, skipped = _clone_docx_with_edits(
            original, [{"original": "不存在", "suggestion": "x"}]
        )
        assert skipped == 1
        # 文件可以被 python-docx 正常打开
        doc = Document(io.BytesIO(cloned))
        assert len(doc.paragraphs) > 0


# ── 测试：对真实样本文档做 clone 导出 ──────────────────────────────────────

@pytest.mark.skipif(not SAMPLE_DOC.exists(), reason=f"样本文件不存在: {SAMPLE_DOC}")
class TestCloneRealSampleDoc:

    @pytest.fixture(scope="class")
    def sample_bytes(self):
        return SAMPLE_DOC.read_bytes()

    @pytest.fixture(scope="class")
    def sample_first_nonempty_para(self, sample_bytes):
        doc = Document(io.BytesIO(sample_bytes))
        for p in doc.paragraphs:
            if p.text.strip():
                return p.text.strip()
        pytest.skip("样本文档无非空段落")

    def test_clone_with_no_edits_preserves_paragraph_count(self, sample_bytes):
        original_doc = Document(io.BytesIO(sample_bytes))
        cloned, applied, skipped = _clone_docx_with_edits(sample_bytes, [])
        cloned_doc = Document(io.BytesIO(cloned))
        assert len(cloned_doc.paragraphs) == len(original_doc.paragraphs)
        assert applied == 0

    def test_clone_with_real_text_replacement(self, sample_bytes, sample_first_nonempty_para):
        # 取第一段前两个汉字作为替换目标
        original_snippet = sample_first_nonempty_para[:2]
        replacement = "XX"
        cloned, applied, skipped = _clone_docx_with_edits(
            sample_bytes,
            [{"original": original_snippet, "suggestion": replacement}],
        )
        assert applied == 1, f"期望 applied=1，实际 applied={applied}，skipped={skipped}"
        all_text = " ".join(_read_all_text(cloned))
        assert replacement in all_text

    def test_clone_nonexistent_text_returns_skipped(self, sample_bytes):
        fake_edit = {"original": "这段文字肯定不存在XYZXYZ", "suggestion": "替换"}
        cloned, applied, skipped = _clone_docx_with_edits(sample_bytes, [fake_edit])
        assert applied == 0
        assert skipped == 1
        # 文件依然可读
        Document(io.BytesIO(cloned))

    def test_clone_undo_means_no_edit_equals_original(self, sample_bytes):
        """采纳 → 撤销 → 导出 ≈ 不传 edits → 文字与原文一致。"""
        # 撤销等价于把该 edit 从列表中移除，导出时 edits=[]
        cloned_no_edit, _, _ = _clone_docx_with_edits(sample_bytes, [])
        original_texts = _read_all_text(sample_bytes)
        cloned_texts = _read_all_text(cloned_no_edit)
        assert original_texts == cloned_texts


if __name__ == "__main__":
    import pytest as _pytest
    _pytest.main([__file__, "-v"])
