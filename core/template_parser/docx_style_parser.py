"""
DOCX 样式解析器 (DOCX Style Parser)

从 .docx 模板文件中提取样式定义，转换为 FormattingRules。
通过分析文档中各段落的实际样式来推断排版规则。
"""

from __future__ import annotations

import re
from collections import defaultdict
from typing import Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from core.formatting_rules import FormattingRules, StyleRules, PageRules, SectionRules
from core.document_model import Alignment, LineSpacingType, PageNumberFormat
from core.template_parser.rule_normalizer import RuleNormalizer
from parsers.docx_parser import DocxParser


class DocxStyleParser:
    """从 .docx 模板文件提取排版规则"""

    def parse(self, file_path: str) -> FormattingRules:
        """解析 .docx 模板，返回排版规则"""
        doc = Document(file_path)

        rules = FormattingRules()
        rules.template_name = self._extract_template_name(doc, file_path)

        # 1. 提取页面设置
        self._extract_page_rules(doc, rules)

        # 2. 提取多节页眉页脚
        self._extract_section_rules(doc, rules)

        # 3. 按样式名分组分析段落
        style_groups = self._group_paragraphs_by_style(doc)

        # 4. 从分组中提取各元素规则
        self._extract_rules_from_groups(style_groups, rules)

        # 5. 填充默认值
        rules = RuleNormalizer.fill_defaults(rules)

        return rules

    def _extract_template_name(self, doc: Document, file_path: str) -> str:
        """提取模板名称"""
        from pathlib import Path
        name = Path(file_path).stem
        # 清理常见前缀
        for prefix in ["template_", "模板_", "_"]:
            if name.startswith(prefix):
                name = name[len(prefix):]
        return name

    def _extract_page_rules(self, doc: Document, rules: FormattingRules):
        """从文档节中提取页面设置"""
        if not doc.sections:
            return
        section = doc.sections[0]

        emu_to_cm = 1.0 / 360000.0
        rules.page.paper_size = DocxParser._detect_paper_size(
            section.page_width * emu_to_cm,
            section.page_height * emu_to_cm,
        )
        rules.page.margin_top_cm = round(section.top_margin * emu_to_cm, 2)
        rules.page.margin_bottom_cm = round(section.bottom_margin * emu_to_cm, 2)
        rules.page.margin_left_cm = round(section.left_margin * emu_to_cm, 2)
        rules.page.margin_right_cm = round(section.right_margin * emu_to_cm, 2)
        rules.page.gutter_cm = round(section.gutter * emu_to_cm, 2)
        rules.page.header_distance_cm = round(section.header_distance * emu_to_cm, 2)
        rules.page.footer_distance_cm = round(section.footer_distance * emu_to_cm, 2)

    def _extract_section_rules(self, doc: Document, rules: FormattingRules):
        """提取多节页眉页脚规则"""
        from docx.oxml.ns import qn

        for i, section in enumerate(doc.sections):
            section_rules = SectionRules()

            # 首页不同
            sect_pr = section._sectPr
            title_pg = sect_pr.find(qn("w:titlePg"))
            section_rules.first_page_different = title_pg is not None

            # 页眉
            header = section.header
            if header and not header.is_linked_to_previous:
                header_text = ""
                for para in header.paragraphs:
                    if para.text.strip():
                        header_text = para.text.strip()
                        break
                if header_text:
                    section_rules.header_text = header_text
                    # 提取页眉字体
                    for para in header.paragraphs:
                        for run in para.runs:
                            if run.font.name:
                                section_rules.header_font = run.font.name
                            if run.font.size:
                                section_rules.header_size_pt = run.font.size.pt
                            if run.font.bold:
                                section_rules.header_bold = True
                            if run.font.name or run.font.size:
                                break
                        if section_rules.header_font:
                            break

            # 首页页眉
            if section_rules.first_page_different:
                first_header = section.first_page_header
                if first_header:
                    for para in first_header.paragraphs:
                        if para.text.strip():
                            section_rules.first_page_header_text = para.text.strip()
                            break

            # 页码
            footer = section.footer
            if footer:
                xml = footer._element.xml
                has_page_num = "PAGE" in xml or "w:fldChar" in xml
                if has_page_num:
                    section_rules.page_number_enabled = True

                    # 页码格式
                    pg_num_type = sect_pr.find(qn("w:pgNumType"))
                    if pg_num_type is not None:
                        fmt = pg_num_type.get(qn("w:fmt"))
                        if fmt:
                            fmt_map = {"upperRoman": PageNumberFormat.UPPER_ROMAN,
                                       "lowerRoman": PageNumberFormat.LOWER_ROMAN,
                                       "decimal": PageNumberFormat.ARABIC}
                            section_rules.page_number_type = fmt_map.get(fmt, PageNumberFormat.ARABIC)
                        start = pg_num_type.get(qn("w:start"))
                        if start:
                            section_rules.page_number_start_from = int(start)

                    # 页码格式模板（检测 － X － 格式）
                    footer_text = ""
                    for para in footer.paragraphs:
                        footer_text += para.text
                    if "－" in footer_text or "—" in footer_text:
                        section_rules.page_number_format = "－ {num} －"

            # 节名称
            if section_rules.header_text:
                section_rules.name = section_rules.header_text.strip()
            else:
                section_rules.name = f"第{i + 1}节"

            rules.sections.append(section_rules)

    def _group_paragraphs_by_style(self, doc: Document) -> dict[str, list]:
        """按样式名分组段落"""
        groups = defaultdict(list)
        for para in doc.paragraphs:
            style_name = para.style.name if para.style else "Normal"
            if para.text.strip():  # 只统计非空段落
                groups[style_name].append(para)
        return dict(groups)

    def _extract_rules_from_groups(self, groups: dict[str, list], rules: FormattingRules):
        """从样式分组中提取排版规则"""
        # 标题样式映射
        heading_style_map = {
            "heading 1": ("heading1", 1),
            "heading 2": ("heading2", 2),
            "heading 3": ("heading3", 3),
            "heading 4": ("heading4", 4),
            "标题 1": ("heading1", 1),
            "标题 2": ("heading2", 2),
            "标题 3": ("heading3", 3),
            "标题 4": ("heading4", 4),
        }

        for style_name, paragraphs in groups.items():
            lower_name = style_name.lower()

            # 标题
            if lower_name in heading_style_map:
                attr, level = heading_style_map[lower_name]
                style = getattr(rules, attr)
                self._infer_style_from_paragraphs(paragraphs, style)
                continue

            # 标题（Title）
            if lower_name in ("title", "标题"):
                self._infer_style_from_paragraphs(paragraphs, rules.title)
                continue

            # 正文
            if lower_name in ("normal", "正文", "body text", "body"):
                self._infer_style_from_paragraphs(paragraphs, rules.body)
                continue

            # 参考文献
            if "bibliography" in lower_name or "参考" in lower_name:
                self._infer_style_from_paragraphs(paragraphs, rules.reference)
                continue

            # 题注
            if "caption" in lower_name or "题注" in lower_name:
                # 判断是图题还是表题
                for para in paragraphs:
                    if para.text.strip().startswith("表"):
                        self._infer_style_from_paragraphs(paragraphs, rules.table_caption)
                        break
                    elif para.text.strip().startswith("图"):
                        self._infer_style_from_paragraphs(paragraphs, rules.figure_caption)
                        break

    def _infer_style_from_paragraphs(self, paragraphs: list, style: StyleRules):
        """从一组段落中推断样式（取出现最多的值）"""
        if not paragraphs:
            return

        # 取前3个段落作为样本
        samples = paragraphs[:3]

        # 字体
        cn_fonts = []
        en_fonts = []
        sizes = []
        bolds = []

        for para in samples:
            for run in para.runs:
                rpr = run._element.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rPr")
                if rpr is not None:
                    ea = rpr.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rFonts")
                    if ea is not None:
                        cn = ea.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia")
                        en = ea.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}ascii")
                        if cn:
                            cn_fonts.append(cn)
                        if en:
                            en_fonts.append(en)
                if run.font.name:
                    en_fonts.append(run.font.name)
                if run.font.size:
                    sizes.append(run.font.size.pt)
                if run.font.bold:
                    bolds.append(True)
                elif run.font.bold is False:
                    bolds.append(False)
                break  # 每段只取第一个 run

        # 取众数
        if cn_fonts:
            style.font_name_cn = max(set(cn_fonts), key=cn_fonts.count)
        if en_fonts:
            style.font_name_en = max(set(en_fonts), key=en_fonts.count)
        if sizes:
            style.font_size_pt = max(set(sizes), key=sizes.count)
        if bolds:
            style.bold = max(set(bolds), key=bolds.count)

        # 对齐
        aligns = []
        align_map = {
            WD_ALIGN_PARAGRAPH.LEFT: Alignment.LEFT,
            WD_ALIGN_PARAGRAPH.CENTER: Alignment.CENTER,
            WD_ALIGN_PARAGRAPH.RIGHT: Alignment.RIGHT,
            WD_ALIGN_PARAGRAPH.JUSTIFY: Alignment.JUSTIFY,
        }
        for para in samples:
            if para.paragraph_format.alignment is not None:
                aligns.append(align_map.get(para.paragraph_format.alignment, Alignment.LEFT))
        if aligns:
            style.alignment = max(set(aligns), key=aligns.count)

        # 行距
        spacings = []
        for para in samples:
            pf = para.paragraph_format
            if pf.line_spacing is not None:
                if pf.line_spacing_rule is not None:
                    rule_name = pf.line_spacing_rule.name if hasattr(pf.line_spacing_rule, "name") else str(pf.line_spacing_rule)
                    if "EXACTLY" in rule_name or "AT_LEAST" in rule_name:
                        spacings.append(("exact", pf.line_spacing.pt if hasattr(pf.line_spacing, "pt") else pf.line_spacing))
                    else:
                        spacings.append(("multiple", pf.line_spacing))
        if spacings:
            # 取最常见的行距类型
            most_common = max(set(s[0] for s in spacings), key=lambda x: sum(1 for s in spacings if s[0] == x))
            values = [s[1] for s in spacings if s[0] == most_common]
            if most_common == "exact":
                style.line_spacing_type = LineSpacingType.EXACT
            else:
                style.line_spacing_type = LineSpacingType.MULTIPLE
            style.line_spacing_value = sum(values) / len(values) if values else 1.0

        # 段前段后
        before_vals = []
        after_vals = []
        for para in samples:
            pf = para.paragraph_format
            if pf.space_before is not None:
                before_vals.append(pf.space_before.pt if hasattr(pf.space_before, "pt") else pf.space_before)
            if pf.space_after is not None:
                after_vals.append(pf.space_after.pt if hasattr(pf.space_after, "pt") else pf.space_after)
        if before_vals:
            style.space_before_pt = round(sum(before_vals) / len(before_vals), 1)
        if after_vals:
            style.space_after_pt = round(sum(after_vals) / len(after_vals), 1)
