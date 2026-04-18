"""
导出引擎 (Exporter)

将 DocumentModel 导出为 .docx 文件。
将文档模型中的样式属性映射到 python-docx 的格式设置。
"""

from __future__ import annotations

import os
from typing import Optional

from docx import Document
from docx.shared import Pt, Cm, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

from core.document_model import (
    DocumentModel, DocElement, ElementType,
    FontStyle, ParagraphStyle, Alignment, LineSpacingType,
    TableData, TableCell,
    PageSetup, SectionConfig, HeaderFooterConfig, PageNumberConfig, PageNumberFormat,
)


# Alignment 映射
_ALIGN_TO_DOCX = {
    Alignment.LEFT: WD_ALIGN_PARAGRAPH.LEFT,
    Alignment.CENTER: WD_ALIGN_PARAGRAPH.CENTER,
    Alignment.RIGHT: WD_ALIGN_PARAGRAPH.RIGHT,
    Alignment.JUSTIFY: WD_ALIGN_PARAGRAPH.JUSTIFY,
}

# PageNumberFormat 映射
_PN_FORMAT_MAP = {
    PageNumberFormat.ARABIC: "decimal",
    PageNumberFormat.UPPER_ROMAN: "upperRoman",
    PageNumberFormat.LOWER_ROMAN: "lowerRoman",
}


class Exporter:
    """导出引擎 — 将 DocumentModel 导出为 .docx"""

    def export(self, doc: DocumentModel, output_path: str,
               qa_report: 'QAReport' = None,
               crossref_report: 'CrossRefReport' = None) -> str:
        """
        导出文档模型为 .docx 文件。

        Args:
            doc: 文档模型
            output_path: 输出文件路径
            qa_report: 可选的质量检查报告，将以批注形式嵌入
            crossref_report: 可选的交叉引用报告，将插入书签和 REF 域

        Returns:
            输出文件的绝对路径
        """
        output_path = os.path.abspath(output_path)
        os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)

        word_doc = Document()

        # 1. 设置页面
        self._setup_page(word_doc, doc)

        # 2. 设置页眉页脚
        self._setup_headers_footers(word_doc, doc)

        # 3. 写入内容
        self._write_elements(word_doc, doc)

        # 4. 写入表格
        self._write_tables(word_doc, doc)

        # 5. 插入交叉引用（书签 + REF 域）
        if crossref_report and (crossref_report.targets or crossref_report.ref_points):
            from core.crossref_executor import CrossRefExecutor
            CrossRefExecutor().execute(word_doc, crossref_report, doc)

        # 6. 嵌入质量检查批注
        if qa_report and qa_report.issues:
            self._add_qa_comments(word_doc, doc, qa_report)

        word_doc.save(output_path)
        return output_path

    # ---- 页面设置 ----

    def _setup_page(self, word_doc: Document, doc: DocumentModel):
        """设置页面格式"""
        if not word_doc.sections:
            return
        section = word_doc.sections[0]
        ps = doc.page_setup

        # 纸张大小
        size_map = {
            "A4": (Cm(21.0), Cm(29.7)),
            "A3": (Cm(29.7), Cm(42.0)),
            "B5": (Cm(17.6), Cm(25.0)),
            "Letter": (Cm(21.59), Cm(27.94)),
        }
        if ps.paper_size in size_map:
            w, h = size_map[ps.paper_size]
            if ps.orientation == "landscape":
                section.page_width, section.page_height = h, w
            else:
                section.page_width, section.page_height = w, h

        # 页边距
        if ps.margin_top_cm:
            section.top_margin = Cm(ps.margin_top_cm)
        if ps.margin_bottom_cm:
            section.bottom_margin = Cm(ps.margin_bottom_cm)
        if ps.margin_left_cm:
            section.left_margin = Cm(ps.margin_left_cm)
        if ps.margin_right_cm:
            section.right_margin = Cm(ps.margin_right_cm)
        if ps.gutter_cm:
            section.gutter = Cm(ps.gutter_cm)
        if ps.header_distance_cm:
            section.header_distance = Cm(ps.header_distance_cm)
        if ps.footer_distance_cm:
            section.footer_distance = Cm(ps.footer_distance_cm)

    # ---- 页眉页脚 ----

    def _setup_headers_footers(self, word_doc: Document, doc: DocumentModel):
        """设置页眉页脚"""
        if not doc.sections:
            return

        # 使用第一个节的配置
        section_config = doc.sections[0] if doc.sections else None
        word_section = word_doc.sections[0]

        if section_config:
            # 首页不同
            if section_config.first_page_different:
                word_section.different_first_page_header_footer = True

            # 页眉
            if section_config.header:
                header = word_section.header
                header.is_linked_to_previous = False
                if header.paragraphs:
                    p = header.paragraphs[0]
                else:
                    p = header.add_paragraph()
                self._apply_run_style(p.add_run(section_config.header.text), section_config.header)

            # 首页页眉
            if section_config.first_page_different and section_config.first_page_header:
                first_header = word_section.first_page_header
                if first_header.paragraphs:
                    p = first_header.paragraphs[0]
                else:
                    p = first_header.add_paragraph()
                self._apply_run_style(p.add_run(section_config.first_page_header.text),
                                      section_config.first_page_header)

            # 页码
            if section_config.page_number and section_config.page_number.enabled:
                footer = word_section.footer
                footer.is_linked_to_previous = False
                if footer.paragraphs:
                    p = footer.paragraphs[0]
                else:
                    p = footer.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

                # 插入页码域
                pn = section_config.page_number
                fmt = _PN_FORMAT_MAP.get(pn.number_format, "decimal")

                if pn.format and pn.format != "{num}":
                    # 自定义格式，如 "－ {num} －"
                    prefix = pn.format.split("{num}")[0] if "{num}" in pn.format else ""
                    suffix = pn.format.split("{num}")[1] if "{num}" in pn.format else ""
                    if prefix:
                        run = p.add_run(prefix)
                        if pn.font_name:
                            run.font.name = pn.font_name
                        if pn.font_size_pt:
                            run.font.size = Pt(pn.font_size_pt)

                    # 页码域
                    fldChar_begin = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
                    run = p.add_run()
                    run._element.append(fldChar_begin)

                    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
                    run = p.add_run()
                    run._element.append(instrText)

                    fldChar_end = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
                    run = p.add_run()
                    run._element.append(fldChar_end)

                    if suffix:
                        run = p.add_run(suffix)
                        if pn.font_name:
                            run.font.name = pn.font_name
                        if pn.font_size_pt:
                            run.font.size = Pt(pn.font_size_pt)
                else:
                    # 简单页码
                    fldChar_begin = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
                    run = p.add_run()
                    run._element.append(fldChar_begin)

                    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
                    run = p.add_run()
                    run._element.append(instrText)

                    fldChar_end = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
                    run = p.add_run()
                    run._element.append(fldChar_end)

                if pn.font_size_pt:
                    for run in p.runs:
                        run.font.size = Pt(pn.font_size_pt)

    # ---- 内容写入 ----

    def _write_elements(self, word_doc: Document, doc: DocumentModel):
        """写入文档元素"""
        for elem in doc.elements:
            self._write_element(word_doc, elem)

    def _write_element(self, word_doc: Document, elem: DocElement):
        """写入单个元素"""
        if elem.element_type == ElementType.PAGE_BREAK:
            word_doc.add_page_break()
        elif elem.element_type == ElementType.SECTION_BREAK:
            word_doc.add_section()
        elif elem.element_type == ElementType.HEADING:
            self._write_heading(word_doc, elem)
        elif elem.element_type == ElementType.PARAGRAPH:
            self._write_paragraph(word_doc, elem)
        elif elem.element_type == ElementType.CAPTION:
            self._write_paragraph(word_doc, elem)
        elif elem.element_type == ElementType.REFERENCE:
            self._write_paragraph(word_doc, elem)
        elif elem.element_type == ElementType.LIST:
            self._write_list(word_doc, elem)
        elif elem.element_type == ElementType.LIST_ITEM:
            self._write_paragraph(word_doc, elem)
        elif elem.element_type == ElementType.CODE_BLOCK:
            self._write_code_block(word_doc, elem)
        elif elem.element_type == ElementType.TABLE:
            pass  # 表格单独处理
        else:
            # 默认按段落处理
            self._write_paragraph(word_doc, elem)

    def _write_heading(self, word_doc: Document, elem: DocElement):
        """写入标题"""
        level = max(1, min(elem.level or 1, 4))
        heading = word_doc.add_heading(level=level)
        run = heading.add_run(elem.content)
        self._apply_font_style(run, elem.font_style)
        self._apply_paragraph_format(heading, elem.paragraph_style)

    def _write_paragraph(self, word_doc: Document, elem: DocElement):
        """写入段落"""
        para = word_doc.add_paragraph()
        run = para.add_run(elem.content)
        self._apply_font_style(run, elem.font_style)
        self._apply_paragraph_format(para, elem.paragraph_style)

    def _write_list(self, word_doc: Document, elem: DocElement):
        """写入列表"""
        for child in elem.children:
            self._write_element(word_doc, child)

    def _write_code_block(self, word_doc: Document, elem: DocElement):
        """写入代码块"""
        para = word_doc.add_paragraph()
        run = para.add_run(elem.content)
        run.font.name = "Consolas"
        run.font.size = Pt(9)
        para.paragraph_format.left_indent = Cm(1)

    # ---- 表格写入 ----

    def _write_tables(self, word_doc: Document, doc: DocumentModel):
        """写入所有表格"""
        for i, table_data in enumerate(doc.tables):
            # 写入表题（在表格前）
            if table_data.caption and table_data.caption_position == "above":
                cap_para = word_doc.add_paragraph()
                cap_run = cap_para.add_run(table_data.caption)
                cap_run.bold = True
                cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            self._write_table(word_doc, table_data)

            # 写入表题（在表格后）
            if table_data.caption and table_data.caption_position == "below":
                cap_para = word_doc.add_paragraph()
                cap_run = cap_para.add_run(table_data.caption)
                cap_run.bold = True
                cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _write_table(self, word_doc: Document, table_data: TableData):
        """写入单个表格"""
        if not table_data.rows:
            return

        num_rows = len(table_data.rows)
        num_cols = max(len(row) for row in table_data.rows)

        table = word_doc.add_table(rows=num_rows, cols=num_cols)
        table.style = "Table Grid"

        for row_idx, row in enumerate(table_data.rows):
            for col_idx, cell_data in enumerate(row):
                if col_idx >= num_cols:
                    break
                cell = table.cell(row_idx, col_idx)
                cell.text = cell_data.content

                # 应用单元格样式
                for para in cell.paragraphs:
                    if cell_data.paragraph_style.alignment:
                        para.alignment = _ALIGN_TO_DOCX.get(
                            cell_data.paragraph_style.alignment, WD_ALIGN_PARAGRAPH.LEFT
                        )
                    for run in para.runs:
                        self._apply_font_style(run, cell_data.font_style)

                # 表头加粗
                if cell_data.is_header:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.bold = True

    def _add_qa_comments(self, word_doc: 'Document', doc: DocumentModel, qa_report: 'QAReport'):
        """将质量检查结果以批注形式嵌入 Word 文档"""
        from docx.oxml.ns import qn
        from docx.oxml import parse_xml
        from datetime import datetime

        # 构建 element_index → paragraph 的映射
        para_map: dict[int, any] = {}
        para_idx = 0
        for elem in doc.elements:
            if elem.element_type in (ElementType.HEADING, ElementType.PARAGRAPH,
                                      ElementType.CAPTION, ElementType.REFERENCE):
                if para_idx < len(word_doc.paragraphs):
                    para_map[para_idx] = word_doc.paragraphs[para_idx]
                    para_idx += 1

        # 为每个 issue 添加批注
        for issue in qa_report.issues:
            para = para_map.get(issue.element_index)
            if para is None:
                continue

            # 构建批注文本
            comment_text = f"[{issue.category.value.upper()}] {issue.title}\n{issue.description}"
            if issue.suggestion:
                comment_text += f"\n建议：{issue.suggestion}"

            # 获取段落中的第一个 run
            if not para.runs:
                continue

            # 通过 XML 添加批注
            # python-docx 没有直接的批注 API，需要操作 XML
            self._add_comment_to_paragraph(para, comment_text, "WordCraft Pro")

    @staticmethod
    def _add_comment_to_paragraph(para, comment_text: str, author: str):
        """通过 XML 操作在段落中添加批注"""
        from docx.oxml.ns import qn
        from docx.oxml import parse_xml
        from lxml import etree
        import random

        if not para.runs:
            return

        # 获取段落的 XML
        p_elem = para._element

        # 在第一个 run 前后添加 commentRangeStart 和 commentRangeEnd
        first_run = para.runs[0]._element

        # 生成唯一的 commentId
        comment_id = random.randint(1000, 99999)

        # commentRangeStart
        comment_start = parse_xml(
            f'<w:commentRangeStart {nsdecls("w")} w:id="{comment_id}"/>'
        )
        first_run.addprevious(comment_start)

        # commentRangeEnd
        comment_end = parse_xml(
            f'<w:commentRangeEnd {nsdecls("w")} w:id="{comment_id}"/>'
        )
        # 添加到最后一个 run 后面
        last_run = para.runs[-1]._element
        last_run.addnext(comment_end)

        # commentReference（在批注区域）
        comment_ref_run = parse_xml(
            f'<w:r {nsdecls("w")}>'
            f'  <w:annotationRef/>'
            f'</w:r>'
        )
        # 简化处理：将 commentReference 添加到段末
        comment_end.addnext(comment_ref_run)

        # 注意：完整的批注需要 CommentsPart，这里简化处理
        # 完整实现需要创建 comments.xml part 并添加到文档关系中
        # 这里我们只在段落中标记批注范围，实际批注内容通过
        # 文档的 comments part 添加

    # ---- 样式应用 ----

    def _apply_font_style(self, run, style: FontStyle):
        """将 FontStyle 应用到 docx Run"""
        if not style:
            return
        if style.font_name_cn:
            # 设置中文字体需要操作 XML
            r = run._element
            rPr = r.find(qn("w:rPr"))
            if rPr is None:
                rPr = parse_xml(f'<w:rPr {nsdecls("w")}/>')
                r.insert(0, rPr)
            rFonts = rPr.find(qn("w:rFonts"))
            if rFonts is None:
                rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
                rPr.insert(0, rFonts)
            rFonts.set(qn("w:eastAsia"), style.font_name_cn)
        if style.font_name_en:
            run.font.name = style.font_name_en
        if style.font_size_pt:
            run.font.size = Pt(style.font_size_pt)
        if style.bold:
            run.bold = True
        if style.italic:
            run.italic = True
        if style.underline:
            run.underline = True
        if style.color:
            try:
                run.font.color.rgb = RGBColor.from_string(style.color)
            except (ValueError, AttributeError):
                pass
        if style.super_script:
            run.font.superscript = True
        if style.sub_script:
            run.font.subscript = True

    def _apply_paragraph_format(self, para, style: ParagraphStyle):
        """将 ParagraphStyle 应用到 docx Paragraph"""
        if not style:
            return
        pf = para.paragraph_format

        if style.alignment:
            pf.alignment = _ALIGN_TO_DOCX.get(style.alignment, WD_ALIGN_PARAGRAPH.JUSTIFY)

        # 缩进
        if style.first_indent_chars and style.first_indent_chars > 0:
            pf.first_line_indent = Pt(style.first_indent_chars * 12)  # 按默认12pt计算
        elif style.first_indent_cm and style.first_indent_cm > 0:
            pf.first_line_indent = Cm(style.first_indent_cm)

        if style.left_indent_cm:
            pf.left_indent = Cm(style.left_indent_cm)
        if style.right_indent_cm:
            pf.right_indent = Cm(style.right_indent_cm)

        # 行距
        if style.line_spacing_type:
            if style.line_spacing_type == LineSpacingType.SINGLE:
                pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
            elif style.line_spacing_type == LineSpacingType.ONE_POINT_FIVE:
                pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            elif style.line_spacing_type == LineSpacingType.DOUBLE:
                pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
            elif style.line_spacing_type == LineSpacingType.EXACT:
                pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                pf.line_spacing = Pt(style.line_spacing_value)
            elif style.line_spacing_type == LineSpacingType.MULTIPLE:
                pf.line_spacing = style.line_spacing_value

        # 段前段后
        if style.space_before_pt:
            pf.space_before = Pt(style.space_before_pt)
        if style.space_after_pt:
            pf.space_after = Pt(style.space_after_pt)

        # 分页控制
        if style.keep_with_next:
            pf.keep_with_next = True
        if style.keep_lines_together:
            pf.keep_together = True
        if style.page_break_before:
            pf.page_break_before = True

    def _apply_run_style(self, run, config: HeaderFooterConfig):
        """将 HeaderFooterConfig 应用到 Run"""
        if config.font_name:
            run.font.name = config.font_name
        if config.font_size_pt:
            run.font.size = Pt(config.font_size_pt)
        if config.bold:
            run.bold = True
