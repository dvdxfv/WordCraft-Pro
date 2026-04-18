#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
WordCraft Pro - WebView 桌面应用

使用 QWebEngineView 嵌入网页版，通过 QWebChannel 实现 JS <-> Python 双向桥接。
"""

from __future__ import annotations

import json
import logging
import os
import re
import sys
from typing import Optional

from PyQt6.QtCore import (
    QCoreApplication,
    QDir,
    QFile,
    QIODevice,
    QObject,
    QUrl,
    pyqtSlot,
)
from PyQt6.QtGui import QAction, QKeySequence, QShortcut
from PyQt6.QtWebChannel import QWebChannel
from PyQt6.QtWebEngineCore import QWebEnginePage, QWebEngineSettings
from PyQt6.QtWidgets import (
    QApplication,
    QFileDialog,
    QMainWindow,
    QMessageBox,
    QStatusBar,
    QVBoxLayout,
    QWidget,
)

logger = logging.getLogger("wordcraft.webview")

# ---------------------------------------------------------------------------
#  常量
# ---------------------------------------------------------------------------
APP_TITLE = "WordCraft Pro"
APP_VERSION = "0.8.0"
MIN_WIDTH, MIN_HEIGHT = 1280, 860


# ---------------------------------------------------------------------------
#  Bridge：JS -> Python 桥接对象
# ---------------------------------------------------------------------------
class Bridge(QObject):
    """暴露给 JavaScript 的 Python 桥接对象。

    JS 端通过 ``window.bridge.xxx()`` 调用下方方法。
    所有公开方法均使用 ``@pyqtSlot`` 声明参数与返回值类型。
    """

    def __init__(self, parent: QObject | None = None) -> None:
        super().__init__(parent)
        self._last_save_path: str = ""

    # ------------------------------------------------------------------
    #  文件操作
    # ------------------------------------------------------------------

    @pyqtSlot(result=str)
    def openFile(self) -> str:
        """打开文件对话框，读取并解析文件内容，返回 JSON。

        Returns:
            JSON 字符串::

                {"path": "...", "name": "...", "content": "<elements JSON>"}
        """
        path, _ = QFileDialog.getOpenFileName(
            None,
            "打开文件",
            "",
            "支持格式 (*.docx *.txt *.md);;Word 文档 (*.docx);;文本文件 (*.txt *.md);;所有文件 (*)",
        )
        if not path:
            return ""
        try:
            elements = self._parse_file(path)
            result = {
                "path": path,
                "name": os.path.basename(path),
                "content": json.dumps(elements, ensure_ascii=False),
            }
            logger.info("打开文件: %s (%d 个元素)", path, len(elements))
            return json.dumps(result, ensure_ascii=False)
        except Exception as exc:
            logger.error("解析文件失败: %s", exc, exc_info=True)
            return json.dumps({"success": False, "error": str(exc)}, ensure_ascii=False)

    @pyqtSlot(str, result=str)
    def saveFile(self, html_content: str) -> str:
        """将 HTML 内容保存到本地文件。

        Args:
            html_content: 文档区域的 innerHTML。

        Returns:
            JSON 字符串 ``{"success": bool, "path": "..."}``
        """
        default_name = self._last_save_path or "文档.html"
        path, _ = QFileDialog.getSaveFileName(
            None,
            "保存文件",
            default_name,
            "HTML 文件 (*.html *.htm);;所有文件 (*)",
        )
        if not path:
            return json.dumps({"success": False})
        try:
            with open(path, "w", encoding="utf-8") as fh:
                fh.write(html_content)
            self._last_save_path = path
            logger.info("文件已保存: %s", path)
            return json.dumps({"success": True, "path": path}, ensure_ascii=False)
        except Exception as exc:
            logger.error("保存失败: %s", exc, exc_info=True)
            return json.dumps({"success": False, "error": str(exc)}, ensure_ascii=False)

    @pyqtSlot(str, result=str)
    def exportDocx(self, html_content: str) -> str:
        """将 HTML 内容导出为 Word (.docx) 文档。

        Args:
            html_content: 文档区域的 innerHTML。

        Returns:
            JSON 字符串 ``{"success": bool, "path": "...", "error": "..."}``
        """
        path, _ = QFileDialog.getSaveFileName(
            None,
            "导出 Word 文档",
            "",
            "Word 文档 (*.docx)",
        )
        if not path:
            return json.dumps({"success": False})
        try:
            from docx import Document
            from docx.shared import Cm, Pt

            doc = Document()
            # 简单解析 HTML 标签，提取文本与结构
            elements = self._html_to_elements(html_content)
            for elem in elements:
                tag = elem.get("type", "p")
                text = elem.get("text", "").strip()
                if not text:
                    continue
                if tag in ("h1", "h2", "h3"):
                    level = int(tag[1])
                    heading = doc.add_heading(text, level=level)
                    # 设置中文字体
                    for run in heading.runs:
                        run.font.name = "黑体"
                        run._element.rPr.rFonts.set(
                            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia",
                            "黑体",
                        )
                elif tag == "caption":
                    p = doc.add_paragraph(text)
                    p.alignment = 3  # CENTER
                    for run in p.runs:
                        run.font.size = Pt(10.5)
                        run.font.name = "黑体"
                        run._element.rPr.rFonts.set(
                            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia",
                            "黑体",
                        )
                elif tag == "ref":
                    p = doc.add_paragraph(text)
                    p.paragraph_format.first_line_indent = Cm(0)
                    for run in p.runs:
                        run.font.size = Pt(10.5)
                else:
                    p = doc.add_paragraph(text)
                    p.paragraph_format.first_line_indent = Cm(0.74)  # 约两个字符
                    for run in p.runs:
                        run.font.name = "宋体"
                        run.font.size = Pt(12)
                        run._element.rPr.rFonts.set(
                            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia",
                            "宋体",
                        )

            doc.save(path)
            logger.info("Word 文档已导出: %s", path)
            return json.dumps({"success": True, "path": path}, ensure_ascii=False)
        except ImportError:
            logger.warning("python-docx 未安装，无法导出 Word 文档")
            return json.dumps(
                {"success": False, "error": "python-docx 未安装，请执行 pip install python-docx"},
                ensure_ascii=False,
            )
        except Exception as exc:
            logger.error("导出 Word 失败: %s", exc, exc_info=True)
            return json.dumps({"success": False, "error": str(exc)}, ensure_ascii=False)

    # ------------------------------------------------------------------
    #  文件解析
    # ------------------------------------------------------------------

    @pyqtSlot(str, result=str)
    def parseFileContent(self, file_path: str) -> str:
        """解析指定文件的内容，返回结构化元素列表 JSON。

        Args:
            file_path: 文件的绝对路径。

        Returns:
            JSON 字符串 ``{"elements": [{"type": "h1", "text": "..."}, ...]}``
        """
        try:
            elements = self._parse_file(file_path)
            return json.dumps({"elements": elements}, ensure_ascii=False)
        except Exception as exc:
            logger.error("解析文件内容失败: %s", exc, exc_info=True)
            return json.dumps(
                {"elements": [], "error": str(exc)}, ensure_ascii=False
            )

    # ------------------------------------------------------------------
    #  AI 分析
    # ------------------------------------------------------------------

    @pyqtSlot(str, result=str)
    def runAIAnalysis(self, prompt: str) -> str:
        """调用 AI 分析（如果可用）。

        Args:
            prompt: 发送给 LLM 的提示文本。

        Returns:
            JSON 字符串 ``{"success": bool, "result": "...", "error": "..."}``
        """
        try:
            # 尝试导入项目内的 LLM 客户端
            project_root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
            if project_root not in sys.path:
                sys.path.insert(0, project_root)
            from llm.client import ChatMessage, create_llm_client  # type: ignore[import-untyped]

            config_path = os.path.join(project_root, "config.yaml")
            client = create_llm_client(config_path=config_path)
            if client.is_available():
                messages = [ChatMessage(role="user", content=prompt)]
                result = client.chat(messages)
                return json.dumps(
                    {"success": True, "result": result}, ensure_ascii=False
                )
            else:
                return json.dumps(
                    {"success": False, "error": "AI 服务未配置 API Key"},
                    ensure_ascii=False,
                )
        except Exception as exc:
            logger.warning("AI 分析不可用: %s", exc)
            return json.dumps(
                {"success": False, "error": f"AI服务不可用: {exc}"},
                ensure_ascii=False,
            )

    # ------------------------------------------------------------------
    #  系统信息
    # ------------------------------------------------------------------

    @pyqtSlot(result=str)
    def getSystemInfo(self) -> str:
        """返回系统与版本信息。"""
        info = {
            "platform": sys.platform,
            "version": APP_VERSION,
            "python": sys.version.split()[0],
            "qt": QCoreApplication.instance().applicationVersion() or "6.x",
        }
        return json.dumps(info, ensure_ascii=False)

    # ------------------------------------------------------------------
    #  内部方法
    # ------------------------------------------------------------------

    @staticmethod
    def _parse_file(path: str) -> list[dict[str, str]]:
        """解析文件内容，返回元素列表。

        支持格式：
        - ``.docx``：使用 python-docx 提取段落，根据样式名识别标题
        - ``.txt`` / ``.md``：按行读取，``# `` 开头为 h1，``## `` 为 h2，其余为段落

        Returns:
            ``[{"type": "h1", "text": "..."}, {"type": "p", "text": "..."}, ...]``
        """
        ext = os.path.splitext(path)[1].lower()

        if ext == ".docx":
            return Bridge._parse_docx(path)
        elif ext in (".txt", ".md"):
            return Bridge._parse_text(path)
        else:
            raise ValueError(f"不支持的文件格式: {ext}")

    @staticmethod
    def _parse_docx(path: str) -> list[dict[str, str]]:
        """解析 .docx 文件，提取段落、表格、列表等。"""
        try:
            from docx import Document
        except ImportError:
            raise RuntimeError("python-docx 未安装，无法解析 .docx 文件")

        doc = Document(path)
        elements: list[dict[str, str]] = []

        # 按文档顺序遍历 body 中的 element
        from docx.oxml.ns import qn
        for element in doc.element.body:
            tag = element.tag.split('}')[-1] if '}' in element.tag else element.tag

            if tag == 'p':
                # 段落
                from docx.text.paragraph import Paragraph
                para = Paragraph(element, doc)
                text = para.text.strip()
                if not text:
                    continue
                style_name = (para.style.name or "").lower()
                # 标题识别（支持更多样式名）
                if any(k in style_name for k in ("heading 1", "标题 1", "title", "toc 1")):
                    elements.append({"type": "h1", "text": text})
                elif any(k in style_name for k in ("heading 2", "标题 2", "toc 2")):
                    elements.append({"type": "h2", "text": text})
                elif any(k in style_name for k in ("heading 3", "标题 3", "toc 3")):
                    elements.append({"type": "h3", "text": text})
                elif any(k in style_name for k in ("heading 4", "标题 4")):
                    elements.append({"type": "h3", "text": text})  # h4 也用 h3 显示
                elif any(k in style_name for k in ("caption", "题注")):
                    elements.append({"type": "caption", "text": text})
                elif any(k in style_name for k in ("list bullet", "无序列表")):
                    elements.append({"type": "li", "text": "• " + text})
                elif any(k in style_name for k in ("list number", "有序列表")):
                    elements.append({"type": "li", "text": text})
                else:
                    # 检测是否为参考文献（以 [数字] 开头）
                    import re as _re
                    if _re.match(r'^\[\d+\]', text):
                        elements.append({"type": "ref", "text": text})
                    else:
                        elements.append({"type": "p", "text": text})

            elif tag == 'tbl':
                # 表格
                from docx.table import Table
                try:
                    table = Table(element, doc)
                    rows = []
                    for row in table.rows:
                        cells = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
                        rows.append(cells)
                    if rows:
                        # 转为 HTML 表格
                        html_parts = ['<table>']
                        for i, row in enumerate(rows):
                            tag = 'th' if i == 0 else 'td'
                            html_parts.append('<tr>' + ''.join(f'<{tag}>{c}</{tag}>' for c in row) + '</tr>')
                        html_parts.append('</table>')
                        elements.append({"type": "table", "text": ''.join(html_parts)})
                except Exception:
                    pass

        # 如果没有解析到任何内容，回退到简单段落解析
        if not elements:
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    elements.append({"type": "p", "text": text})

        return elements

    @staticmethod
    def _parse_text(path: str) -> list[dict[str, str]]:
        """解析 .txt / .md 文件。"""
        content: str | None = None
        for enc in ("utf-8", "utf-8-sig", "gbk", "gb2312", "latin-1"):
            try:
                with open(path, "r", encoding=enc) as fh:
                    content = fh.read()
                break
            except (UnicodeDecodeError, UnicodeError):
                continue
        if content is None:
            raise RuntimeError(f"无法解码文件: {path}")

        elements: list[dict[str, str]] = []
        lines = content.splitlines()
        i = 0
        while i < len(lines):
            line = lines[i].strip()
            if not line:
                i += 1
                continue
            # Markdown 标题
            if line.startswith("#### "):
                elements.append({"type": "h3", "text": line[5:].strip()})
            elif line.startswith("### "):
                elements.append({"type": "h3", "text": line[4:].strip()})
            elif line.startswith("## "):
                elements.append({"type": "h2", "text": line[3:].strip()})
            elif line.startswith("# "):
                elements.append({"type": "h1", "text": line[2:].strip()})
            # Markdown 列表
            elif re.match(r'^[\-\*]\s+', line):
                elements.append({"type": "li", "text": "• " + re.sub(r'^[\-\*]\s+', '', line)})
            elif re.match(r'^\d+\.\s+', line):
                elements.append({"type": "li", "text": line})
            # Markdown 表格
            elif line.startswith('|') and i + 1 < len(lines) and lines[i + 1].strip().startswith('|---'):
                table_lines = [line]
                i += 1
                # 跳过分隔行
                while i < len(lines) and lines[i].strip().startswith('|'):
                    if not lines[i].strip().startswith('|---'):
                        table_lines.append(lines[i].strip())
                    i += 1
                # 解析表格
                rows = []
                for tl in table_lines:
                    cells = [c.strip() for c in tl.split('|')[1:-1]]
                    if cells:
                        rows.append(cells)
                if rows:
                    html_parts = ['<table>']
                    for ri, row in enumerate(rows):
                        tag = 'th' if ri == 0 else 'td'
                        html_parts.append('<tr>' + ''.join(f'<{tag}>{c}</{tag}>' for c in row) + '</tr>')
                    html_parts.append('</table>')
                    elements.append({"type": "table", "text": ''.join(html_parts)})
                continue
            # 参考文献
            elif re.match(r'^\[[\d]+\]', line):
                elements.append({"type": "ref", "text": line})
            # 题注
            elif re.match(r'^(图|表|Figure|Table)\s*[\d]', line):
                elements.append({"type": "caption", "text": line})
            else:
                elements.append({"type": "p", "text": line})
            i += 1
        return elements

    @staticmethod
    def _html_to_elements(html_content: str) -> list[dict[str, str]]:
        """将 innerHTML 简单解析为元素列表。

        使用正则提取标签和文本内容，用于导出 Word。
        """
        elements: list[dict[str, str]] = []
        # 移除 script/style 标签
        cleaned = re.sub(r"<(script|style)[^>]*>.*?</\1>", "", html_content, flags=re.DOTALL | re.IGNORECASE)
        # 匹配块级元素
        pattern = re.compile(
            r"<(h[1-6]|p|div|caption|li)[^>]*>(.*?)</\1>|(<table>.*?</table>)",
            re.DOTALL | re.IGNORECASE,
        )
        for match in pattern.finditer(cleaned):
            tag = match.group(1)
            html = match.group(3)  # table 的完整 HTML
            if html:
                elements.append({"type": "table", "text": html})
                continue
            text = re.sub(r"<[^>]+>", "", match.group(2)).strip()
            if not text:
                continue
            if tag in ("h1", "h2", "h3", "h4", "h5", "h6"):
                elements.append({"type": tag, "text": text})
            elif tag == "caption":
                elements.append({"type": "caption", "text": text})
            elif tag == "li":
                elements.append({"type": "li", "text": text})
            else:
                elements.append({"type": "p", "text": text})
        return elements


# ---------------------------------------------------------------------------
#  WebView 主窗口
# ---------------------------------------------------------------------------
class WebViewApp(QMainWindow):
    """基于 QWebEngineView 的 WordCraft Pro 桌面主窗口。"""

    def __init__(self) -> None:
        super().__init__()
        self._bridge = Bridge(self)
        self._init_ui()
        self._init_menu()
        self._init_shortcuts()
        self._init_statusbar()
        self._load_web_page()

    # ------------------------------------------------------------------
    #  UI 初始化
    # ------------------------------------------------------------------

    def _init_ui(self) -> None:
        """初始化窗口基本属性和中央控件。"""
        self.setWindowTitle(APP_TITLE)
        self.setMinimumSize(MIN_WIDTH, MIN_HEIGHT)
        self.resize(1440, 900)

        # 中央容器
        central = QWidget()
        self.setCentralWidget(central)
        layout = QVBoxLayout(central)
        layout.setContentsMargins(0, 0, 0, 0)
        layout.setSpacing(0)

        # QWebEngineView
        from PyQt6.QtWebEngineWidgets import QWebEngineView

        self._web_view = QWebEngineView()

        # 启用开发者工具
        page = self._web_view.page()
        page.settings().setAttribute(
            QWebEngineSettings.WebAttribute.LocalContentCanAccessRemoteUrls, True
        )
        page.settings().setAttribute(
            QWebEngineSettings.WebAttribute.JavascriptEnabled, True
        )
        page.settings().setAttribute(
            QWebEngineSettings.WebAttribute.DevToolsEnabled, True
        )

        # 设置 QWebChannel
        channel = QWebChannel(page)
        channel.registerObject("bridge", self._bridge)
        page.setWebChannel(channel)

        layout.addWidget(self._web_view)

        # 页面加载完成后注入桥接 JS
        self._web_view.page().loadFinished.connect(self._on_page_loaded)

    def _init_menu(self) -> None:
        """创建菜单栏。"""
        menubar = self.menuBar()

        # ---- 文件菜单 ----
        file_menu = menubar.addMenu("文件(&F)")

        act_open = QAction("打开(&O)", self)
        act_open.setShortcut(QKeySequence("Ctrl+O"))
        act_open.triggered.connect(self._on_action_open)
        file_menu.addAction(act_open)

        act_save = QAction("保存(&S)", self)
        act_save.setShortcut(QKeySequence("Ctrl+S"))
        act_save.triggered.connect(self._on_action_save)
        file_menu.addAction(act_save)

        act_export = QAction("导出 Word(&E)", self)
        act_export.setShortcut(QKeySequence("Ctrl+E"))
        act_export.triggered.connect(self._on_action_export)
        file_menu.addAction(act_export)

        file_menu.addSeparator()

        act_exit = QAction("退出(&Q)", self)
        act_exit.setShortcut(QKeySequence("Ctrl+Q"))
        act_exit.triggered.connect(self.close)
        file_menu.addAction(act_exit)

        # ---- 视图菜单 ----
        view_menu = menubar.addMenu("视图(&V)")

        act_refresh = QAction("刷新(&R)", self)
        act_refresh.setShortcut(QKeySequence("F5"))
        act_refresh.triggered.connect(self._on_action_refresh)
        view_menu.addAction(act_refresh)

        act_devtools = QAction("开发者工具(&D)", self)
        act_devtools.setShortcut(QKeySequence("F12"))
        act_devtools.triggered.connect(self._on_action_devtools)
        view_menu.addAction(act_devtools)

        # ---- 帮助菜单 ----
        help_menu = menubar.addMenu("帮助(&H)")

        act_about = QAction("关于(&A)", self)
        act_about.triggered.connect(self._on_action_about)
        help_menu.addAction(act_about)

    def _init_shortcuts(self) -> None:
        """注册额外快捷键（防止被 WebView 拦截）。"""
        # F5 刷新
        refresh = QShortcut(QKeySequence("F5"), self)
        refresh.setContext(Qt.ShortcutContext.ApplicationShortcut)  # type: ignore[attr-defined]
        refresh.activated.connect(self._on_action_refresh)

        # F12 开发者工具
        devtools = QShortcut(QKeySequence("F12"), self)
        devtools.setContext(Qt.ShortcutContext.ApplicationShortcut)  # type: ignore[attr-defined]
        devtools.activated.connect(self._on_action_devtools)

    def _init_statusbar(self) -> None:
        """初始化状态栏。"""
        self._statusbar = QStatusBar()
        self.setStatusBar(self._statusbar)
        self._status_label = self._statusbar.showMessage("就绪")

    # ------------------------------------------------------------------
    #  加载网页
    # ------------------------------------------------------------------

    def _load_web_page(self) -> None:
        """读取本地 index.html 并通过 setHtml 加载。"""
        web_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "web")
        index_path = os.path.join(web_dir, "index.html")

        qfile = QFile(index_path)
        if not qfile.open(QIODevice.OpenModeFlag.ReadOnly | QIODevice.OpenModeFlag.Text):
            QMessageBox.critical(
                self,
                "加载失败",
                f"无法打开网页文件:\n{index_path}\n\n{qfile.errorString()}",
            )
            return

        html_content = bytes(qfile.readAll()).decode("utf-8")
        qfile.close()

        # 将 baseUrl 设置为 web 目录，以便正确加载相对资源
        base_url = QUrl.fromLocalFile(os.path.abspath(web_dir) + "/")
        self._web_view.setHtml(html_content, baseUrl=base_url)
        logger.info("正在加载网页: %s", index_path)

    # ------------------------------------------------------------------
    #  页面加载完成 -> 注入桥接 JS
    # ------------------------------------------------------------------

    def _on_page_loaded(self, ok: bool) -> None:
        """页面加载完成后注入 QWebChannel 初始化脚本和桥接覆盖函数。"""
        if not ok:
            logger.error("网页加载失败")
            self._statusbar.showMessage("网页加载失败")
            return

        logger.info("网页加载完成，正在注入桥接脚本")

        # 1) 注入 QWebChannel 初始化（让 JS 端能拿到 bridge 对象）
        channel_init_js = """
        (function() {
            'use strict';
            if (typeof QWebChannel === 'undefined') {
                // 内联最小化 QWebChannel polyfill（仅用于 new QWebChannel(qt.webChannelTransport, cb)）
                // 完整版由 Qt 在 baseUrl 下提供；如果不可用则使用简化版
                console.warn('QWebChannel not loaded, trying fallback...');
            }
            new QWebChannel(qt.webChannelTransport, function(channel) {
                window.bridge = channel.objects.bridge;
                console.log('QWebChannel connected, bridge ready.');
                // 通知页面桥接已就绪
                if (typeof window.onBridgeReady === 'function') {
                    window.onBridgeReady();
                }
            });
        })();
        """
        self._web_view.page().runJavaScript(channel_init_js)

        # 2) 注入桥接覆盖函数
        self._inject_bridge_js()

    def _inject_bridge_js(self) -> None:
        """注入 JS 代码，覆盖网页版中的文件操作函数，连接 Python 桥接。"""
        js = r'''
// ================================================================
//  WordCraft Pro - Python Bridge 注入脚本
//  在页面加载完成后由 WebViewApp._inject_bridge_js() 注入
// ================================================================
(function() {
    'use strict';

    // 等待 bridge 对象可用
    function waitForBridge(callback) {
        if (window.bridge) {
            callback();
        } else {
            var checkInterval = setInterval(function() {
                if (window.bridge) {
                    clearInterval(checkInterval);
                    callback();
                }
            }, 100);
            // 超时 5 秒
            setTimeout(function() {
                clearInterval(checkInterval);
                if (!window.bridge) {
                    console.warn('Bridge injection timed out');
                }
            }, 5000);
        }
    }

    waitForBridge(function() {

        // ---- 覆盖 openFile ----
        window.openFile = async function() {
            try {
                var result = await window.bridge.openFile();
                if (result) {
                    var data = JSON.parse(result);
                    if (data.error) {
                        window.showToast('打开失败: ' + data.error);
                        return;
                    }
                    if (data.content) {
                        window.renderParsedContent(data.name, data.content);
                        window.showToast('已打开: ' + data.name);
                    }
                }
            } catch (e) {
                console.error('openFile bridge error:', e);
                window.showToast('打开文件时出错');
            }
        };

        // ---- 覆盖 saveFile ----
        window.saveFile = function() {
            var content = '';
            var docPage = document.getElementById('docPage');
            if (docPage) {
                content = docPage.innerHTML;
            }
            window.bridge.saveFile(content).then(function(result) {
                var data = JSON.parse(result);
                if (data.success) {
                    window.showToast('已保存: ' + data.path);
                } else {
                    window.showToast('保存失败' + (data.error ? ': ' + data.error : ''));
                }
            }).catch(function(e) {
                console.error('saveFile bridge error:', e);
                window.showToast('保存时出错');
            });
        };

        // ---- 覆盖 exportDoc ----
        window.exportDoc = function() {
            var content = '';
            var docPage = document.getElementById('docPage');
            if (docPage) {
                content = docPage.innerHTML;
            }
            window.bridge.exportDocx(content).then(function(result) {
                var data = JSON.parse(result);
                if (data.success) {
                    window.showToast('已导出: ' + data.path);
                } else {
                    window.showToast('导出失败: ' + (data.error || ''));
                }
            }).catch(function(e) {
                console.error('exportDoc bridge error:', e);
                window.showToast('导出时出错');
            });
        };

        // ---- 添加 renderParsedContent ----
        // 将解析后的文件内容注入到网页版的数据结构中并打开标签页
        window.renderParsedContent = function(fileName, contentJson) {
            try {
                var elements = JSON.parse(contentJson);
                // 确保 elements 是数组
                if (!Array.isArray(elements)) {
                    console.warn('renderParsedContent: elements is not an array');
                    return;
                }

                // 检查是否已存在同名文件
                var existIdx = -1;
                for (var i = 0; i < window.filesData.length; i++) {
                    if (window.filesData[i].name === fileName) {
                        existIdx = i;
                        break;
                    }
                }

                // 确定文件类型
                var ext = fileName.split('.').pop().toLowerCase();
                var fileType = (['docx', 'pdf', 'xlsx', 'txt', 'md'].indexOf(ext) >= 0) ? ext : 'txt';

                if (existIdx >= 0) {
                    // 更新已有文件内容
                    window.docContents[fileName] = elements;
                    // 切换到该标签
                    var tabIdx = -1;
                    for (var j = 0; j < window.openTabs.length; j++) {
                        if (window.openTabs[j].name === fileName) {
                            tabIdx = j;
                            break;
                        }
                    }
                    if (tabIdx >= 0) {
                        window.switchTab(tabIdx);
                    } else {
                        window.openFileTab(existIdx);
                    }
                } else {
                    // 添加新文件
                    window.filesData.push({
                        name: fileName,
                        size: _estimateSize(elements),
                        type: fileType,
                        modified: false
                    });
                    window.docContents[fileName] = elements;
                    window.openFileTab(window.filesData.length - 1);
                }

                // 刷新文件树
                if (typeof window.renderFileTree === 'function') {
                    window.renderFileTree();
                }
            } catch (e) {
                console.error('renderParsedContent error:', e);
                window.showToast('加载文件内容失败');
            }
        };

        // 辅助：估算文件大小
        function _estimateSize(elements) {
            var totalChars = 0;
            for (var i = 0; i < elements.length; i++) {
                totalChars += (elements[i].text || '').length;
            }
            var bytes = totalChars * 3; // 粗略估算 UTF-8 字节数
            if (bytes < 1024) return bytes + ' B';
            if (bytes < 1024 * 1024) return (bytes / 1024).toFixed(1) + ' KB';
            return (bytes / (1024 * 1024)).toFixed(1) + ' MB';
        }

        console.log('WordCraft Pro bridge injected successfully');
    });
})();
'''
        self._web_view.page().runJavaScript(js)

    # ------------------------------------------------------------------
    #  菜单动作
    # ------------------------------------------------------------------

    def _on_action_open(self) -> None:
        """文件 -> 打开"""
        self._web_view.page().runJavaScript("if(typeof openFile==='function') openFile();")

    def _on_action_save(self) -> None:
        """文件 -> 保存"""
        self._web_view.page().runJavaScript("if(typeof saveFile==='function') saveFile();")

    def _on_action_export(self) -> None:
        """文件 -> 导出 Word"""
        self._web_view.page().runJavaScript("if(typeof exportDoc==='function') exportDoc();")

    def _on_action_refresh(self) -> None:
        """视图 -> 刷新"""
        self._load_web_page()

    def _on_action_devtools(self) -> None:
        """视图 -> 开发者工具"""
        page = self._web_view.page()
        page.setDevToolsPage(page)  # 在同一窗口打开
        # 备选方案：使用 inspectElement
        try:
            page.runJavaScript("1", 0)  # 触发 DevTools
            # Qt6 没有直接的 API，通过触发 inspected 页面
            self._web_view.page().triggerAction(
                QWebEnginePage.WebAction.InspectElement
            )
        except Exception as exc:
            logger.warning("无法打开开发者工具: %s", exc)

    def _on_action_about(self) -> None:
        """帮助 -> 关于"""
        QMessageBox.about(
            self,
            "关于 WordCraft Pro",
            f"<h3>WordCraft Pro</h3>"
            f"<p>版本 {APP_VERSION}</p>"
            f"<p>智能 Word 排版桌面应用</p>"
            f"<p>基于 PyQt6 + QWebEngineView 构建</p>"
            f"<p>Python {sys.version.split()[0]} | Qt 6</p>",
        )

    # ------------------------------------------------------------------
    #  状态栏更新（供 JS 调用）
    # ------------------------------------------------------------------

    def update_status(self, message: str) -> None:
        """更新状态栏文本。"""
        self._statusbar.showMessage(message)


# ---------------------------------------------------------------------------
#  Qt.ShortcutContext 兼容
# ---------------------------------------------------------------------------
# PyQt6 中 Qt.ShortcutContext 需要从 QtCore 导入
from PyQt6.QtCore import Qt  # noqa: E402  (already imported at top, re-import for clarity)


# ---------------------------------------------------------------------------
#  入口
# ---------------------------------------------------------------------------
def main() -> None:
    """应用启动入口。"""
    # 高 DPI 支持
    os.environ.setdefault("QT_ENABLE_HIGHDPI_SCALING", "1")

    app = QApplication(sys.argv)
    app.setStyle("Fusion")
    app.setApplicationName(APP_TITLE)
    app.setApplicationVersion(APP_VERSION)

    # 日志配置
    logging.basicConfig(
        level=logging.INFO,
        format="%(asctime)s [%(levelname)s] %(name)s: %(message)s",
        datefmt="%H:%M:%S",
    )

    # 预检测 WebEngine 依赖
    try:
        from PyQt6.QtWebEngineWidgets import QWebEngineView as _test
        del _test
    except Exception as exc:
        logging.error("WebEngine import failed: %s", exc)
        from PyQt6.QtWidgets import QMessageBox
        QMessageBox.critical(
            None,
            "缺少依赖",
            f"未找到 PyQt6-WebEngine，请执行:\n\n"
            f"  pip install PyQt6-WebEngine\n\n"
            f"错误详情: {exc}\n\n"
            f"后重试。",
        )
        sys.exit(1)

    window = WebViewApp()
    window.show()

    logger.info("WordCraft Pro WebView 启动完成")
    sys.exit(app.exec())


if __name__ == "__main__":
    main()
